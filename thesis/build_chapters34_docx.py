#!/usr/bin/env python3
"""Expand THESIS_DRAFT_FCU_v1.docx chapters 3–4 with Lab/Phase5 sections and figures."""

from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches
from docx.text.paragraph import Paragraph

THESIS = Path(__file__).resolve().parent
ROOT = THESIS.parent
DOC_PATH = THESIS / "THESIS_DRAFT_FCU_v1.docx"

FIG = {
    "4.1": ROOT
    / "runtime/outputs/phase3_rtx_pra_comparison_fixed_tcp_repeatability_v1/rtx_amplitude_max_vs_distance.png",
    "4.2": THESIS / "figures/fig4_2_rtx_early_energy_vs_distance.png",
    "4.3": ROOT
    / "runtime/outputs/phase3_rtx_pra_comparison_fixed_tcp_repeatability_v1/pra_features_vs_distance.png",
    "4.4": THESIS / "figures/fig4_4_material_early_energy_0p5m.png",
    "4.5": ROOT / "runtime/outputs/lab_dynamic_smoke_v1/lab_target_trajectory_xy.png",
    "4.6": ROOT / "runtime/outputs/lab_dynamic_smoke_v1/lab_obs_vs_gt_distance.png",
    "4.7": ROOT / "runtime/outputs/lab_sl_distance_v1/sl_sim_to_lab_pred_vs_gt.png",
    "4.8": ROOT / "runtime/outputs/lab_sl_distance_v1/sl_sim_to_lab_trajectory.png",
}

ABSTRACT_ZH_ADD = (
    "此外，本研究將同一 Geometry Passport 與 RTX 特徵工廠延伸至 Isaac Lab，於固定 TCP 與正弦移動目標之動態場景中連續擷取 128 步觀測；"
    "並以 Sim 靜態標定資料訓練之線性回歸器，在 Lab 動態測試集上達趨勢級距離追蹤（Pearson r≈0.47，MAE≈0.41 m），示範 Sim→Lab 訓練前管線之可行性。"
    "進一步地，本研究以 Isaac Lab DirectRLEnv 搭配 RSL-RL PPO 於模擬器內閉環訓練距離估計策略；"
    "修正 GMO 擷取管線後可達趨勢級追蹤，惟聲學特徵對距離之變異解釋量仍有限，故 in-sim RL 僅作訓練迴路可行性示範，不宣稱優於監督基線。"
    "上述延伸結果均為趨勢級代理指標，非部署級測距精度。"
)

ABSTRACT_EN_ADD = (
    " The same geometry passport and RTX feature factory were extended in Isaac Lab under a fixed-TCP, sinusoidally moving-target dynamic scenario (128 steps). "
    "A linear regressor trained on static Isaac Sim captures achieved trend-level distance tracking on the Lab dynamic hold-out set (Pearson r≈0.47, MAE≈0.41 m). "
    "An in-simulation closed-loop extension with DirectRLEnv and RSL-RL PPO demonstrated a feasible training loop, without claiming superiority over the supervised baseline. "
    "These extensions are trend-level proxies, not deployment-grade ranging."
)

# Each item: ("Header2"|"Content"|"Image"|"TableCaption", ...)
CH34: list[tuple[Any, ...]] = [
    ("Header2", "3.1 研究架構與流程"),
    (
        "Content",
        "本研究採三階段可審計管線：Phase 1 建立 Geometry Passport（房間、UR10 掛載、距離 waypoints）；"
        "Phase 2 建立 Material Passport（非視覺材質條件 A/B/C）；Phase 3 於 Isaac Sim host standalone 執行 fixed_tcp_moving_target 正式實驗，"
        "擷取 RTX GMO signal-way 並與 PyRoomAcoustics 趨勢對照。Phase 4 將同一 Passport 與特徵工廠延伸至 Isaac Lab 動態觀測與 Sim→Lab 監督學習；"
        "Phase 5 再示範 DirectRLEnv + RSL-RL 之 in-sim 閉環訓練（方法延伸，非主貢獻）。",
    ),
    ("Header2", "3.2 實驗平台與軟體環境"),
    (
        "Content",
        "模擬平台為 Isaac Sim 6.0 host standalone（6.0.0-rc.59），透過 scripts/run_host_python.sh 與 env_host_isolated.sh 隔離執行。"
        "機器人採官方 Universal Robots UR10 USD 資產；RTX Acoustic 掛載於 ee_link。Isaac Lab（v3.0.0-beta2）安裝於 IsaacLab/，"
        "以 _isaac_sim 符號連結共用同一 Sim 安裝，避免重複解析 GMO。",
    ),
    ("Header2", "3.3 實驗設計：fixed_tcp_moving_target"),
    (
        "Content",
        "考量 UR10 有效 reach 約 1.3 m，本研究採固定 TCP（半徑 xy≈0.816 m、高度 0.65 m）、僅移動目標之設計，"
        "使 sensor→target 距離定義一致且 ee_link 運動量為 0 m（max_ee_motion_m=0），排除手臂運動引入之額外變異。"
        "距離 waypoints 為 0.5–3.0 m（6 點）；目標板 8 cm×8 cm×2 cm。已廢棄之 IK 移動手臂方案不納入本論文。",
    ),
    ("Header2", "3.4 Geometry Passport 與 Material Passport"),
    (
        "Content",
        "Geometry Passport v1.0 定義房間 4.5 m×3.0 m×2.8 m、感測器掛載與目標軌跡；Material Passport v1.0 定義 NonVisualMaterial 條件 A（低吸收）、"
        "B（中吸收，主實驗）、C（高吸收）。兩份 Passport 以 JSON/YAML 可版本化，支撐 30/30 可重複性協定與 replication package。",
    ),
    ("Header2", "3.5 RTX 擷取、特徵萃取與 PRA 對照"),
    (
        "Content",
        "RTX 數據透過 Replicator Writer 擷取，並以 timeline.play() 驅動模擬步進。rtx_acoustic_factory.py 驗證 GMO 結構並解析 signal-way，"
        "萃取 primary_sgw_early_energy（前 25% samples 能量和）等特徵。PyRoomAcoustics 以同幾何對齊之 image-source 模型提供 RIR、RT60 與 early energy 趨勢參考。",
    ),
    ("Header2", "3.9 Isaac Lab 動態環境與 Sim 管線延伸"),
    (
        "Content",
        "Isaac Lab 端不重新實作 RTX 感測解析，而是直接呼叫 Phase 3 已驗證之 geometry_passport_v1.py、rtx_acoustic_factory.py 與 rtx_material_passport_v1.py。"
        "動態環境 Ur10RtxAcousticDynamicEnv（lab/ur10_rtx_acoustic_env.py）維持與 Sim 一致之場景：官方 UR10、關節鎖定、材質 B、感測器掛載於 ee_link/official_rtx_acoustic。",
    ),
    (
        "Content",
        "目標沿感測器 +X 軸正弦運動：d(t)=1.5+0.5·sin(2π·step/64) m，step=0…127。"
        "觀測含 primary_sgw_early_energy、primary_sgw_peak、gmo_valid 與 GT 距離 target_distance_m_gt；動作為零動作（fixed TCP）。"
        "每 4 步擷取一次 GMO（decimation=4）。啟動使用 AppLauncher 與 isaacsim.exp.base.python.kit experience。",
    ),
    ("Header2", "3.10 In-sim 強化學習環境（DirectRLEnv + RSL-RL）"),
    (
        "Content",
        "§4.6 與離線 REINFORCE smoke 在已記錄 GMO 軌跡上更新策略，環境不隨 policy 重算 RTX。"
        "為驗證 Isaac Lab + RSL-RL 閉環，本節採 DirectRLEnv（Isaac-Ur10RtxAcousticDistance-Direct-v0）："
        "每一步 policy 輸出距離預測，模擬器推進目標軌跡、觸發 GMO 並回傳 reward。",
    ),
    (
        "Content",
        "觀測 6 維：early_energy×1e-4、peak×1e-2、gmo_valid、sensor_x,y,z；動作 1 維映射至 [1.0,2.0] m。"
        "Episode 32 policy steps；GMO 每 2 steps 擷取。獎勵基線 r=−|pred−gt|；v5 加入能量先驗與趨勢一致獎勵（塑形獎勵）。"
        "PPO 以 RSL-RL 實作，v3/v4 各 200 iter，v5 為 500 iter + empirical_normalization。",
    ),
    ("Header1", "第四章、實證結果與分析"),
    ("Header2", "4.1 實驗可行性與可重複性"),
    (
        "Content",
        "正式實驗 30/30 PASS（6 距離點×5 repeats）；gmo_valid_rate=1.0；鎖定 TCP 下 max_ee_motion_m=0。"
        "跨 repeat 變異係數（CV）於 amplitude_max 與 early_energy 皆極低，2.0–3.0 m 區間 CV 近乎 0，顯示管線具高可重複性。",
    ),
    ("Header2", "4.2 距離趨勢分析（材質 B）"),
    (
        "Content",
        "表4.1 彙整材質 B、5 repeats 平均之 RTX 與 PRA 特徵。amplitude_max 於 1.0 m 後飽和，不適合作為距離單調 proxy；"
        "primary_sgw_early_energy 與距離呈負相關（Spearman ρ≈−0.66）。RTX 與 PRA early energy 趨勢一致（ρ≈+0.66），但振幅尺度不同，屬趨勢級對照。",
    ),
    ("Image", "4.1", "圖4.1  RTX amplitude_max 與目標距離之關係（材質 B，5 repeats 平均）"),
    ("Image", "4.2", "圖4.2  RTX primary_sgw_early_energy 與目標距離之關係（材質 B）"),
    ("Image", "4.3", "圖4.3  PyRoomAcoustics 特徵與目標距離之趨勢對照"),
    ("Header2", "4.3 材質敏感度分析（A/B/C）"),
    (
        "Content",
        "跨材質比較顯示 primary_sgw_peak 於 0.5 m 與 3.0 m 幾乎無差（飽和）；primary_sgw_early_energy 於 0.5 m 可區分條件 C（≈190.6）與 A/B（≈165.4）。"
        "PRA RT60 對材質吸收差異更敏感，支持「RTX early energy 近距離具材質線索、peak 不敏感」之結論。",
    ),
    ("Image", "4.4", "圖4.4  材質條件 A/B/C 於 0.5 m 之 early_energy 比較"),
    ("Header2", "4.5 Isaac Lab 動態觀測原型"),
    (
        "Content",
        "動態 smoke（128 steps，GMO decimation=4）PASS。共 128 列逐步觀測，27 步有效 GMO（擷取率 84%），gmo_valid_rate=1.0。"
        "鎖定 TCP 下 sensor 位置不變（max_sensor_position_motion_m=0）。primary_sgw_early_energy 與 GT 距離 Pearson ρ=−0.475（n=27），"
        "方向與 Sim 靜態掃描一致，惟樣本數與距離範圍（1–2 m）有限，不宣稱統計顯著性。",
    ),
    ("Image", "4.5", "圖4.5  動態場景 GT 距離隨 step 變化（正弦軌跡，1.0–2.0 m）"),
    ("Image", "4.6", "圖4.6  Lab 動態場景 early_energy 與 GT 距離散佈（ρ≈−0.48）"),
    ("Header2", "4.6 監督學習距離估計（Sim→Lab 遷移）"),
    (
        "Content",
        "以 Sim fixed_tcp_repeatability_v1 之 GMO 列（材質 B，n=125）訓練單變量線性回歸（特徵：primary_sgw_early_energy），"
        "於 Lab 動態 27 筆有效 GMO 上測試。Sim→Lab 主結果：MAE=0.41 m、RMSE=0.52 m、Pearson r=0.47。"
        "Lab-only 5-fold CV（n=27）MAE=0.31 m、r=0.27，樣本少僅供參考。",
    ),
    (
        "Content",
        "雙特徵消融（early_energy + peak）使 Sim→Lab MAE 升至 0.44 m，因 Lab 1–2 m 區間 peak 飽和於 ≈5171，與 Sim Phase 3 一致；"
        "故論文主模型採 early_energy 單特徵。此節為趨勢級可行性示範，MAE 0.41 m 不可解讀為部署級測距精度。",
    ),
    ("Image", "4.7", "圖4.7  Sim 訓練→Lab 測試：預測距離 vs GT"),
    ("Image", "4.8", "圖4.8  Sim→Lab 預測與 GT 軌跡對照（依 GT 排序）"),
    ("Header2", "4.7 綜合討論"),
    (
        "Content",
        "（1）Sim 主結果：primary_sgw_early_energy 優於飽和之 amplitude_max（ρ≈−0.66，30/30 PASS）。"
        "（2）Lab 延伸：動態場景趨勢方向一致（ρ≈−0.48），Sim→Lab SL r≈0.47。"
        "（3）RTX×PRA：趨勢級一致，非波形等價。（4）CH201 實機應採 task-level 指標；本研究輸出為協定與資料格式參考。",
    ),
    ("Header2", "4.7.1 In-sim RSL-RL 距離估計實驗"),
    (
        "Content",
        "本節示範 RTX GMO 管線可支撐 in-sim PPO 訓練。初期長訓練出現 pred 常數化，根因為 GMO writer 於 sim.reset() 後失效；"
        "修正要點：set_simulation_app(app_launcher.app)、DirectRLEnv 初始化後 rebind_rtx_gmo_writer()、"
        "capture_rtx_gmo() 採 timeline.play() 模式。修正後 gmo_init_captured=True，obs0 有效。",
    ),
    (
        "Content",
        "表4.7 彙整 hold-out 64 steps 評估。v3 det MAE=0.115 m 但 pred 近常數（std≈0）；v5 塑形獎勵使 pred 具變異（std≈0.48 m）"
        "惟 MAE=0.441 m，未優於 §4.6 SL（MAE=0.41 m，r=0.47）。r(gt,raw_E)≈0 反映 1–2 m 區間 early_energy 變化 <0.5%。"
        "RL 價值在閉環管線驗證，非精度超越監督基線。",
    ),
    ("TableCaption", "表4.7  In-sim RSL-RL hold-out 評估（64 steps，材質 B）"),
    (
        "Table",
        ["Run", "Checkpoint", "模式", "MAE(m)", "r(gt,pred)", "pred mean±std(m)"],
        [
            ["long_v3", "model_199", "det", "0.115", "0.335", "1.576±0.000"],
            ["long_v4", "model_199", "stoch", "0.126", "0.247", "1.534±0.083"],
            ["long_v5", "model_499", "det", "0.441", "0.229", "1.637±0.477"],
            ["long_v5", "model_499", "stoch", "0.457", "0.078", "1.559±0.481"],
        ],
    ),
    ("TableCaption", "表4.8  與 §4.6 監督基線對照"),
    (
        "Table",
        ["方法", "MAE (m)", "Pearson r", "備註"],
        [
            ["Sim→Lab 線性 SL（§4.6）", "0.41", "0.47", "主學習基線"],
            ["in-sim PPO v3 det", "0.115", "0.34", "pred 近常數"],
            ["in-sim PPO v5 det", "0.441", "0.229", "有變異，未優於 SL"],
        ],
    ),
]

CH5_52 = [
    ("Header2", "5.2 研究建議"),
    ("Header2", "5.2.1 已完成之 Lab 階段"),
    ("Content", "Phase 4 動態 smoke（128 steps）與 Sim→Lab 監督學習線性遷移已完成。"),
    ("Header2", "5.2.2 已完成之 in-sim RL 階段"),
    (
        "Content",
        "DirectRLEnv 註冊、GMO writer 生命週期修正、v3/v4 200-iter 與 v5 500-iter 塑形獎勵長訓練及 hold-out eval 已完成；"
        "結果支持閉環可行性，未優於 SL 基線。",
    ),
    ("Header2", "5.2.3 後續建議"),
    (
        "Content",
        "建議後續：（1）補齊論文圖表與 CH201 task-level 實機協定；（2）擴大距離範圍或離軸點位以增加 early_energy 動態；"
        "（3）多 env 平行加速 PPO；（4）以 replication package 支援口試重現。",
    ),
]


def insert_paragraph_after(paragraph: Paragraph, style: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.style = style
    return new_para


def remove_paragraph(paragraph: Paragraph) -> None:
    paragraph._element.getparent().remove(paragraph._element)


def set_paragraph_text(paragraph: Paragraph, text: str, style: str, align_center: bool = False) -> None:
    paragraph.style = style
    if align_center:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in paragraph.runs:
        run._element.getparent().remove(run._element)
    if text:
        paragraph.add_run(text)


def paragraph_after_element(parent, element) -> Paragraph:
    after = OxmlElement("w:p")
    element.addnext(after)
    return Paragraph(after, parent)


def append_content(doc: Document, anchor: Paragraph, items: list[tuple[Any, ...]]) -> Paragraph:
    prev = anchor
    for item in items:
        kind = item[0]
        if kind == "Table":
            _, headers, rows = item
            table = doc.add_table(rows=1 + len(rows), cols=len(headers))
            for j, h in enumerate(headers):
                table.rows[0].cells[j].text = h
            for i, row in enumerate(rows, start=1):
                for j, val in enumerate(row):
                    table.rows[i].cells[j].text = val
            prev._element.addnext(table._element)
            prev = paragraph_after_element(prev._parent, table._element)
            continue

        new_para = insert_paragraph_after(prev, "Content")
        if kind == "Header1":
            set_paragraph_text(new_para, item[1], "Header1")
        elif kind == "Header2":
            set_paragraph_text(new_para, item[1], "Header2")
        elif kind == "Content":
            set_paragraph_text(new_para, item[1], "Content")
        elif kind == "Image":
            _, key, caption = item
            path = FIG[key]
            if not path.exists():
                raise FileNotFoundError(path)
            new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            new_para.add_run().add_picture(str(path), width=Inches(5.2))
            cap_para = insert_paragraph_after(new_para, "Content")
            set_paragraph_text(cap_para, caption, "Content", align_center=True)
            prev = cap_para
            continue
        elif kind == "TableCaption":
            set_paragraph_text(new_para, item[1], "Content", align_center=True)
        prev = new_para
    return prev


def replace_chapter_block(doc: Document, start_title: str, end_title: str, items: list[tuple[Any, ...]]) -> None:
    start_idx = end_idx = None
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t == start_title:
            start_idx = i
        if start_idx is not None and t == end_title:
            end_idx = i
            break
    if start_idx is None or end_idx is None:
        raise RuntimeError(f"Block not found: {start_title} -> {end_title}")

    anchor = doc.paragraphs[start_idx]
    for idx in range(end_idx - 1, start_idx, -1):
        remove_paragraph(doc.paragraphs[idx])
    append_content(doc, anchor, items)


def replace_section_from_header(doc: Document, header_text: str, until_text: str, items: list[tuple[Any, ...]]) -> None:
    start_idx = until_idx = None
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t == header_text:
            start_idx = i
        if start_idx is not None and t == until_text and i > start_idx:
            until_idx = i
            break
    if start_idx is None or until_idx is None:
        raise RuntimeError(f"Section not found: {header_text} until {until_text}")

    anchor = doc.paragraphs[start_idx - 1] if start_idx > 0 else doc.paragraphs[start_idx]
    for idx in range(until_idx - 1, start_idx - 1, -1):
        remove_paragraph(doc.paragraphs[idx])
    append_content(doc, anchor, items)


def update_abstract(doc: Document) -> None:
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.startswith("本研究評估 NVIDIA") and ABSTRACT_ZH_ADD not in t:
            p.text = t + ABSTRACT_ZH_ADD
        if t.startswith("This study evaluates RTX Acoustic") and ABSTRACT_EN_ADD.strip() not in t:
            p.text = t.rstrip() + ABSTRACT_EN_ADD


def main() -> None:
    for key, path in FIG.items():
        if not path.exists():
            raise FileNotFoundError(f"Missing figure {key}: {path}")

    doc = Document(str(DOC_PATH))
    update_abstract(doc)
    replace_chapter_block(doc, "第三章、研究流程與方法", "第五章、結論與建議", CH34)
    replace_section_from_header(doc, "5.2 研究建議", "參考文獻", CH5_52)
    doc.save(str(DOC_PATH))
    print(f"Updated {DOC_PATH}")


if __name__ == "__main__":
    main()