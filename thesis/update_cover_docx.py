#!/usr/bin/env python3
"""Update thesis cover metadata for 電聲碩士學位學程."""

from pathlib import Path

from docx import Document

DOC = Path(__file__).resolve().parent / "THESIS_DRAFT_FCU_v1.docx"

PROGRAM = "電聲碩士學位學程"
TITLE_ZH = (
    "基於 RTX Acoustic 超音波感測之機械手臂閉迴路接近控制"
    "與 Physical AI 狀態判斷"
)
DATE = "中華民國 115 年 7 月"
ADVISOR_LINE = "指導教授：蔡鈺鼎 教授"


def replace_paragraph_text(paragraph, new_text: str) -> None:
    for run in paragraph.runs:
        run._element.getparent().remove(run._element)
    paragraph.add_run(new_text)


def main() -> None:
    doc = Document(str(DOC))
    for p in doc.paragraphs:
        t = p.text.strip()
        if t == "智能製造與工程管理碩士在職學位學程":
            replace_paragraph_text(p, PROGRAM)
        elif t.startswith("UR10 末端 RTX") or t.startswith("基於 RTX Acoustic"):
            replace_paragraph_text(p, TITLE_ZH)
        elif t.startswith("中華民國") and "年" in t:
            replace_paragraph_text(p, DATE)

    # Insert advisor line after date if not present
    has_advisor = any("蔡鈺鼎" in p.text for p in doc.paragraphs)
    if not has_advisor:
        for i, p in enumerate(doc.paragraphs):
            if p.text.strip() == DATE:
                new_p = p._element
                from docx.oxml import OxmlElement
                from docx.text.paragraph import Paragraph

                advisor_xml = OxmlElement("w:p")
                new_p.addnext(advisor_xml)
                advisor_para = Paragraph(advisor_xml, p._parent)
                replace_paragraph_text(advisor_para, ADVISOR_LINE)
                break

    doc.save(str(DOC))
    print(f"Updated cover: {DOC}")


if __name__ == "__main__":
    main()