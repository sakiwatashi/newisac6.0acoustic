#!/usr/bin/env python3
"""Remove duplicate history sections/images introduced by re-running update script."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph

PATHS = [
    Path("/home/lab109/song/isaacsim6.0/thesis/THESIS_DRAFT_FCU_v1.docx"),
    Path("/home/lab109/song/isaacsim6.0/thesis/2THESIS_DRAFT_FCU_v1.docx"),
    Path("/home/lab109/下載/2THESIS_DRAFT_FCU_v1_no_pyroom.docx"),
]

TOC_FIXES = [
    (
        r"2\.5 閉迴路感知、機器人接近與視覺語義操作對照\d+2\.5 感測回授式接近與視覺語義操作對照\d+2\.5 感測回授式接近與視覺語義操作對照\d+",
        "2.5 感測回授式接近與視覺語義操作對照11",
    ),
    (
        r"2\.5 閉迴路感知、機器人接近與視覺語義操作對照\d+2\.5 感測回授式接近與視覺語義操作對照\d+",
        "2.5 感測回授式接近與視覺語義操作對照11",
    ),
    (
        r"3\.6 第二階段：超音波閉迴路接近控制\d+3\.6 第二階段：超音波感測回授接近\d+3\.6 第二階段：超音波感測回授接近\d+",
        "3.6 第二階段：超音波感測回授接近15",
    ),
    (
        r"3\.6 第二階段：超音波閉迴路接近控制\d+3\.6 第二階段：超音波感測回授接近\d+",
        "3.6 第二階段：超音波感測回授接近15",
    ),
    (
        r"第五章、閉迴路接近、離線狀態判斷與夾取評估（第二、三階段）\d+第五章、感測回授接近、離線狀態判斷與夾取評估（第二、三階段）\d+第五章、感測回授接近、離線狀態判斷與夾取評估（第二、三階段）\d+",
        "第五章、感測回授接近、離線狀態判斷與夾取評估（第二、三階段）18",
    ),
    (
        r"第五章、閉迴路接近、離線狀態判斷與夾取評估（第二、三階段）\d+第五章、感測回授接近、離線狀態判斷與夾取評估（第二、三階段）\d+",
        "第五章、感測回授接近、離線狀態判斷與夾取評估（第二、三階段）18",
    ),
]


def remove_paragraph(paragraph: Paragraph) -> None:
    paragraph._element.getparent().remove(paragraph._element)


def has_image(paragraph: Paragraph) -> bool:
    return any(run._element.xpath(".//a:blip") for run in paragraph.runs)


def fix_text(text: str) -> str:
    out = text
    for pat, repl in TOC_FIXES:
        out = re.sub(pat, repl, out)
    return out


def repair(path: Path) -> None:
    doc = Document(str(path))
    seen_history = 0
    seen_captions: set[str] = set()
    to_remove: list[Paragraph] = []

    for para in doc.paragraphs:
        t = para.text.strip()
        fixed = fix_text(t)
        if fixed != t:
            for run in list(para.runs):
                run._element.getparent().remove(run._element)
            if fixed:
                para.add_run(fixed)

        if t == "1.3 研究歷程與系統演進":
            seen_history += 1
            if seen_history > 1:
                to_remove.append(para)
                continue

        if t.startswith("圖") and "  " in t:
            if t in seen_captions:
                # remove caption and previous image paragraph if present
                prev = para._element.getprevious()
                if prev is not None:
                    prev_para = Paragraph(prev, para._parent)
                    if has_image(prev_para):
                        to_remove.append(prev_para)
                to_remove.append(para)
                continue
            seen_captions.add(t)

    for para in to_remove:
        remove_paragraph(para)

    doc.save(str(path))
    imgs = sum(1 for p in doc.paragraphs for r in p.runs if r._element.xpath(".//a:blip"))
    print(f"Repaired {path.name}: images={imgs}, removed={len(to_remove)}")


def main() -> None:
    for path in PATHS:
        if path.exists():
            repair(path)


if __name__ == "__main__":
    main()