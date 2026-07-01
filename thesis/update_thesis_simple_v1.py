#!/usr/bin/env python3
"""Update advisor-facing thesis draft: simpler title, 感測回授 wording, history section, figures."""

from __future__ import annotations

import re
import subprocess
import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches
from docx.text.paragraph import Paragraph

THESIS = Path(__file__).resolve().parent
ROOT = THESIS.parent
SRC = Path("/home/lab109/下載/2THESIS_DRAFT_FCU_v1_no_pyroom.docx")
OUT_REPO = THESIS / "THESIS_DRAFT_FCU_v1.docx"
OUT_ALT = THESIS / "2THESIS_DRAFT_FCU_v1.docx"
OUT_DL = SRC

TITLE_ZH = "Isaac Sim 機械手臂超音波感測回授接近之模擬驗證"
TITLE_OLD = "基於 RTX Acoustic 超音波感測之機械手臂閉迴路接近控制與狀態判斷"

ABSTRACT_ZH = (
    "本研究在 NVIDIA Isaac Sim 6.0 中，使用 RTX Acoustic 超音波感測模組，"
    "輔助 UR10 系列機械手臂進行接近任務。研究分三階段："
    "先確認感測資料在固定手臂條件下可重複擷取（30 次皆有效）；"
    "再以感測回授驅動逐步接近，並與未使用回授的對照組比較；"
    "最後檢查感測資料能否支援「是否應停止前進」之離線判斷。"
    "結果顯示，感測回授組進入 0.45 m 目標區比例為 84.0%，對照組為 29.2%；"
    "僅使用感測特徵時，停止區域判斷 F1 約 0.598。"
    "最終夾取成功率兩組皆約 20%，顯示接近與夾取應分開評估。"
    "本文並說明由環境建置、試錯到協定修正之研究歷程，供後續實機驗證參考。"
)

ABSTRACT_EN = (
    "This thesis evaluates RTX Acoustic ultrasonic sensor-feedback approach for UR10-class "
    "manipulators in NVIDIA Isaac Sim 6.0. The work proceeds in three stages: repeatable "
    "sensor capture under fixed-arm conditions (30 valid trials), sensor-feedback approach "
    "compared with a no-feedback baseline, and offline stop-region state judgment from "
    "sensor features. Sensor-feedback trials reached the 0.45 m zone in 84.0% of runs "
    "versus 29.2% for the baseline. Acoustic-only features achieved F1≈0.598 for stop-region "
    "classification. Final success remained near 20% in both groups, indicating that grasping "
    "should be evaluated separately from approach. The thesis also documents the iterative "
    "research process from environment setup to protocol refinement."
)

KEYWORDS_ZH = "關鍵詞：Isaac Sim；超音波感測；RTX Acoustic；感測回授；機械手臂接近"
KEYWORDS_EN = "Keywords: Isaac Sim; ultrasonic sensing; RTX Acoustic; sensor feedback; robotic approach"

HISTORY_HEADER = "1.3 研究歷程與系統演進"
HISTORY_BODY = (
    "本研究並非一次完成，而是依序累積而來。2026 年 3 至 4 月先完成 Isaac Sim 環境建置，"
    "確認 UR10 官方資產與腕部 RTX Acoustic 超音波感測模組可穩定擷取資料；"
    "5 月以固定機械手臂、移動目標方式完成 30 次重複實驗，確認感測特徵可重現且具距離趨勢。"
    "6 月將問題延伸至 UR10e 與 Robotiq 夾爪，嘗試以感測回授驅動接近，並與未使用回授的對照組比較。"
    "初期實驗發現，最終夾取成功率並未隨感測回授明顯提升；進一步排查後，"
    "將原因歸於接觸物理與抬升流程不穩定，而非接近階段本身。"
    "因此 7 月改採「只評估接觸、不要求穩定抬升」與隨機化目標位置，"
    "得到較乾淨的主結果。"
    "本文所稱感測回授接近，即利用即時感測特徵調整運動；控制概念上等同閉迴路，"
    "但正文統稱感測回授，以避免與一般軟體開發流程混淆。"
)

FIG = {
    "1": THESIS / "figures/fig1_research_timeline.png",
    "2": THESIS / "figures/fig2_scene_schematic.png",
    "3": THESIS / "figures/fig3_1_research_architecture.png",
    "4": THESIS / "figures/fig4_2_rtx_early_energy_vs_distance.png",
    "5": THESIS / "figures/fig5_approach_success_bar.png",
    "6": THESIS / "figures/fig6_state_f1_bar.png",
}

REPLACEMENTS = [
    (TITLE_OLD, TITLE_ZH),
    ("第五章、閉迴路接近、離線狀態判斷與夾取評估（第二、三階段）", "第五章、感測回授接近、離線狀態判斷與夾取評估（第二、三階段）"),
    ("3.6 第二階段：超音波閉迴路接近控制", "3.6 第二階段：超音波感測回授接近"),
    ("2.5 閉迴路感知、機器人接近與視覺語義操作對照", "2.5 感測回授式接近與視覺語義操作對照"),
    ("5.2 有無聲學回授之接近成功率比較", "5.2 感測回授與對照組之接近成功率比較"),
    ("閉迴路接近", "感測回授接近"),
    ("閉迴路控制", "感測回授控制"),
    ("閉迴路組", "感測回授組"),
    ("閉迴路超音波", "感測回授"),
    ("超音波閉迴路", "超音波感測回授"),
    ("非視覺超聲閉迴路接近", "非視覺超音波感測回授接近"),
    ("閉迴路運動", "感測回授運動"),
    ("閉迴路融合", "感測回授融合"),
    ("閉迴路狀態", "感測回授狀態"),
    ("閉迴路之", "感測回授之"),
    ("閉迴路", "感測回授"),
    ("聲學回授的感測回授控制", "感測回授控制"),
    ("聲學回授", "感測回授"),
    ("聲學特徵", "感測特徵"),
    ("聲學觀測", "感測觀測"),
    ("聲學融合距離", "感測融合距離"),
    ("聲學資料", "感測資料"),
    ("只使用聲學特徵", "只使用感測特徵"),
    ("僅使用聲學特徵", "僅使用感測特徵"),
    ("合併所有特徵", "合併所有特徵"),  # no-op anchor
    ("RQ2：加入感測回授的感測回授控制", "RQ2：感測回授控制是否能改善目標區到達率？"),
    ("RQ3：離線狀態判斷是否能從感測特徵取得可測量訊號？", "RQ3：感測特徵能否支援離線狀態判斷？"),
    ("貢獻二（次）：建立超音波感測回授接近與安全監督流程", "貢獻二（次）：建立超音波感測回授接近與安全監督流程"),
    ("last-meter 非視覺感測回授", "last-meter 非視覺感測回授"),
    ("第二階段", "第二階段"),
]


def insert_paragraph_after(paragraph: Paragraph, style: str = "Content") -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.style = style
    return new_para


def set_text(p: Paragraph, text: str, style: str | None = None, center: bool = False) -> None:
    if style:
        p.style = style
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    if text:
        p.add_run(text)


def insert_image_after(paragraph: Paragraph, image_path: Path, caption: str, width_in: float = 5.2) -> Paragraph:
    img_p = insert_paragraph_after(paragraph)
    img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_p.add_run().add_picture(str(image_path), width=Inches(width_in))
    cap_p = insert_paragraph_after(img_p)
    set_text(cap_p, caption, "Content", center=True)
    return cap_p


def cleanup_mangled_toc(text: str) -> str:
    """Fix duplicated TOC lines like '舊標題頁碼新標題頁碼'."""
    patterns = [
        (r"2\.5 閉迴路感知、機器人接近與視覺語義操作對照\d+2\.5 感測回授式接近與視覺語義操作對照\d+", "2.5 感測回授式接近與視覺語義操作對照11"),
        (r"3\.6 第二階段：超音波閉迴路接近控制\d+3\.6 第二階段：超音波感測回授接近\d+", "3.6 第二階段：超音波感測回授接近15"),
        (
            r"第五章、閉迴路接近、離線狀態判斷與夾取評估（第二、三階段）\d+第五章、感測回授接近、離線狀態判斷與夾取評估（第二、三階段）\d+",
            "第五章、感測回授接近、離線狀態判斷與夾取評估（第二、三階段）18",
        ),
    ]
    out = text
    for pat, repl in patterns:
        out = re.sub(pat, repl, out)
    return out


def replace_in_doc(doc: Document) -> None:
    for para in doc.paragraphs:
        t = para.text
        if not t:
            continue
        new = cleanup_mangled_toc(t)
        for old, new_val in REPLACEMENTS:
            if old in new:
                new = new.replace(old, new_val)
        if new.startswith("本研究從應用電聲角度"):
            new = ABSTRACT_ZH
        elif new.startswith("This thesis evaluates RTX Acoustic") or new.startswith("This thesis examines"):
            new = ABSTRACT_EN
        elif new.startswith("關鍵詞："):
            new = KEYWORDS_ZH
        elif new.startswith("Keywords:"):
            new = KEYWORDS_EN
        elif new.startswith("RQ2：加入感測回授"):
            new = "RQ2：感測回授控制是否能改善目標區到達率？"
        if new != t:
            set_text(para, new, para.style.name if para.style else "Content")

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    t = para.text
                    new = cleanup_mangled_toc(t)
                    for old, new_val in REPLACEMENTS:
                        if old in new:
                            new = new.replace(old, new_val)
                    if new != t:
                        set_text(para, new)


def find_paragraph(doc: Document, exact: str) -> Paragraph | None:
    for para in doc.paragraphs:
        if para.text.strip() == exact:
            return para
    return None


def insert_history_section(doc: Document) -> Paragraph | None:
    if find_paragraph(doc, HISTORY_HEADER):
        return None
    anchor = find_paragraph(doc, "RQ4：在僅評估接觸、不要求穩定抬升的條件下，夾取流程是否具趨勢級可行性？")
    if anchor is None:
        return None
    h = insert_paragraph_after(anchor, "Header2")
    set_text(h, HISTORY_HEADER, "Header2")
    c = insert_paragraph_after(h, "Content")
    set_text(c, HISTORY_BODY, "Content")
    cap = c
    # renumber following headers in body only
    for para in doc.paragraphs:
        t = para.text.strip()
        if t == "1.3 研究範圍與限制":
            set_text(para, "1.4 研究範圍與限制", para.style.name)
        elif t == "1.4 名詞解釋與研究貢獻":
            set_text(para, "1.5 名詞解釋與研究貢獻", para.style.name)
    return cap


def insert_extra_figures(doc: Document) -> None:
    """Skip Chinese matplotlib figures — they render as boxes without CJK fonts."""
    return


def ensure_figures() -> None:
    subprocess.run([sys.executable, str(THESIS / "generate_thesis_simple_figures.py")], check=True)
    subprocess.run([sys.executable, str(THESIS / "generate_fig31.py")], check=True)
    subprocess.run([sys.executable, str(THESIS / "generate_thesis_figures.py")], check=True)
    for p in FIG.values():
        if not p.exists():
            raise FileNotFoundError(p)


def process(src: Path, dst: Path) -> None:
    doc = Document(str(src))
    replace_in_doc(doc)
    insert_history_section(doc)
    insert_extra_figures(doc)
    doc.save(str(dst))
    print(f"Wrote {dst}")


def main() -> None:
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    ensure_figures()
    process(SRC, OUT_REPO)
    process(SRC, OUT_ALT)
    process(SRC, OUT_DL)


if __name__ == "__main__":
    main()