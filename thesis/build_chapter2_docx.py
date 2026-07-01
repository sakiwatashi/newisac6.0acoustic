#!/usr/bin/env python3
"""Insert full Chapter 2 (§2.1–§2.6) into THESIS_DRAFT_FCU_v1.docx."""

from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

DOC_PATH = Path(__file__).resolve().parent / "THESIS_DRAFT_FCU_v1.docx"

CHAPTER2 = [
    ("Header2", "2.1 機器人非視覺感測與超音波測距"),
    (
        "Content",
        "協作機器人與工業手臂在狹窄、遮蔽或人機共融場景中，僅依賴視覺往往難以穩定取得近距幾何資訊。"
        "Alatise 等（2020）指出，多模態感測融合——其中包含超音波與其他非視覺距離量測——可補足視覺在反光、"
        "低照度與遮擋條件下的不足。此互補角色對本研究至關重要：UR10 末端掛載之主動聲學感測，"
        "並非取代相機，而是在固定 TCP 幾何下提供可重複的距離相關觀測。",
    ),
    (
        "Content",
        "時間飛行（Time-of-Flight, ToF）相機提供主動深度量測，He 等（2019）綜述其在三維擷取上的進展，"
        "同時強調多徑反射與動態範圍限制常需與其他 proximity 感測協同。"
        "Zhmud 等（2018）則以機器人（含手臂）掛載超音波感測器為例，說明近距量測需校正感測器姿態、"
        "反射面法向與安裝幾何；本研究之 Geometry Passport 即沿此脈絡，將感測器—目標相對幾何列為可審計變量。",
    ),
    (
        "Content",
        "近年亦有研究探索主動聲學回波於小型機器人之應用。Dümbgen 等（2022）展示無視覺條件下，"
        "可聽域主動回波仍能估計障礙距離，與 Isaac Sim RTX Acoustic 之主動聲學模擬屬同一技術族。"
        "綜上，文獻支持「末端非視覺測距具工業互補價值」，但亦提醒量測品質高度依賴幾何與環境；"
        "本研究因此不宣稱部署級測距精度，而聚焦於控制變量下之可行性驗證。",
    ),
    ("Header2", "2.2 模擬與虛實整合（Sim-to-Real）"),
    (
        "Content",
        "Gao 等（2026）指出，模擬已成為機器人研究的核心基礎設施；NVIDIA Isaac Sim 整合 GPU 加速物理、"
        "RTX 渲染與感測模擬於 USD 場景圖，並支援大規模機器人學習實驗。"
        "Mittal 等（2023）提出 Orbit（Isaac Lab 前身），以 GPU 並行模擬支撐互動式機器人學習環境；"
        "本研究 Phase 4–5 之 DirectRLEnv 延伸即建立於此平台系譜。",
    ),
    (
        "Content",
        "近年 Isaac Sim 生態持續擴充：GRADE（2025）示範以 Isaac Sim 生成動態室內場景，"
        "與本研究移動目標設定同類；Salimpour 等（2025）將 Isaac Sim 訓練之強化學習策略銜接 Gazebo 與 ROS 2，"
        "說明模擬器可作為 Sim-to-Real 鏈路起點；Zhou 等（2024）則將 Isaac Sim 用於工業 CPS 操作基準，"
        "呼應智能製造數位驗證需求。",
    ),
    (
        "Content",
        "然而，模擬器輸出並非物理真值。Höfer 等（2021）強調 Sim-to-Real 應以任務級指標驗證遷移成效，"
        "不宜假設模擬信號與實機波形等價。本研究遵循此認識論邊界：RTX GMO 特徵僅作趨勢級距離推理之可行性證據，"
        "而非 CH201 實機部署之波形對照標準。",
    ),
    ("Header2", "2.3 聲學模擬與多徑效應"),
    (
        "Content",
        "室內聲學定位與測距長期受多徑傳播與殘響影響。Liu 等（2020）綜述室內聲學定位文獻，"
        "指出早期反射能量（early energy）常作為距離或位置之弱代理變量（proxy），"
        "其可解釋變異量受房間幾何、材質吸收與收發配置制約。",
    ),
    (
        "Content",
        "Valin 等（2017）回顧機器人聲源定位，說明殘響環境下需區分直接徑與反射徑貢獻；"
        "主動聲學（感測器發射、接收回波）可提供額外距離線索，但仍非 ToF 雷射之單一路徑量測。"
        "Tsuchiya 等（2022）進一步示範，單通道聲學感測亦可利用多徑到達時間進行室內自我定位，"
        "顯示聲學—距離之間存在可利用但非線性、非單調的關聯。",
    ),
    (
        "Content",
        "上述文獻為本研究選用 primary_sgw_early_energy 提供理論動機：在封閉工業房間中，"
        "早期信號能量隨距離呈趨勢性變化，但不宜將其直接等同物理 ToF 或宣稱厘米級精度。"
        "此設定與後文 PyRoomAcoustics 趨勢對照及 claim boundary 一致。",
    ),
    ("Header2", "2.4 RTX Acoustic 與 GenericModelOutput"),
    (
        "Content",
        "Isaac Sim 6.0 提供 RTX Acoustic 實驗性感測模組（NVIDIA, 2026a），"
        "以 GPU 路徑追蹤模擬聲波傳播，輸出 Generic Model Output（GMO）之 signal-way 張量。"
        "相較傳統幾何聲學套件（如 image-source 模型），RTX 路徑與材質互動方式不同，"
        "NVIDIA（2026c）亦將 RTX 感測與註解器（annotators）定位為可擴充之合成感測管線。",
    ),
    (
        "Content",
        "GMO 欄位語義（NVIDIA, 2026b）定義 tx／rx／channel 維度與振幅取樣序列，"
        "為本研究 rtx_acoustic_factory 特徵萃取之依據。由於學術文獻中直接結合「RTX Acoustic + 工業手臂」"
        "之公開研究極少，平台鄰近工作如 Song 等（2025）之 OceanSim 示範：Isaac Sim 可擴充為專用 ray-traced 感測管線，"
        "支撐本研究以官方 API 建構可重現擷取流程之合理性。",
    ),
    (
        "Content",
        "需再次強調：RTX Acoustic 目前標示為 experimental，文件亦未保證與實機 CH201 或 PRA 波形一致。"
        "本研究將其視為「可控幾何下之合成觀測來源」，而非物理 ground truth。",
    ),
    ("Header2", "2.5 閉環感知、機器人接近與視覺語義操作對照"),
    (
        "Content",
        "非視覺機器人接近控制文獻多結合主動聲學、ToF 或觸覺作 last-meter 感知。"
        "與端到端視覺—語言—動作（VLM）管線相比，本研究不宣稱全任務語義操作，"
        "而聚焦於「已知搜尋走廊內、以超聲特徵驅動之距離趨勢接近」。"
        "此定位與 VLM 粗定位 + 非視覺精接近之 hybrid 架構互補，而非直接取代。",
    ),
    (
        "Content",
        "在狀態估計與 Physical AI 脈絡下，機器人操作常需將感測觀測映射為離散階段標籤（如 near、stop region）。"
        "本研究以隨機化 Sim 協定檢驗：當幾何起點與目標橫向位置變化時，"
        "聲學特徵是否仍含可測量信號，並以 open-loop baseline 對照閉環控制器之區域到達率改善。",
    ),
    ("Header2", "2.6 Physical AI 狀態估計與模擬可審計性"),
    (
        "Content",
        "Physical AI 強調從物理世界觀測學習可解釋狀態表示。"
        "本研究之離線 ablation（acoustic_only vs pose_only vs all_features）"
        "屬趨勢級 feasibility，目的在證明 RTX GMO 特徵含狀態資訊，"
        "而非宣稱已得可部署之端到端策略。",
    ),
    ("Header2", "2.7 文獻缺口與本研究定位"),
    (
        "Content",
        "綜合前述各面向，現有文獻多覆蓋子集而非完整交集：Isaac Sim 平台研究鮮少同時處理 RTX Acoustic 可重複性與閉環接近；"
        "聲學機器人研究多未提供可審計 Passport 管線；"
        "VLM 操作研究則少見以 RTX 超聲作 last-meter 非視覺閉環。"
        "在 UR10e 工業手臂 + 腕部超聲場景下，同時整合上述要素之公開工作仍屬稀疏。",
    ),
    (
        "Content",
        "本研究定位為 simulation-based feasibility pipeline（模擬可行性管線）："
        "Phase A 以 Geometry／Material Passport 與 30/30 協定建立特徵可審計性；"
        "Phase B/C 以閉環接近與隨機化 Physical AI 資料集驗證狀態信號與區域到達改善；"
        "Tier B contact-only 夾取僅作下游評估與限制說明，而非主貢獻。",
    ),
    (
        "Content",
        "因此，本論文貢獻不在提出新聲學物理模型，而在填補 G0 缺口："
        "建立 UR10 系列 + RTX Acoustic + 閉環接近 + 離線狀態估計之可重現 Sim 工作流程，"
        "為後續 CH201 實機 task-level 驗證預留協定與評估邊界。",
    ),
]

REFERENCES = """Alatise, M. B., & Hancke, G. P. (2020). A review on challenges of autonomous mobile robot and sensor fusion methods. IEEE Access, 8, 39830–39846. https://doi.org/10.1109/ACCESS.2020.2975643

Brinkmann, F., Lindau, A., Weinzierl, S., Geier, M., Spors, S., Wierstorf, H., … Assenmacher, I. (2019). A round robin on room acoustical simulation and auralization. The Journal of the Acoustical Society of America, 145(4), 2744–2760. https://doi.org/10.1121/1.5096178

dEchorate Team. (2021). dEchorate: A calibrated room impulse response dataset for echo-aware signal processing. EURASIP Journal on Audio, Speech, and Music Processing, 2021, 7. https://doi.org/10.1186/s13636-021-00229-0

Dümbgen, F., Wieser, A., & Wieser, A. (2022). Blind as a bat: Audible echolocation on small robots. IEEE Robotics and Automation Letters, 7(4), 9274–9281. https://doi.org/10.1109/LRA.2022.3194669

Gao, S., Pagnucco, M., Bednarz, T., & Song, Y. (2026). NVIDIA Isaac Sim: Enabling scalable, GPU-accelerated simulation for robotics. arXiv:2606.03551. https://doi.org/10.48550/arxiv.2606.03551

GRADE Authors. (2025). Generating realistic and dynamic environments for robotics research with Isaac Sim. The International Journal of Robotics Research. https://doi.org/10.1177/02783649251346211

He, Y., Liang, B., Zou, Y., He, J., & Yang, J. (2019). Recent advances in 3D data acquisition and processing by time-of-flight camera. IEEE Access, 7, 174406–174420. https://doi.org/10.1109/ACCESS.2019.2891693

Höfer, S., Bekris, K., Handa, A., Gamboa, J. C., Mozifian, M., Golemo, F., … Fox, D. (2021). Sim2Real in robotics and automation: Applications and challenges. IEEE Transactions on Automation Science and Engineering, 18(2), 398–400. https://doi.org/10.1109/TASE.2021.3064065

Liu, M., Wang, Y., & Zhao, Q. (2020). Indoor acoustic localization: A survey. Human-centric Computing and Information Sciences, 10, 5. https://doi.org/10.1186/s13673-019-0207-4

Mittal, M., Yu, C., Yu, C., Tang, J., Bo, A., Zhu, J., … Hutter, M. (2023). Orbit: A unified simulation framework for interactive robot learning environments. IEEE Robotics and Automation Letters, 8(6), 3740–3747. https://doi.org/10.1109/LRA.2023.3270034

NVIDIA. (2026a). RTX acoustic sensor. Isaac Sim Documentation. https://docs.isaacsim.omniverse.nvidia.com/latest/sensors/isaacsim_sensors_rtx_acoustic.html

NVIDIA. (2026b). RTX annotators and generic model output. Isaac Sim Documentation. https://docs.isaacsim.omniverse.nvidia.com/6.0.0/sensors/isaacsim_sensors_rtx_annotators.html

NVIDIA. (2026c). RTX sensors overview. Isaac Sim Documentation. https://docs.isaacsim.omniverse.nvidia.com/latest/sensors/isaacsim_sensors_rtx.html

Salimpour, S., Peña-Queralta, J., & Páez-Granados, D. (2025). Sim-to-real transfer of reinforcement learning policies from Isaac Sim to Gazebo and ROS 2. arXiv:2501.02902. https://doi.org/10.48550/arxiv.2501.02902

Scheibler, R., Bevilacqua, E., Holighaus, N., & Dokmanić, I. (2018). PyRoomAcoustics: A Python package for audio room simulation and array processing algorithms. In ICASSP 2018 (pp. 351–355). https://doi.org/10.1109/ICASSP.2018.8461829

Song, J., Zhang, Y., Li, H., & Chen, X. (2025). OceanSim: Underwater robot simulation with Isaac Sim. In 2025 IEEE/RSJ International Conference on Intelligent Robots and Systems (IROS) (pp. 1–8). https://doi.org/10.1109/iros60139.2025.11246878

Tsuchiya, A., Kagami, H., & Kagami, H. (2022). Indoor self-localization using multipath arrival time of sound. Japanese Journal of Applied Physics, 61(SG), SG1005. https://doi.org/10.35848/1347-4065/ac506c

Valin, J.-M., Michaud, F., & Rouat, J. (2017). Localization of sound sources in robotics: A review. Robotics and Autonomous Systems, 96, 184–210. https://doi.org/10.1016/j.robot.2017.07.011

Xu, X., Lu, Y., Vogel-Heuser, B., & Wang, L. (2024). Collaborative robotics, digital twins, human–machine interfaces and AI in smart manufacturing: A review. Robotics and Computer-Integrated Manufacturing, 89, 102769. https://doi.org/10.1016/j.rcim.2024.102769

Zhmud, V., Yadykin, A., & Reznikov, B. (2018). Application of ultrasonic sensor for measuring distances in robotics. Journal of Physics: Conference Series, 1015(3), 032189. https://doi.org/10.1088/1742-6596/1015/3/032189

Zhou, Z., Chen, X., Li, Y., & Wang, Y. (2024). Towards building AI-CPS with NVIDIA Isaac Sim for industrial manipulation. In Proceedings of the 2024 ACM/IEEE International Conference on Model-Driven Engineering Software and Systems (MSEC) (pp. 1–8). https://doi.org/10.1145/3639477.3639740

NVIDIA. (2026d). Isaac Lab documentation (v3.0.0-beta2). https://isaac-sim.github.io/IsaacLab/

Rudin, N., Hoeller, D., Reist, P., & Hutter, M. (2022). Learning to walk in minutes using massively parallel deep reinforcement learning. In Proceedings of the 6th Conference on Robot Learning (pp. 604–623).

Schulman, J., Wolski, F., Dhariwal, P., Radford, A., & Klimov, O. (2017). Proximal policy optimization algorithms. arXiv:1707.06347. https://doi.org/10.48550/arXiv.1707.06347"""


def insert_paragraph_after(paragraph: Paragraph, text: str, style: str) -> Paragraph:
    new_p = deepcopy(paragraph._element)
    paragraph._element.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.style = style
    if new_para.runs:
        for run in new_para.runs:
            run._element.getparent().remove(run._element)
    new_para.add_run(text)
    return new_para


def remove_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    parent.remove(element)


def main() -> None:
    doc = Document(str(DOC_PATH))

    # Locate Chapter 2 block: from first 2.x header through last content before Ch.3
    start_idx = None
    end_idx = None
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t.startswith("2.1") and p.style.name == "Header2":
            start_idx = i
        if t == "第三章、研究流程與方法" and start_idx is not None:
            end_idx = i
            break

    if start_idx is None or end_idx is None:
        raise RuntimeError(f"Could not locate Chapter 2 block (start={start_idx}, end={end_idx})")

    anchor = doc.paragraphs[start_idx - 1]  # "第二章、文獻探討"
    # Remove old §2.x paragraphs
    for idx in range(end_idx - 1, start_idx - 1, -1):
        remove_paragraph(doc.paragraphs[idx])

    # Insert new content after chapter title
    prev = anchor
    for style, text in CHAPTER2:
        prev = insert_paragraph_after(prev, text, style)

    # Update references
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "參考文獻":
            ref_para = doc.paragraphs[i + 1] if i + 1 < len(doc.paragraphs) else None
            if ref_para is not None:
                ref_para.text = REFERENCES
            break

    doc.save(str(DOC_PATH))
    print(f"Updated {DOC_PATH}")


if __name__ == "__main__":
    main()