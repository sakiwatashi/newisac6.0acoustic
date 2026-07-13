#!/usr/bin/env python3
"""Rebuild thesis docx (V2): six chapters aligned with 論文敘事對齊_V2_2026-07-09.md.

Reads THESIS_DRAFT_FCU_v1.docx (never modified) and writes a new
THESIS_DRAFT_FCU_v2.docx. Mechanism (helper functions, Content list format)
is copied from rebuild_thesis_six_chapters.py; that script and the v1 docx
are read-only inputs and are never touched by this script.

All quantitative claims below are sourced from:
  - docs/plan_v2/reports/S1_envelope_report.md
  - docs/plan_v2/reports/S2_datasheet_report.md
  - docs/plan_v2/reports/D1_approach_report.md
  - docs/plan_v2/reports/D15_arm_approach_report.md
  - docs/handoff/E3_target_invisibility_and_v9_audit.md
  - docs/WPM_EXPERIMENT_RULES.md
No number is invented here.
"""

from __future__ import annotations

import sys
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

THESIS = Path(__file__).resolve().parent
ROOT = THESIS.parent
DOC_IN = THESIS / "THESIS_DRAFT_FCU_v1.docx"
DOC_OUT = THESIS / "THESIS_DRAFT_FCU_v2.docx"

FIG: dict[str, Path] = {}  # V2 introduces no new figures; Image items unused.

sys.path.insert(0, str(THESIS))
from build_chapter2_docx import CHAPTER2, REFERENCES  # noqa: E402

# ---------------------------------------------------------------------------
# Front matter: abstract / keywords (V2 narrative)
# ---------------------------------------------------------------------------

# 2026-07-12 摘要可讀性重修:降密度、術語先白話、主數字與宣稱邊界不變。
ABSTRACT_ZH = (
    "本研究於 NVIDIA Isaac Sim 6.0 機器人模擬平台之 RTX Acoustic 超音波感測模組中，"
    "建立並驗證一套「包絡優先」之感測回授接近方法：先量測感測器在何種幾何條件下讀得到目標，"
    "再將機械手臂任務設計於已驗證之可感測範圍（感測包絡）內。"
    "感測特性方面，以五十二格「配對移除」掃描——同姿態下先量測有目標之波形，再將目標自場景移除後量測背景——"
    "界定可偵測包絡：五十二格中三十六格可偵測，指向角度為主宰因子，"
    "感測器後方手臂對前向偵測路徑之貢獻為零。"
    "距離由多影格平均回波之峰值位置，經已知距離線性自校後換算；"
    "視軸距離編碼之皮爾森相關係數 r＝0.9994，桌面高度目標之均方根誤差（RMSE）為 5.3 mm。"
    "在此包絡內，以三組對照臂（以下簡稱三臂）驗證純聲學閉環："
    "聲學臂依估距決定停止；盲走臂保留相同量測程序，但將估距置換為無資訊值（資訊消融對照）；"
    "開環臂完全不量測，僅執行固定行程。三臂使用相同隨機種子與相同目標組。"
    "主結果（感測器掛載於 UR10e 六軸機械手臂）顯示：控制器不讀取目標之真實世界座標時，"
    "停止位置與目標位置之相關 r＝0.9856、停止誤差 RMSE 2.8 cm，"
    "三十個試驗回合全數由聲學條件觸發停止，九十回合姿態稽核零違規；"
    "盲走臂於同一管線完全失能，確認接近行為來自聲學資訊，而非場景幾何巧合。"
    "在接近能力成立後，進一步完成端到端夾取：以聲學估距決定夾取位置，"
    "夾取中心與目標之相關 r＝0.9885，對位（夾爪中心與目標中心水平吻合）誤差 RMSE 1.9 cm；"
    "對位成功率 60% 顯著優於盲走對照之 23%（費雪精確檢定 p＝0.004）；"
    "對位成功後之升舉成功率為 83%。"
    "夾持於模擬器中以接觸觸發之附著機制實現，本文不宣稱物理摩擦夾持。"
    "側向（左右）感知方面，四項獨立實驗均顯示感測器單次輸出不含左右方向資訊；"
    "改以手臂移至五個量測位置，由距離交會（多點定位）恢復目標二維座標，並經完整三臂驗證："
    "側向估計 r＝0.950（RMSE 3.3 cm），二維閉環停止誤差 RMSE 1.9 cm，"
    "盲走臂同流程失能（15.0 cm，p＜0.001）。"
    "研究範圍限於單一隨機種子與確定性模擬引擎；摩擦夾持與實機驗證列為未來工作。"
)

ABSTRACT_EN = (
    "This thesis develops and validates an envelope-first methodology for ultrasonic "
    "sensor-feedback robotic approach in NVIDIA Isaac Sim 6.0 with the RTX Acoustic module. "
    "The method first measures under which geometries the sensor can detect a target, then "
    "places the manipulator task only inside that validated sensing envelope. "
    "A fifty-two-cell paired-removal scan—measuring the scene with the target present, then "
    "again after physically removing the target—maps detectability: thirty-six of fifty-two "
    "cells are detectable, pointing angle dominates, and an arm behind the sensor contributes "
    "nothing to the forward path. Distance is obtained from the peak location of a multi-frame "
    "averaged echo after linear self-calibration on known ranges (boresight encoding r=0.9994; "
    "tabletop-height target RMSE 5.3 mm). "
    "Inside this envelope, a three-arm design tests a purely acoustic closed loop: the acoustic "
    "arm stops on estimated range; the blind arm keeps the same measurement pipeline but replaces "
    "range with a non-informative value (information ablation); the open arm measures nothing and "
    "follows a fixed stroke. Arms share seeds and targets. With the sensor on a UR10e arm, and "
    "without reading target world coordinates, stop position correlates with the target at "
    "r=0.9856 (RMSE 2.8 cm); all thirty trials stop on the acoustic trigger, with zero posture "
    "violations across ninety trials. The blind arm fails on the same pipeline, so success is "
    "attributable to acoustic information rather than scene geometry. "
    "An end-to-end grasp stage then places the gripper from acoustic range alone "
    "(grasp-center-to-target r=0.9885; alignment RMSE 1.9 cm). Alignment rate is 60% versus 23% "
    "for the blind arm (Fisher's exact test, p=0.004); lift succeeds in 83% of aligned trials. "
    "Attachment in simulation uses a contact-triggered constraint; friction grasping is not claimed. "
    "Four independent tests show the sensor output carries no left-right cue. Recovering 2-D "
    "position by multilateration from five arm-moved vantage ranges, and validating with the same "
    "three-arm design, yields lateral r=0.950 (RMSE 3.3 cm), 2-D stop RMSE 1.9 cm, and blind-arm "
    "failure (15.0 cm, p<0.001). "
    "Scope is limited to a single random seed and a deterministic simulation engine; physical "
    "friction grasping and hardware validation remain future work."
)

KEYWORDS_ZH = "關鍵詞：RTX Acoustic；感測包絡；聲學閉環接近；資訊消融對照；實驗效度設計；Isaac Sim"
KEYWORDS_EN = (
    "Keywords: RTX Acoustic; sensing envelope; acoustic closed-loop approach; "
    "ablation control; experimental validity design; Isaac Sim"
)

# ---------------------------------------------------------------------------
# Chapter 2: reused verbatim from build_chapter2_docx (no changes per spec)
# ---------------------------------------------------------------------------

CH2_FIXED = list(CHAPTER2)
# 首次出現術語就地加註(僅改本腳本之副本,不動 build_chapter2_docx.py 原文):
# CH2 之「固定 TCP」先於第四章之全稱解釋出現,於此補全稱。
CH2_FIXED = [
    (item[0], item[1].replace("在固定 TCP 幾何下", "在固定工具中心點（Tool Center Point, TCP）幾何下", 1))
    if len(item) >= 2 and isinstance(item[1], str) and "在固定 TCP 幾何下" in item[1]
    else item
    for item in CH2_FIXED
]
# 「審計」一詞在全文最終檢查中須為 0（見任務規格）；CH2 之「可審計變量」與已移除之
# 舊管線審計敘事無關（純粹指「可稽核之變量」），僅換用同義詞避免誤判，不動原文語意。
CH2_FIXED = [
    (item[0], item[1].replace("列為可審計變量", "列為可稽核變量", 1))
    if len(item) >= 2 and isinstance(item[1], str) and "列為可審計變量" in item[1]
    else item
    for item in CH2_FIXED
]

# 2026-07-10 逐節審查(用戶要求:字數過少之節擴寫、舊敘事殘句更新、名詞白話化;
# 僅改本腳本之 CH2 副本,不動 build_chapter2_docx.py 原文):

# 2.3(原 146 字)追加白話解釋段
_EXP_23 = (
    "所謂多路徑（multipath），指聲波自發射器出發後，除了「直達目標、反射回來」這一條最短路徑外，"
    "還會經由牆面、桌面等其他表面多次反彈後抵達接收器；殘響（reverberation）則是這些較晚抵達之"
    "回波在時間上疊加拖尾之現象。兩者使接收波形中同時混雜多種幾何來源之訊號，"
    "單看波形上之一個峰值並不足以斷定其對應之物理路徑。"
    "基於此，本研究採取保守之量測立場：不宣稱模擬波形具備與實體量測相同之絕對精度，"
    "而是以「配對移除」（同場景下有目標與移除目標之成對量測，兩者之差即為目標本身之貢獻）"
    "隔離目標訊號，並以趨勢性指標（峰值位置隨距離之線性關係）而非單點絕對值建立距離編碼；"
    "完整之量測程序與稽核設計見第三章。"
)
CH2_FIXED = [
    (item[0], item[1] + _EXP_23)
    if len(item) >= 2 and isinstance(item[1], str) and "特徵定義與擷取流程見第三章" in item[1]
    else item
    for item in CH2_FIXED
]

# 2.4(原 374 字)追加:光線追蹤式聲學與參數化模型之差異白話版
_EXP_24 = (
    "此處值得以白話說明兩類模擬取徑之差異：參數化模型以數學公式直接產生「看起來合理」之回波"
    "（例如依距離套用固定之衰減曲線），計算快速但對場景中實際擺放之物體不敏感——"
    "目標移動了、模擬輸出卻可能不變；光線追蹤式模型則實際追蹤聲波與場景幾何（目標、桌面、手臂連桿）"
    "之碰撞與反射路徑，因此輸出會隨物體之位置、大小、角度而變化。"
    "唯有後者，「量測感測器在什麼幾何下讀得到目標」這一問題才有意義——"
    "此為本研究選用 RTX Acoustic 之核心理由，亦是第四章感測包絡量測之技術前提。"
)
CH2_FIXED = [
    (item[0], item[1] + _EXP_24)
    if len(item) >= 2 and isinstance(item[1], str)
       and "用來檢驗感測回授接近與離線狀態判斷是否可行" in item[1]
    else item
    for item in CH2_FIXED
]

# 2.5 第一段(原 150 字)追加:VLM 白話
_EXP_25A = (
    "此處之視覺語言模型（Vision-Language Model, VLM）指同時理解影像與自然語言之大型神經網路模型，"
    "其強項在於「理解場景中有什麼、大致在哪裡」之語義層次；"
    "但其空間定位精度受影像解析度與訓練資料限制，且在低照度、反光或遮蔽下失效。"
    "本研究之超音波感測回授恰好補足此弱點：不依賴光學條件、對「已知方向上之距離」提供公分級之連續回授。"
    "兩者組成之階層式架構中，語義理解與精細接近各司其職，此為本研究「互補而非取代」定位之具體含義。"
)
CH2_FIXED = [
    (item[0], item[1] + _EXP_25A)
    if len(item) >= 2 and isinstance(item[1], str) and "hybrid 架構互補，而非直接取代" in item[1]
    else item
    for item in CH2_FIXED
]

# 2.5 第二段(原 150 字,含過時敘事)整段改寫
_NEW_25B = (
    "在狀態估計脈絡下，機器人操作常需將連續之感測觀測映射為離散之任務狀態"
    "（例如「目標存在與否」「已進入可夾取範圍與否」），再由狀態驅動行為決策。"
    "本研究對此脈絡之立場是：狀態判斷之前提是感測訊號確實含有對應資訊，"
    "故本文先以第四章之包絡量測與第五章之消融對照確立「訊號中有什麼、沒有什麼」"
    "（距離趨勢存在、單次量測之側向資訊不存在），"
    "將學習式狀態估計留待此一因果地基確立之後（第六章未來工作），"
    "避免在訊號是否含資訊尚屬未知時即投入模型訓練、產生難以歸因之結果。"
)
CH2_FIXED = [
    (item[0], _NEW_25B)
    if len(item) >= 2 and isinstance(item[1], str) and "在狀態估計與 Physical AI 脈絡下" in item[1]
    else item
    for item in CH2_FIXED
]

# 2.6(原 133 字,含過時敘事)整段改寫並擴寫
_NEW_26 = (
    "模擬驗證有其明確之能與不能，本研究對此採取「邊界先行」之立場。"
    "模擬環境之獨特價值在於變因之完全可控與量測之完全可重現："
    "研究者可以精確地只移動一個物體、只改變一個角度，並取得逐位元可重複之量測結果，"
    "這使得「拿掉某一項資訊、觀察系統是否失能」之消融式因果檢驗得以嚴格執行——"
    "此類檢驗在實體環境中因雜訊與不可重現性而難以達到相同之嚴格程度。"
    "反之，模擬結論之效力邊界亦須明確劃定：本研究之確定性引擎不含實體感測器之熱雜訊與元件變異，"
    "亦不對頻率相關之聲學物理完整建模（詳見第四章之負結果記錄），"
    "故所有結論皆明確標註為模擬環境內之因果驗證，"
    "其對實體系統之可移轉性須另經實機實驗確認（第六章）。"
    "本文第六章之「可宣稱與不可宣稱範圍界定」即為此一立場之系統性落實。"
)
CH2_FIXED = [
    (item[0], _NEW_26)
    if len(item) >= 2 and isinstance(item[1], str) and "本研究進一步以離線特徵消融檢查" in item[1]
    else item
    for item in CH2_FIXED
]

# ---------------------------------------------------------------------------
# References: numeric citation style (advisor request). REFERENCES (imported
# from build_chapter2_docx) is a single string with entries separated by a
# blank line. We parse entries, sort by (year asc, first-author surname),
# number them [1..N], and rewrite in-text "Surname 等（YYYY）" / "（Author,
# YYYY）" citations to "Surname 等 [n]" / "[n]". No entry text is altered
# beyond prefixing "[n] "; the blank-line separation used by replace_body's
# write mechanism is preserved.
# ---------------------------------------------------------------------------

import re as _re


def _parse_reference_entries(refs: str) -> list[str]:
    return [e.strip() for e in refs.split("\n\n") if e.strip()]


def _ref_author_surname(entry: str) -> str:
    author_seg = entry.split("(")[0].strip(" .,")
    return author_seg.split(",")[0].strip()


def _ref_year_suffix(entry: str) -> str:
    m = _re.search(r"\((\d{4}[a-z]?)\)", entry)
    return m.group(1) if m else ""


def _ref_year_int(entry: str) -> int:
    m = _re.search(r"\((\d{4})", entry)
    return int(m.group(1)) if m else 9999


def _ref_sort_key(entry: str) -> tuple[int, str, str]:
    return (_ref_year_int(entry), _ref_author_surname(entry).lower(), _ref_year_suffix(entry))


# 2026-07-11 新增 3 筆(第二章新增 2.7 節之文獻依據;作者/年份/出處均經檢索核實):
# 2026-07-12 新增 2 筆(方法學支柱文獻:預註冊、消融;作者/年份/出處均經檢索核實):
# 2026-07-12 新增 3 筆(1.5 實驗平台官方文件:Isaac Sim 總覽、UR10e、Robotiq 2F-85):
REFERENCES = REFERENCES + "\n\n" + "\n\n".join([
    "Kerstens, R., Laurijssen, D., & Steckel, J. (2019). eRTIS: A fully embedded real time 3D "
    "imaging sonar sensor for robotic applications. In 2019 International Conference on Robotics "
    "and Automation (ICRA) (pp. 1438\u20131443). IEEE.",
    "Hayes, M. P., & Gough, P. T. (2009). Synthetic aperture sonar: A review of current status. "
    "IEEE Journal of Oceanic Engineering, 34(3), 207\u2013224.",
    "Kapoor, R., Ramasamy, S., Gardi, A., Bieber, C., Silverberg, L., & Sabatini, R. (2016). "
    "A novel 3D multilateration sensor using distributed ultrasonic beacons for indoor navigation. "
    "Sensors, 16(10), 1637. https://doi.org/10.3390/s16101637",
    "Nosek, B. A., Ebersole, C. R., DeHaven, A. C., & Mellor, D. T. (2018). "
    "The preregistration revolution. Proceedings of the National Academy of Sciences, "
    "115(11), 2600\u20132606. https://doi.org/10.1073/pnas.1708274114",
    "Meyes, R., Lu, M., de Puiseau, C. W., & Meisen, T. (2019). "
    "Ablation studies in artificial neural networks. arXiv:1901.08644. "
    "https://doi.org/10.48550/arXiv.1901.08644",
    "Robotiq. (2019). 2F-85 & 2F-140 instruction manual (e-Series). "
    "https://assets.robotiq.com/website-assets/support_documents/document/"
    "2F-85_2F-140_Instruction_Manual_e-Series_PDF_20190206.pdf",
    "Universal Robots. (2023). UR10e product factsheet. "
    "https://www.universal-robots.com/media/1828411/"
    "ur-product-factsheets-ur10e-06_2023-en-version-5-1-4-3.pdf",
    "NVIDIA. (2026e). Isaac Sim documentation. "
    "https://docs.isaacsim.omniverse.nvidia.com/latest/index.html",
])

_REF_ENTRIES = _parse_reference_entries(REFERENCES)
_REF_ENTRIES_SORTED = sorted(_REF_ENTRIES, key=_ref_sort_key)
_REF_NUMBER: dict[str, int] = {e: i for i, e in enumerate(_REF_ENTRIES_SORTED, start=1)}

REFERENCES_NUMBERED = "\n\n".join(f"[{_REF_NUMBER[e]}] {e}" for e in _REF_ENTRIES_SORTED)

# (surname_lower, "YYYY" or "YYYYa") -> number, for citations that include a
# disambiguating letter suffix (e.g. NVIDIA, 2026a).
_CITATION_LOOKUP: dict[tuple[str, str], int] = {}
# (surname_lower, year_int) -> list of candidate entries, to detect and warn
# on genuine ambiguity (same surname + same bare year, multiple references).
_BARE_YEAR_GROUPS: dict[tuple[str, int], list[str]] = {}
for _entry in _REF_ENTRIES_SORTED:
    _surname_full = _ref_author_surname(_entry)
    # Register both the full first-author segment ("GRADE Authors") and its
    # first whitespace token ("GRADE") as lookup keys: some in-text citations
    # abbreviate a group name to its first word (e.g. "（GRADE, 2025）" citing
    # "GRADE Authors. (2025)").
    _surname_keys = {_surname_full.lower(), _surname_full.split()[0].lower()}
    _yr_suffix = _ref_year_suffix(_entry)
    _yr_int = _ref_year_int(_entry)
    for _surname_l in _surname_keys:
        if _yr_suffix:
            _CITATION_LOOKUP[(_surname_l, _yr_suffix)] = _REF_NUMBER[_entry]
        _BARE_YEAR_GROUPS.setdefault((_surname_l, _yr_int), []).append(_entry)
        _bare_key = (_surname_l, str(_yr_int))
        _CITATION_LOOKUP.setdefault(_bare_key, _REF_NUMBER[_entry])

_CITATION_AMBIGUOUS: set[tuple[str, int]] = {
    k for k, v in _BARE_YEAR_GROUPS.items() if len(v) > 1
}

_CITATION_STATS = {"pattern1": 0, "pattern2": 0, "unresolved": []}

_PAT1 = _re.compile(r"([A-Za-zÀ-ÖØ-öø-ÿ]+)\s?(等)?（(\d{4}[a-z]?)）")
# Parenthetical citations: "（Author, YYYY）" and "（Author 等, YYYY）" (the
# author/等/year are all inside the full-width parens together, as opposed to
# _PAT1 where only the year is parenthesized).
_PAT2 = _re.compile(r"（([A-Za-zÀ-ÖØ-öø-ÿ]+)\s?(?:等)?, ?(\d{4}[a-z]?)）")


def _resolve_citation(surname: str, year: str) -> int | None:
    key = (surname.lower(), year)
    if key in _CITATION_LOOKUP:
        # Guard against a bare-year lookup landing on an ambiguous group.
        if not year[-1].isalpha() and (surname.lower(), int(year)) in _CITATION_AMBIGUOUS:
            return None
        return _CITATION_LOOKUP[key]
    return None


def _replace_citations_in_text(text: str) -> str:
    def _sub1(m: _re.Match) -> str:
        surname, deng, year = m.group(1), m.group(2), m.group(3)
        num = _resolve_citation(surname, year)
        if num is None:
            _CITATION_STATS["unresolved"].append(m.group(0))
            print(f"WARNING: could not map citation to a reference entry: {m.group(0)!r}")
            return m.group(0)
        _CITATION_STATS["pattern1"] += 1
        return f"{surname} 等 [{num}]" if deng else f"{surname} [{num}]"

    def _sub2(m: _re.Match) -> str:
        surname, year = m.group(1), m.group(2)
        num = _resolve_citation(surname, year)
        if num is None:
            _CITATION_STATS["unresolved"].append(m.group(0))
            print(f"WARNING: could not map citation to a reference entry: {m.group(0)!r}")
            return m.group(0)
        _CITATION_STATS["pattern2"] += 1
        return f"[{num}]"

    text = _PAT1.sub(_sub1, text)
    text = _PAT2.sub(_sub2, text)
    return text


def _replace_citations_in_items(items: list[tuple[Any, ...]]) -> list[tuple[Any, ...]]:
    new_items = []
    for item in items:
        kind = item[0]
        if kind == "Table":
            _, headers, rows = item
            new_rows = [[_replace_citations_in_text(c) for c in row] for row in rows]
            new_items.append((kind, headers, new_rows))
        elif len(item) >= 2 and isinstance(item[1], str):
            new_items.append((item[0], _replace_citations_in_text(item[1]), *item[2:]))
        else:
            new_items.append(item)
    return new_items


CH2_FIXED = _replace_citations_in_items(CH2_FIXED)


def _insert_after_containing(items: list, needle: str, new_items: list) -> list:
    """Splice new_items into items, right after the first item whose text
    contains `needle`. Raises if no match (fail loud rather than silently
    dropping content in the wrong place)."""
    out = list(items)
    for idx, it in enumerate(out):
        if len(it) >= 2 and isinstance(it[1], str) and needle in it[1]:
            return out[: idx + 1] + list(new_items) + out[idx + 1 :]
    raise RuntimeError(f"anchor text not found in items: {needle!r}")


# 補充引用(第二章原文未涵蓋之 11 篇參考文獻,依主題插入既有小節內):
CH2_FIXED = _insert_after_containing(
    CH2_FIXED,
    "與 Isaac Sim RTX Acoustic 之主動聲學模擬屬同一技術族",
    [
        (
            "Content",
            "聲源與距離定位之機器人文獻脈絡更廣。Valin 等（2017）系統性回顧機器人聲源定位方法，"
            "指出陣列幾何、多路徑與殘響為距離與方位估計之共同挑戰；"
            "Tsuchiya 等（2022）進一步以室內多路徑到達時間實現無地圖自我定位，"
            "顯示即使在強反射室內環境，時間域到達資訊仍可萃取空間資訊。"
            "此二文獻共同支持本研究以「配對移除」量測分離目標回波與場景背景多路徑之設計動機。",
        ),
    ],
)
CH2_FIXED = _insert_after_containing(
    CH2_FIXED,
    "不宜假設模擬信號與實機波形等價",
    [
        (
            "Content",
            "模擬平台之另一價值在於支援大規模強化學習訓練。Rudin 等（2022）展示以 GPU 大量並行模擬"
            "（數千環境同時執行）可將四足機器人行走策略之訓練時間縮短至數分鐘，"
            "此類大規模並行能力所依託之 GPU 加速模擬架構與 Isaac Sim 同源；"
            "策略最佳化演算法方面，Schulman 等（2017）提出之近端策略最佳化"
            "（Proximal Policy Optimization, PPO）為現行機器人強化學習之主流方法之一。"
            "本研究現行控制器為規則式比例步進控制，未涉及學習型策略；"
            "上述基礎設施顯示：若後續以強化學習取代規則式控制器，"
            "Isaac Sim／Isaac Lab 之 GPU 並行模擬能力已提供可行路徑（詳見第六章未來工作）。",
        ),
    ],
)
CH2_FIXED = _insert_after_containing(
    CH2_FIXED,
    "以早期能量等摘要特徵檢驗「距離趨勢是否可用於感測回授」",
    [
        (
            "Content",
            "室內聲學模擬之驗證方法學方面，Brinkmann 等（2019）之跨實驗室輪測（round-robin）"
            "比較多套房間聲學模擬工具，發現不同幾何聲學引擎對早期反射與殘響時間之預測存在系統性差異，"
            "凸顯「模擬—實機」比對須謹慎詮釋，不宜視為等價。"
            "Scheibler 等（2018）之 PyRoomAcoustics 為公開之房間聲學模擬與陣列訊號處理工具，"
            "提供影像來源法（image-source method）之開放實作；"
            "dEchorate 等（2021）建立之資料集則提供已校正之房間脈衝響應，供回聲感知演算法之基準測試。"
            "此三項工作共同顯示：房間聲學模擬之保真度評估已有成熟方法學基礎，"
            "然多聚焦於固定陣列之訊號處理，較少涵蓋機械手臂載具下之動態幾何場景——"
            "此為本研究以配對移除協定系統性量測動態場景可偵測度之切入點。",
        ),
    ],
)

# Re-run numeric-citation replacement: the three insertions above were
# spliced in AFTER CH2_FIXED's first replacement pass (line ~237), so their
# "Surname 等（YYYY）" citations were never converted. Re-running is safe/
# idempotent — already-converted "Surname 等 [n]" text does not match the
# author-year regex again.
# 2026-07-11 新增 2.7 節:聲學陣列、多點定位與合成孔徑(本文第五、六章之
# 側向感知與多點定位內容此前於文獻章無對應;插於 2.6 之後、原 2.7 之前):
CH2_FIXED = _insert_after_containing(
    CH2_FIXED,
    "即為此一立場之系統性落實",
    [
        ("Header2", "2.7 聲學陣列、多點定位與合成孔徑"),
        (
            "Content",
            "單一收發器之超音波感測僅能量測「指向方向上之距離」，無法分辨目標之左右方位；"
            "文獻對此提供兩大類解法。第一類為多元件陣列：以多個接收器同時接收回波，"
            "利用各接收器間之到達時間差或波束成形（beamforming，將多路訊號依幾何延遲相加以合成指向性）"
            "解算方位。Kerstens 等（2019）之嵌入式即時三維成像聲納即為此路線之代表，"
            "以麥克風陣列在空氣中實現三維超音波成像，其研究群並發展出受蝙蝠迴聲定位啟發之一系列仿生聲納系統。"
            "此路線之前提是各接收器之訊號可被獨立取得——本研究第四章之負結果顯示，"
            "現行模擬引擎之輸出不區分接收器身分，故此路線於本研究之模擬環境中不可行，"
            "此為本研究轉向第二類解法之直接原因。",
        ),
        (
            "Content",
            "第二類解法不依賴多接收器，而是以「多個量測位置」取代「多個接收器」："
            "多點定位（multilateration）自多個已知位置分別量測與目標之距離，以幾何交會解算目標座標，"
            "Kapoor 等（2016）以分佈式超音波信標實現室內三維定位即為典型應用；"
            "合成孔徑（synthetic aperture）則更進一步，讓單一感測器在移動中連續量測，"
            "等效合成一具大孔徑陣列，Hayes 等（2009）對此技術在聲納領域之發展有系統性回顧。"
            "本研究第五章之側向定位方案屬此一族系：以機械手臂本身作為感測器之移動平台，"
            "自五個橫向錯開之量測位置各取一次距離，以最小平方交會解算目標之二維座標——"
            "相較於信標方案需在環境中預先佈設多個發射器，"
            "本研究之方案僅需手臂既有之運動自由度，不增加任何硬體元件。",
        ),
    ],
)
CH2_FIXED = [
    (item[0], item[1].replace("2.7 文獻缺口與本研究定位", "2.8 文獻缺口與本研究定位", 1))
    if len(item) >= 2 and isinstance(item[1], str) and item[1].startswith("2.7 文獻缺口")
    else item
    for item in CH2_FIXED
]
CH2_FIXED = _replace_citations_in_items(CH2_FIXED)

# ---------------------------------------------------------------------------
# 2026-07-13 CH2 深修（對齊第一章節奏）：整章重寫正文；引用仍用作者—年份，再轉 [n]。
# 舊 CH2 拼接鏈僅作文獻來源考古；定稿以本節 CH2_FIXED 為準。
# ---------------------------------------------------------------------------
CH2_FIXED = [
    ("Header2", "2.1 機器人非視覺感測與超音波測距"),
    (
        "Content",
        "工業手臂近距工作中，非視覺距離感測仍有其位置。"
        "協作機器人與工業手臂常進入狹窄、遮蔽或人機共處的空間；此時只靠相機，"
        "往往難以穩定取得「還有多遠」這類近距幾何資訊。"
        "Alatise 等（2020）指出，多模態感測融合——其中包含超音波與其他非視覺距離量測——"
        "可補足視覺在反光、低照度與遮擋下的不足。"
        "對本研究而言，UR10e 腕部附近的主動聲學感測並非要取代相機，"
        "而是在固定的工具中心點（Tool Center Point, TCP：手臂末端工具的參考點）幾何下，"
        "提供可重複的距離相關觀測。",
    ),
    (
        "Content",
        "時間飛行（Time-of-Flight, ToF）相機以主動打光量測深度，He 等（2019）綜述其進展，"
        "也提醒多徑反射與動態範圍限制，常需搭配其他近距（proximity）感測。"
        "Zhmud 等（2018）以機器人（含手臂）掛載超音波為例，強調近距量測必須校正"
        "感測器姿態、反射面法向與安裝幾何。"
        "本研究沿此脈絡，把感測器安裝姿態與「感測器—目標」相對幾何列為可稽核變量"
        "（實驗記錄稱 Geometry Passport，即「幾何護照」：掛載與場景幾何寫清楚、可核對）。",
    ),
    (
        "Content",
        "主動聲學回波亦見於小型機器人：Dümbgen 等（2022）展示在無視覺條件下，"
        "可聽域（人耳可聽頻率範圍附近）主動回波仍能估計障礙距離；"
        "這與 Isaac Sim 中 RTX Acoustic 的主動聲學模擬屬同一技術族。"
        "文獻共識可概括為兩點：末端非視覺測距具工業互補價值；"
        "但量測品質高度依賴幾何與環境。"
        "因此本文不宣稱部署級測距精度，而聚焦「在控制好的變量下是否可行，以及因果能否歸因到聲學」。",
    ),
    (
        "Content",
        "更廣的聲源與距離定位文獻，進一步支持本研究的量測設計。"
        "Valin 等（2017）回顧機器人聲源定位，指出陣列幾何、多路徑與殘響是距離與方位估計的共同挑戰；"
        "Tsuchiya 等（2022）則以室內多路徑到達時間做無地圖自我定位，"
        "顯示即使在強反射室內，時間域到達資訊仍可能帶有空間訊息。"
        "這兩條線共同支持以「配對移除」分離目標回波與背景多路徑（程序見 3.3 節），"
        "而不是把整條波形直接當成距離真理。",
    ),
    ("Header2", "2.2 模擬與虛實整合（Sim-to-Real）"),
    (
        "Content",
        "本節討論：為何在模擬裡做，以及模擬結論能講到哪裡。"
        "Gao 等（2026）為 NVIDIA Isaac Sim 的專題綜述，指出模擬已成機器人研究的基礎設施。"
        "平台建立在 Omniverse（三維協作與模擬基礎環境）上，"
        "整合 GPU 加速 PhysX 物理引擎、RTX 光線追蹤渲染與 USD 場景描述，"
        "並內建彩色影像／深度、光達（LiDAR）、慣性量測單元（IMU）等感測模擬。"
        "該綜述亦指出 Isaac Lab 可支援大規模並行強化學習；"
        "Mittal 等（2023）的 Orbit 是其系譜前身，後續併入 Isaac Lab 生態。"
        "本研究以 Isaac Sim 6.0 為主實驗平台；Isaac Lab 僅保留為未來延伸路徑（見 1.5、6.3 節）。",
    ),
    (
        "Content",
        "近年 Isaac Sim 被用於動態場景建置（GRADE, 2025）、"
        "Sim-to-Real 策略遷移（Salimpour 等, 2025）與工業操作驗證（Zhou 等, 2024）。"
        "但 Gao 等（2026）並未專論 RTX Acoustic；模擬器輸出也不是物理真值。"
        "Höfer 等（2021）強調應以任務級指標談遷移，不宜假設「模擬波形＝實機波形」。"
        "本研究採同一邊界：RTX Acoustic 特徵只當趨勢級距離推理的可行性證據，"
        "不是 TDK CH201 等實機波形的對照標準。",
    ),
    (
        "Content",
        "模擬的另一價值是大規模強化學習。"
        "Rudin 等（2022）展示以 GPU 並行數千環境，可在短時間內訓練四足行走策略；"
        "Schulman 等（2017）的近端策略最佳化（Proximal Policy Optimization, PPO）"
        "則是常見的策略優化演算法。"
        "本文現行控制器是規則式比例步進，並未使用 PPO 或任何學習型策略；"
        "上述文獻只說明：若未來要以強化學習取代規則控制器，"
        "Isaac Sim／Isaac Lab 已提供可行的訓練基礎設施，且仍須沿用本文的三臂效度框架"
        "以防「學到場景捷徑」（詳見第六章）。",
    ),
    ("Header2", "2.3 室內環境與距離特徵"),
    (
        "Content",
        "室內超聲不能「看一個峰就當距離」；本節說明原因與本文的保守立場。"
        "封閉工業環境中，回波會受牆面與工件反射影響，距離很少是單一路徑量測。"
        "Liu 等（2020）指出，早期反射能量可作距離的弱趨勢指標，但仍受房間幾何與材質制約。"
        "因此本文不宣稱厘米級絕對測距，而以摘要特徵檢驗"
        "「距離趨勢是否足以支撐感測回授」；特徵定義見第三章。",
    ),
    (
        "Content",
        "所謂多路徑（multipath），指聲波除了「打到目標再回來」的最短路徑外，"
        "還會經牆面、桌面等多次反彈才到達接收器；"
        "殘響（reverberation）則是較晚到的回波在時間上拖尾疊加。"
        "兩者使波形同時混有多種幾何來源，單看一個峰值不足以斷定物理路徑。"
        "本文因此採取保守量測：不宣稱模擬波形與實體量測絕對精度相同；"
        "而以配對移除隔離目標貢獻（程序見 3.3 節），"
        "並以峰值位置隨距離的線性趨勢建立距離編碼。",
    ),
    (
        "Content",
        "室內聲學模擬的保真評估已有方法學傳統。"
        "Brinkmann 等（2019）的跨實驗室輪測（round-robin：同一題目多實驗室各自跑再比較）"
        "顯示不同引擎對早期反射與殘響時間的預測常有系統差，"
        "故「模擬—實機」不宜直接等價。"
        "Scheibler 等（2018）的 PyRoomAcoustics 提供影像來源法"
        "（image-source：以鏡像聲源近似房間反射）的開放實作；"
        "dEchorate 等（2021）則提供已校正的房間脈衝響應資料集，供演算法基準測試。"
        "這些工作多聚焦固定陣列與房間聲學；"
        "機械手臂載具下、幾何隨時間改變的動態場景可偵測度，仍是較少被系統量測的切入點——"
        "此即第四章包絡實驗的文獻位置。",
    ),
    ("Header2", "2.4 RTX Acoustic 與感測輸出（GMO）"),
    (
        "Content",
        "以下說明 RTX Acoustic 是什麼、為何選它，以及它不能被當成什麼。"
        "Isaac Sim 6.0 提供 RTX Acoustic 超音波感測模組（NVIDIA, 2026a），屬實驗性功能："
        "在虛擬場景中以 GPU 產生回波資料（以信號路徑 signal way 組織），不輸出點雲。"
        "本研究把它視為「可控幾何下的合成超音波觀測」，"
        "用來檢驗感測回授接近是否可行，以及訊號裡到底有沒有側向等資訊"
        "（後者見第四章；閉環與消融見第五章）。",
    ),
    (
        "Content",
        "可用白話區分兩類模擬。"
        "參數化模型用公式直接產生「看起來合理」的回波（例如依距離套固定衰減），"
        "算得快，但對場景裡真實擺放的物體可能不敏感——目標移了，輸出卻幾乎不變。"
        "光線追蹤式模型則追蹤聲波與場景幾何（目標、桌面、手臂連桿）的反射路徑，"
        "輸出會隨位置、大小、角度改變。"
        "只有後者，「量測感測器在什麼幾何下讀得到目標」才有意義——"
        "此為選用 RTX Acoustic／WPM 的主要理由，也是第四章包絡量測的技術前提。",
    ),
    (
        "Content",
        "官方說明輸出包含發射端、接收端、通道與振幅取樣（NVIDIA, 2026a）；"
        "本研究由此整理峰值位置與早期能量等特徵，定義與流程見第三章。"
        "Gao 等（2026）平台綜述未涵蓋此模組；"
        "學術上相近的延伸案例包括 Song 等（2025）的 OceanSim，"
        "顯示 Isaac Sim 可做成專用光線追蹤式感測管線。"
        "再次強調邊界：RTX Acoustic 不保證與實機 CH201 波形一致；"
        "全文結論以任務級、趨勢級指標表述，不宣稱部署級測距精度。",
    ),
    ("Header2", "2.5 感測回授式接近與視覺語義操作對照"),
    (
        "Content",
        "本文與視覺—語言操作路線是分工，而不是互相取代。"
        "視覺語言模型（Vision-Language Model, VLM）同時理解影像與自然語言，"
        "強項是「場景裡有什麼、大概在哪」的語義層次；"
        "但空間精度受影像與訓練資料限制，且在低照度、反光或遮蔽下容易失效。"
        "非視覺接近文獻則常結合主動聲學、ToF 或觸覺，做最後一公尺（last-meter）精細靠近。",
    ),
    (
        "Content",
        "本研究不宣稱端到端語義操作，而聚焦："
        "在已知搜尋走廊（目標被允許出現的工作帶）內，以超聲距離趨勢驅動接近。"
        "這與「VLM 粗定位＋非視覺精接近」的混合式（hybrid）階層互補："
        "上層管場景理解與大範圍導航，下層管最後一段非視覺距離回授。"
        "超聲的互補點很具體：不依賴光學條件，並對已知方向上的距離提供公分級連續回授。",
    ),
    (
        "Content",
        "機器人操作有時還會把連續感測映射成離散任務狀態"
        "（例如「目標在不在」「是否已進入可夾取範圍」），再驅動行為。"
        "本文立場是：先證明訊號裡「有什麼、沒有什麼」，再談學習式狀態估計。"
        "因此第四章先做包絡與側向證偽，第五章以消融對照確立距離趨勢能否驅動閉環；"
        "學習式狀態估計留待因果地基之後（第六章），"
        "避免在資訊是否存在都未查清時就訓練模型、結果難以歸因。",
    ),
    ("Header2", "2.6 模擬驗證邊界與消融式因果檢驗"),
    (
        "Content",
        "模擬特別適合回答某些因果問題，但也有明確效力邊界。"
        "模擬的長處是變因可控、量測可重現："
        "可以只移動一個物體、只改一個角度，並取得可逐位重複的量測。"
        "因此「拿掉某一項資訊，看系統是否失能」的消融式檢驗可以做得很乾淨；"
        "在實體環境裡，雜訊與不可重現性往往使同等嚴格的對照難以成立。",
    ),
    (
        "Content",
        "模擬的短處也必須先講清楚。"
        "本研究所用確定性引擎不含實體感測器的熱雜訊與元件變異，"
        "也不完整建模頻率相關的聲學物理（見第四章負結果）。"
        "故所有正負結果都標註為「模擬環境內的因果驗證」；"
        "能否轉移到實體系統，須另以實機任務級實驗確認（第六章）。"
        "第六章的可宣稱／不可宣稱清單，即是把本節立場寫成全文邊界。",
    ),
    ("Header2", "2.7 聲學陣列、多點定位與合成孔徑"),
    (
        "Content",
        "單次超聲測距沒有左右資訊時，文獻上有兩大類解法；本文選其中一條。"
        "單一收發器大致只能量「指向方向上的距離」，無法分辨目標在左或在右。"
        "第一類解法是多元件陣列：多個接收器同時收音，"
        "用到達時間差或波束成形（beamforming：依幾何延遲把多路訊號對齊相加，形成指向性）解方位。"
        "Kerstens 等（2019）的嵌入式即時三維成像聲納是此路線代表。"
        "前提是各接收器訊號能被獨立取得——"
        "第四章顯示現行模擬輸出不區分接收器身分，故陣列路線在本模擬環境中不可行。",
    ),
    (
        "Content",
        "第二類解法不增加接收器，而增加量測位置。"
        "多點定位（multilateration）在多個已知位置各量一次距離，以幾何交會解座標；"
        "Kapoor 等（2016）的分散式超音波信標室內定位是典型例子。"
        "合成孔徑（synthetic aperture）則讓單一感測器在移動中連續量測，等效成更大孔徑；"
        "Hayes 等（2009）回顧其在聲納領域的發展。"
        "本研究第五章的側向方案屬此一族："
        "手臂本身當移動平台，在五個橫向錯開位置各取距離，再以最小平方交會解二維座標。"
        "相較信標方案須在環境預佈發射器，本方案只使用手臂既有運動自由度，不新增硬體。",
    ),
    ("Header2", "2.8 文獻缺口與本研究定位"),
    (
        "Content",
        "綜合前述，現有文獻多覆蓋子集，而少見完整交集："
        "Gao 等（2026）整理 Isaac Sim，但未涵蓋 RTX Acoustic 與工業手臂感測回授接近；"
        "聲學機器人研究常缺少完整幾何、材質與任務設定紀錄；"
        "VLM 操作研究則少見以 RTX 超音波做最後一公尺非視覺回授。"
        "在 UR10e 與腕部超音波的場景下，同時把平台、感測、閉環接近與效度對照做完整的公開工作仍有限。",
    ),
    (
        "Content",
        "本研究定位為模擬可行性研究（simulation-based feasibility study）。"
        "先檢查感測特徵是否穩定、可重複；"
        "再以閉環接近與資訊消融對照，檢查特徵能否真正驅動行為；"
        "夾取是任務鏈下游延伸與限制分析；"
        "學習式狀態估計留待因果地基之後（第六章），不是本文主貢獻。",
    ),
    (
        "Content",
        "因此，貢獻不在提出新的聲學物理模型，而在填補文獻上的交集缺口："
        "把 UR10e、RTX Acoustic、感測回授接近與可重複效度流程放在同一條可重跑鏈上，"
        "並為後續 CH201 等實機任務級驗證保留協定與評估邊界。"
        "第三章說明方法與場景；第四章給感測包絡與負結果；"
        "第五章給閉環、夾取與二維定位；第六章收回宣稱範圍。",
    ),
]
CH2_FIXED = _replace_citations_in_items(CH2_FIXED)

# ---------------------------------------------------------------------------
# Chapter 1
# ---------------------------------------------------------------------------

# 2026-07-12 CH1 可讀性重修:術語白話、1.3 與 5.4 敘事對齊、名詞表分段。
CH1 = [
    ("Header1", "第一章、緒論"),
    ("Header2", "1.1 研究背景與問題意識"),
    (
        "Content",
        "協作機器人與工業手臂之應用，日益要求人機能安全共處、設備能快速改線佈署。"
        "Xu 等（2024）之綜述指出，數位孿生、人機介面與人工智慧之整合已成智慧製造共同趨勢；"
        "其中末端（手臂最前端）感測是否穩健——尤其在遮蔽、反光或動態變化的近距工作空間——"
        "仍是限制彈性佈署的關鍵瓶頸。"
        "室內主動聲學系統中，多徑傳播（multipath：聲波經牆面等表面多次反彈後才抵達接收器）"
        "與殘響（較晚到達的回波在時間上拖尾疊加）使接收信號同時混有幾何、材質與距離訊息。"
        "此一特性既是超音波測距的挑戰，也是它相對相機與雷射的互補優勢："
        "超音波不受光學反光、低照度或煙塵遮蔽影響，且感測硬體成本相對低廉。",
    ),
    (
        "Content",
        "在機器人末端之非視覺「最後一公尺」（last-meter）接近中——"
        "亦即目標大致已進入工作區後、仍須靠距離回授完成的精細靠近——"
        "如何從可控的模擬條件取得可重現的聲學觀測，並用來驅動閉環運動，"
        "是應用電聲與機器人感知的交叉問題。"
        "NVIDIA Isaac Sim 6.0 提供 RTX Acoustic 實驗性超音波模組，"
        "可輸出名為 Generic Model Output（GMO）的原始感測資料；"
        "資料以「信號路徑」（signal way：單一發射器—接收器組合的回波時間序列）為單位記錄振幅。"
        "平台與模組的白話說明見 1.5 節。"
        "相較於實體超音波晶片（如 TDK CH201）須搭配實體場地與硬體，"
        "模擬可精確控制幾何、材質與感測器擺位，使因果關係較易被隔離與驗證——"
        "此即本研究以模擬為主要驗證場域的理由。",
    ),
    (
        "Content",
        "與「相機加視覺語言模型（Vision-Language Model, VLM）」的端到端操作相比，"
        "本研究聚焦非視覺超聲閉環接近，定位為互補而非取代。"
        "VLM 擅長理解場景語義與粗定位；超聲則擅長在已知工作帶"
        "（本文稱搜尋走廊：目標被允許隨機出現的縱向／側向範圍）內，"
        "依距離趨勢做精細接近。"
        "兩者可組成階層：上層負責場景理解與大範圍導航，下層負責最後一段非視覺精細接近。"
        "本研究不宣稱取代視覺方案，而是為此類混合架構提供一段可驗證、可量化效度的非視覺感測回授元件。",
    ),
    ("Header2", "1.2 研究目的與問題陳述"),
    (
        "Content",
        "本研究以下列四項研究問題（Research Question, RQ）組織全文之實驗設計與驗證邏輯。"
        "每一問題皆對應獨立實驗，並在執行前寫定通過判準，避免事後依結果回頭改標準。",
    ),
    (
        "Content",
        "RQ1（感測包絡）：超音波感測器在何種「感測器—目標」幾何下能穩定偵測目標？包絡邊界在哪？"
        "若不先回答此問，後續手臂任務可能建在未經驗證的感測假設上；"
        "一旦失敗，難以分辨是控制策略問題，還是該幾何下根本讀不到訊號。",
    ),
    (
        "Content",
        "RQ2（聲學閉環接近）：不讀取目標真實世界座標的純聲學控制迴路，"
        "能否驅動機械手臂接近隨機擺放的目標？"
        "且其因果能否以資訊消融對照（盲走臂失能：保留量測動作、拿掉估距資訊）證實來自聲學？"
        "中心問題是：僅有「看起來接近成功」不夠；"
        "必須在拿掉聲學資訊後對照組仍失敗，才能排除「只是場景幾何造成成功」。",
    ),
    (
        "Content",
        "RQ3（實驗效度）：如何用實驗設計排除「成功來自場景幾何或隱含目標真值」？"
        "這無法只靠單一統計檢定；需要一組對照與稽核（詳見第三章四支柱）。",
    ),
    (
        "Content",
        "RQ4（範圍邊界）：夾取整合與側向（左右）定位的可行邊界各到哪裡？"
        "夾取方面：聲學對位能力已驗證；夾持機制與摩擦限制之宣稱邊界見 5.3、6.2 節。"
        "側向方面：單次輸出經四項實驗證偽後，以多點定位恢復左右資訊並完成正式三臂（見 5.4 節）。"
        "兩項答案共同標出本文的限制與後續工作方向。",
    ),
    ("Header2", "1.3 研究範圍與限制"),
    (
        "Content",
        "本研究實驗分層遞進、互相引用前層結論。"
        "納入項目如下（括號中的 S1、D1 等為全文統一實驗代號，便於對照章節與數據目錄）："
        "（一）Isaac Sim 6.0 本機獨立執行環境與 UR10e 官方機械手臂資產；"
        "（二）感測包絡地圖（S1：五十二格配對移除掃描）；"
        "（三）感測器量化特性表（S2：距離、桌面高度、側向、重複性）；"
        "（四）聲學閉環接近三臂對照——D1 為未掛臂、可自由移動之感測器隔離驗證，"
        "D1.5 為感測器掛於手臂之主結果；"
        "（五）端到端夾取三臂對照（D3：對位與附著升舉）；"
        "（六）側向資訊之系統性檢驗，以及二維多點定位之正式三臂對照（D2，見 5.4 節）；"
        "（七）貫穿各實驗的三臂對照與姿態／量測稽核。",
    ),
    (
        "Content",
        "排除項目包括："
        "商用超音波測距晶片（如 TDK CH201）之實機量測；"
        "夾爪物理摩擦夾持之部署級宣稱（詳見 5.3、6.2 節）；"
        "以及跨多組隨機種子的統計穩健性完整驗證。"
        "二維場景下「定位之後再夾取」仍屬未來工作（D2 已完成定位與二維閉環，見 5.4、6.3 節）。",
    ),
    (
        "Content",
        "限制方面："
        "採單一隨機種子與確定性物理引擎，尚未完成跨種子補跑；"
        "控制決策以單自由度前向接近為主（沿接近軸前進或停止），"
        "D2 雖含側向掃描定位，接近段仍由估距驅動前向運動；"
        "預定停止距離固定為 0.35 m；"
        "桌面高度目標在約 0.32 m 以內距離編碼失效，故不宣稱該近距內之閉環量測。"
        "控制器全程不讀取目標真實世界座標；全部結論僅在模擬中成立，不含實機部署評估。",
    ),
    ("Header2", "1.4 名詞解釋與研究貢獻"),
    (
        "Content",
        "英文術語與程式相關名稱於首次出現時附中文；其後以中文為主，必要時括注英文。"
        "常用縮寫與符號彙整如下（詳細操作定義見第三章）。"
        "GMO：Generic Model Output，RTX Acoustic 的原始輸出資料結構。"
        "WPM：Wave Propagation Model，光線追蹤式波傳播模型。"
        "SNR：本文之偵測訊噪比，指配對移除中目標貢獻相對量測雜訊底的比值（公式見 3.3 節），"
        "不是通訊工程的功率譜 SNR。"
        "RMSE：均方根誤差。r：皮爾森相關係數；ρ：斯皮爾曼等級相關係數。"
        "IK：逆向運動學，由末端目標位姿反解關節角。"
        "DOF：自由度。seed：隨機種子，決定目標位置序列。"
        "session：一次模擬器啟動至關閉。episode：一次完整試驗回合。"
        "standoff：預定停止距離，本文為 0.35 m。"
        "對位：夾取前夾爪中心與目標中心在水平方向是否落在預先鎖定容差內。"
        "多點定位：在多個已知位置各量一次距離，以幾何交會解座標"
        "（文獻見 Kapoor 等（2016）；合成孔徑譜系見 Hayes 等（2009））。",
    ),
    (
        "Content",
        "感測包絡：使目標可被穩定偵測的感測器—目標相對幾何範圍，以配對移除與 SNR 地圖界定。"
        "指標失效：僅看「到達率」等表面數字時，成功可能被走廊寬度等場景幾何撐起來，"
        "並非感測真的有用；本文改以停止位置相關性與盲走消融為主指標。"
        "盲走臂：量測管線與聲學臂相同，只把控制器可用的估距換成無資訊值，"
        "用以檢驗閉環是否真依賴聲學（消融邏輯見 Meyes 等（2019））。",
    ),
    (
        "Content",
        "貢獻一（感測包絡量測方法）：以配對移除與 SNR 地圖系統性界定 WPM 超音波感測器之可偵測幾何域"
        "（五十二格中三十六格可偵測），作為後續擺位與任務設計的依據。",
    ),
    (
        "Content",
        "貢獻二（聲學閉環接近主結果）：在已驗證包絡內，以三臂資訊消融對照顯示："
        "純聲學控制迴路可驅動 UR10e 接近隨機桌面目標；"
        "主結果 r＝0.9856、RMSE 2.8 cm、三十回合全數聲學觸發停止，九十回合姿態稽核零違規。",
    ),
    (
        "Content",
        "貢獻三（端到端夾取整合）：將任務由「接近後停止」延伸至「停止後夾取」；"
        "以聲學估距決定夾取位置，夾取中心與目標 r＝0.9885、對位 RMSE 1.9 cm，"
        "對位因果經消融對照確認（費雪精確檢定 p＝0.004）。",
    ),
    (
        "Content",
        "貢獻四（側向證偽與演算法恢復）：以四項獨立實驗證偽感測器原生左右資訊"
        "（能量差、時間差、身分欄位、接收器分組），"
        "再以手臂運動合成多視點之多點定位恢復二維座標，並完成正式三臂對照"
        "（側向 r＝0.950、二維停止 RMSE 1.9 cm、消融組失能 p＜0.001），"
        "示範「單軸測距能力 × 運動自由度」的補償路徑。",
    ),
    (
        "Content",
        "貢獻五（實驗效度設計）：提出三臂資訊消融對照（Meyes 等（2019））、"
        "預註冊判準（Nosek 等（2018））與量測／姿態稽核之組合，"
        "使模擬閉環結果的因果可歸屬，並可推廣為以模擬為基礎的機器人感測驗證流程。",
    ),
    ("Header2", "1.5 實驗平台與軟硬體環境"),
    (
        "Content",
        "為讓跨領域讀者先建立共同語境，本節以白話說明模擬平台、聲學模組與手臂／夾爪，並標註官方文件。"
        "聲學資料結構與場景尺寸見第三章；平台在文獻中的位置見第二章。"
        "本節只當「環境地圖」，不重複後續實驗結果。",
    ),
    (
        "Content",
        "NVIDIA Isaac Sim 6.0 是本研究所用的機器人模擬平台官方發行版（NVIDIA, 2026e）。"
        "它建立在 Omniverse（NVIDIA 的三維協作與模擬基礎平台）之上，"
        "整合 GPU 加速物理、光線追蹤渲染，以及 USD（Universal Scene Description，場景描述格式），"
        "並可載入工業機械手臂等官方資產，使運動與感測能在可重複、可程式化的虛擬場景中同時進行"
        "（平台綜述見 Gao 等（2026））。"
        "本研究採本機獨立執行，目的是在可控幾何與材質下驗證超聲感測回授接近，"
        "不宣稱與任一真實工廠現場等價。"
        "Isaac Lab 是同一生態系、偏強化學習訓練的延伸框架（NVIDIA, 2026d）；"
        "本文控制器為規則式、未做強化學習，Isaac Lab 僅列未來路徑（6.3 節）。",
    ),
    (
        "Content",
        "RTX Acoustic 是 Isaac Sim 內建的實驗性超音波感測模組（NVIDIA, 2026a）："
        "在虛擬場景中以 GPU 產生回波波形，供距離與偵測等特徵使用；輸出不是點雲，也不是深度圖。"
        "波傳播模型 WPM 屬幾何光線追蹤式路徑模擬（NVIDIA, 2026c），"
        "會與目標、桌面、手臂連桿等幾何交互，因此「什麼擺位讀得到目標」才有量測意義。"
        "原始輸出結構 GMO 以信號路徑序列記錄振幅（NVIDIA, 2026b）；欄位限制與重建細節見 3.1 節。"
        "重要邊界：不保證與商用晶片（如 TDK CH201）實機波形一致；"
        "結論皆為模擬內任務級、趨勢級表述（2.2、2.4、6.2 節）。",
    ),
    (
        "Content",
        "機械手臂為 Universal Robots 的 UR10e 六軸協作臂官方資產。"
        "依原廠規格，約為六個旋轉關節、工作半徑 1,300 mm、額定負載 12.5 kg 的 e-Series 機型"
        "（Universal, 2023），常見於桌面尺度工業與研究任務。"
        "本研究在模擬中載入資產並以逆向運動學驅動，不涵蓋實體控制器或實機安全認證。"
        "末端夾爪為 Robotiq 2F-85 二指自適應夾爪，最大開口約 85 mm（Robotiq, 2019）；"
        "與 UR10e 一併載入，但僅在夾取實驗（第五章 D3）啟用閉合與升舉。"
        "夾持相關限制見 5.3 與 6.2 節。",
    ),
    (
        "Content",
        "可記成一條鏈："
        "Isaac Sim 提供場景與物理 → RTX Acoustic（WPM／GMO）提供虛擬超聲觀測 → "
        "UR10e 提供六軸運動 → Robotiq 2F-85 在夾取段提供二指執行器。"
        "後續各章的包絡、閉環與消融都建在此鏈上。"
        "官方文件最新版以網址為準："
        "Isaac Sim（NVIDIA, 2026e）；RTX Acoustic（NVIDIA, 2026a）；"
        "GMO（NVIDIA, 2026b）；RTX 感測總覽（NVIDIA, 2026c）；"
        "UR10e（Universal, 2023）；2F-85（Robotiq, 2019）。",
    ),
]

# ---------------------------------------------------------------------------
# Chapter 3 (rewritten)
# ---------------------------------------------------------------------------


# 2026-07-13 CH3 深修（對齊第一、二章）：節首「本節回答」、拆長段、程式詞降級、飛行感測器改寫。
CH3 = [
    ("Header1", "第三章、研究方法"),
    ("Header2", "3.1 WPM 感測模型與 GMO 資料結構"),
    (
        "Content",
        "本章先交代模擬裡的超聲訊號從哪來、長什麼樣子，以及距離怎麼從波形算出來。"
        "RTX Acoustic 的波傳播模型（Wave Propagation Model, WPM）屬幾何光線追蹤式聲波模擬"
        "（NVIDIA, 2026c），不是只靠公式套衰減的參數化近場模型。"
        "WPM 會追蹤聲波與場景幾何（立方體目標、桌面、手臂連桿等）的反射交互，"
        "因此感測器—目標相對幾何一變，輸出就可能變——"
        "這也是第四章能用配對移除系統性畫「感測包絡」的技術前提。",
    ),
    (
        "Content",
        "WPM 的原始輸出稱為 Generic Model Output（GMO）。"
        "依官方格式（NVIDIA, 2026b），GMO 以信號路徑（signal way）為單位，"
        "記錄各發射器（TX）／接收器（RX）／通道組合下的振幅樣本序列。"
        "本研究所用組態為兩個聲學掛載點（相距 0.10 m）、單一接收群組，中心頻率 40 kHz。",
    ),
    (
        "Content",
        "把緩衝區還原成時間波形時，必須依「每條路徑的樣本數」把資料切開再接成一條時間軸；"
        "不能把通道編號欄位誤當成時間索引——否則峰值位置會壞掉（例如恆為零）。"
        "這是管線正確性的硬條件，細節屬實作層，正文只保留上述原則。",
    ),
    (
        "Content",
        "GMO 的時間偏移欄位（timeOffsetNs）在 Isaac Sim 6.0 恆為 0，不能拿來算飛行時間。"
        "距離改由：峰值樣本索引（多幀平均波形中，振幅絕對值最大的取樣點）"
        "× 樣本週期 × 聲速 ÷ 2，"
        "此即超音波時間飛行（Time-of-Flight, ToF）測距的標準形式"
        "（Zhmud 等（2018）；主動深度脈絡見 He 等（2019））。"
        "此外，發射／接收／通道身分在本組態下亦不區分接收器"
        "（每幀恆為 (0, 0, 0)）；第四章側向實驗以實測資料記錄此限制，而非只憑文件。"
        "場景建立後還有確定性啟動暫態，且逐幀能量常有小幅跳動，"
        "故量測一律多幀平均，並配合 3.3 節稽核剔除不穩定量測。",
    ),
    ("Header2", "3.2 樣本週期自校"),
    (
        "Content",
        "距離公式裡的「樣本週期」必須每批實驗當場自校，理由如下。"
        "樣本週期若用錯，全部距離估計會一起偏掉。"
        "本研究不沿用外部或舊批次常數，而在每一批新實驗中當場自校："
        "感測器固定，目標沿視軸放到 20 個已知距離（約 0.15–1.20 m）；"
        "每點先等 40 個模擬影格穩定，再取 24 影格平均波形讀峰值位置；"
        "以峰值位置對已知距離做最小平方法（OLS）迴歸，"
        "由斜率反推樣本週期（樣本週期 = 2 ÷（斜率 × 聲速））。"
        "斜率與截距即為該批實驗的距離預測模型："
        "之後任一段峰值位置，都以（峰值位置 − 截距）÷ 斜率換成距離。"
        "此「已知距離掃描自校」與機器人超聲應用中的校正脈絡一致（Zhmud 等（2018））。",
    ),
    (
        "Content",
        "自校以兩種幾何交叉檢查："
        "一是視軸（boresight，感測器指向軸）高度掃描；"
        "二是目標放在桌面高度的掃描（須另計垂直高度差）。"
        "兩者獨立得到的樣本週期分別約 103.09 µs 與 100.77 µs，"
        "接近引擎預設約 102.4 µs，顯示跨幾何大致一致。"
        "但後續換算仍一律採用當輪自校值，不假設固定常數可跨場景沿用。"
        "重要邊界：自校階段使用已知距離；進入控制回合後，目標真值只寫進記錄欄、不進控制分支"
        "（見 3.5 節與第五章）。",
    ),
    ("Header2", "3.3 實驗方法學：四支柱"),
    (
        "Content",
        "後續實驗共同遵守四條方法紀律，用以保證「成功」能歸因到聲學資訊。",
    ),
    (
        "Content",
        "支柱一：包絡優先（envelope-first）。"
        "先用配對移除與偵測訊噪比（SNR）地圖回答「在什麼幾何下讀得到目標」，"
        "再把手臂任務放進已驗證包絡——先驗證感測，再建控制。"
        "配對移除程序：同場景、同姿態，先量有目標波形，再物理移除目標量背景，"
        "並與緊接的重複量測（雜訊底）比較。"
        "動機是分離目標回波與背景多路徑"
        "（Valin 等（2017）；Tsuchiya 等（2022）；Liu 等（2020））。"
        "本文 SNR 定義為 "
        "max|W有目標 − W無目標| ÷ max|W有目標 − W雜訊參考|，"
        "其中 W 為多幀平均波形；SNR 大於 10 判為可偵測。"
        "此 SNR 是操作型偵測指標，不是通訊工程的功率譜 SNR。",
    ),
    (
        "Content",
        "支柱二：三臂對照。"
        "每一組閉環實驗都配三條臂，且共用同一組隨機種子與同一批目標位置："
        "聲學臂用估距做接近與停止；"
        "盲走臂量測管線完全相同，只把控制器可用的估距換成無資訊值（正無窮），"
        "使停止條件永不成立；"
        "開環臂完全不量測，只走固定行程。"
        "盲走屬資訊消融（ablation）：拿掉單一資訊通道、保留其餘管線，"
        "以檢驗該通道是否必要（Meyes 等（2019））。"
        "任何閉環宣稱都必須通過「盲走失能」：若盲走也差不多成功，"
        "就表示成功可能來自走廊幾何，而非聲學。",
    ),
    (
        "Content",
        "支柱三：預註冊判準（preregistration）。"
        "通過／失敗標準寫在執行規格中，執行前鎖定，"
        "避免看完結果再改門檻（資料窺探）；"
        "此對應開放科學中「先鎖分析計畫、再看結果」的原則（Nosek 等（2018））。"
        "主指標固定為停止位置與目標位置的相關 r 與誤差 RMSE，"
        "刻意不用單純到達率——在窄目標帶上，固定行程也可能「看起來常到達」"
        "（第六章以實測說明）。",
    ),
    (
        "Content",
        "支柱四：量測與姿態稽核。"
        "不合格量測標成無效並排除，不做事後插補。"
        "三類稽核："
        "（一）平穩性——場景啟動暫態與慢震盪，以穩定等待影格與事後漂移檢查剔除；"
        "（二）手臂姿態——逐步檢查連桿是否穿桌、穿地等非物理姿態；"
        "（三）感測器位姿——確認水平視軸與掛載高度等校正前提仍成立。"
        "第三類是所有距離校正結果有效的前提。",
    ),
    ("Header2", "3.4 場景建構與物件操作"),
    (
        "Content",
        "實驗場景與物件／手臂的移動方式如下，目的是維持聲學可比性。"
        "全部實驗共用同一基準場景，每次模擬器啟動以程式重建："
        "工作桌約 1.2 m × 0.8 m × 0.40 m，置於手臂基座前方（桌面中心距基座約 1.05 m）；"
        "UR10e 固定於原點，並載入 Robotiq 2F-85 夾爪（僅第五章 D3 真正啟用夾取動作）。"
        "超音波感測器有兩種掛法："
        "第四章與 D1 可為獨立放置、未掛臂的感測器物件；"
        "D1.5 與 D3 則掛於腕部連桿並前伸 0.25 m，以離開夾爪機構的聲學陰影。"
        "兩種情形下感測器均離地約 0.65 m、視軸水平朝前，並由位姿稽核逐步確認。"
        "目標為桌面上的長方體（尺寸見各實驗）；"
        "隨機位置由固定種子在預劃範圍內均勻抽取，三臂共用同一批目標。",
    ),
    (
        "Content",
        "在模擬中移動物件的方式是直接改寫空間位置，使下一影格出現在指定處，而非以物理力推動。"
        "此操作對聲學是否有效，已由前置探針驗證："
        "移動目標後，波形變化遠大於背景雜訊；"
        "「刪除目標」與「移到遠處」效果一致。"
        "配對移除中的「移除目標」即自場景刪除該物件。"
        "每次位置變動後的標準量測時序為：先等 40 影格消化暫態，"
        "再取 24 影格並拆成前後兩段各 12 影格平均；"
        "若兩段早期能量相差超過 5%，判定不平稳、該點無效。"
        "手臂移動則直接寫入關節角（運動學寫入），每步後同樣等待與量測；"
        "感測器隨臂移動的聲學有效性，另由第五章各實驗的啟動閘門探針確認。",
    ),
    ("Header2", "3.5 聲學特徵定義與閉環控制律"),
    (
        "Content",
        "從波形取出的特徵，以及如何變成手臂動作，定義如下。"
        "特徵一：峰值樣本索引——多幀平均波形中振幅絕對值最大的取樣點位置。"
        "物理意義是最強回波的到達時間，與往返距離成正比，故隨距離線性後移"
        "（ToF 標準解釋見 Zhmud 等（2018）；He 等（2019））。"
        "這是全部距離估計的依據；第四章亦顯示它對能量慢震盪相對免疫。"
        "特徵二：早期能量——平均波形前 20 個取樣點的振幅平方和，代表近程回波總強度。"
        "它只用於配對移除的 SNR（4.1 節）與平穩性稽核（3.4 節），"
        "不用來推距離：能量同時受距離、尺寸、材質與場景耦合，複雜場景不可靠。",
    ),
    (
        "Content",
        "閉環控制律以 D1.5（感測器掛於手臂）為例，因果鏈分五步："
        "（一）量測——依 3.4 節取得平均波形，讀峰值位置；"
        "（二）預測——用 3.2 節自校模型換成感測器至目標的三維距離，"
        "再以已知垂直高度差（約 0.19–0.20 m）換成水平距離估計；"
        "（三）決策——若水平估距 ≤ 預定停止距離 0.35 m 則停止，否則沿接近軸前進 0.05 m；"
        "（四）執行——以逆向運動學把末端目標解成六軸關節角並寫入，"
        "以上一步解為初值（暖啟動）並限制單步關節變化，避免多解之間亂跳；"
        "（五）稽核——檢查手臂姿態與感測器位姿後，回到（一）。"
        "另有兩道與聲學無關的安全護欄：最多 40 步，以及走廊端點強制停止；"
        "觸發時在記錄中與「聲學觸發停止」分開標示。"
        "目標真實世界座標只進記錄欄，不進任何控制判斷。",
    ),
    (
        "Content",
        "三臂差異只在步驟（二）（三），定義見 3.3 節支柱二："
        "聲學臂跑滿五步；盲走臂仍量測但估距作廢；開環臂完全不量測、固定行程。"
        "控制器保持最簡：固定步長前進、無濾波、無多步預測，"
        "使每一次停止都能追溯到單次量測的單一特徵。"
        "這是「可歸因」優先於「控制效能」的取捨；更強控制器列第六章未來工作。",
    ),
    ("Header2", "3.6 章節間資料鏈"),
    (
        "Content",
        "後續實驗一層接一層，而不是各自假設；資料鏈如下。"
        "資料鏈為："
        "S1 感測包絡界定可偵測幾何 →"
        "S2 在包絡內建立距離編碼量化關係 →"
        "D1 以未掛臂、可自由移動的感測器隔離驗證「感測＋控制」因果 →"
        "D1.5 把感測器掛上 UR10e，加入真實手臂運動學 →"
        "D3 延伸到夾取對位與附著升舉；"
        "另有 D2 在側向證偽後，以多點定位恢復二維並做三臂閉環（見 5.4 節）。"
        "每一層都有預先寫定判準；閉環層皆有消融對照。"
        "D1／D1.5 場景直接沿用 S1 包絡與 S2 校正；"
        "D3 更換目標物前，另以三道前置閘門重測可偵測性、測距與夾取動作力學，"
        "避免「換了物體就偷用舊結論」。",
    ),
    (
        "Content",
        "校正斜率的可信度另做歸因。"
        "不同批次斜率表面上可差約一成，但統計上落在小樣本迴歸不確定度內"
        "（點數有限，且峰值索引為整數造成量化）。"
        "直接對照實驗亦顯示：同一場景「移目標、感測器固定」與「手臂載感測器移動、目標固定」"
        "所得斜率無系統差。"
        "結論：斜率主要反映傳播與取樣機制，目標外形主要影響截距；"
        "同幾何下理論上可跨目標沿用，但本研究更換目標物時仍一律重做自校以保守。",
    ),
]


# ---------------------------------------------------------------------------
# Chapter 4 (rewritten)
# ---------------------------------------------------------------------------


# 2026-07-13 CH4 深修（對齊前三章）：節首「本節回答」、結果拆段、停損／datasheet／arm-free 白話化。
CH4 = [
    ("Header1", "第四章、感測特性化與包絡"),
    ("Header2", "4.1 S1 感測包絡地圖"),
    (
        "Content",
        "S1 對應第一章 RQ1：感測器在哪些幾何下讀得到目標，包絡邊界在哪。"
        "實驗代號 S1。設計為四因子配對移除掃描："
        "距離五個水準（0.15、0.3、0.5、0.8、1.2 m）、"
        "目標邊長三個水準（0.04、0.10、0.20 m）、"
        "俯仰角三個水準（0°、30°、60°，俯視為正）、"
        "干擾物三個水準（無干擾、僅桌面、桌面加手臂）。"
        "因子組合裁成 52 個量測格點，分 A–D 四區執行；全數完成，最終 0 格無效"
        "（少數格因平穩性未過而重測後通過）。",
    ),
    (
        "Content",
        "每一格獨立啟動一次模擬器，避免格點間狀態殘留。"
        "感測器固定離地 0.65 m，俯仰依格點設定；"
        "立方體目標放在視軸延伸線上指定距離（俯仰非零時目標隨視軸放低），"
        "使目標中心始終在指向正前方。"
        "「桌面」水準把 1.2×0.8 m 實心桌放在目標下方；"
        "「桌面加手臂」再把手臂以固定姿態立在感測器正後方約 0.1 m，"
        "模擬腕載時背後就是手臂本體的幾何。"
        "量測分三段，皆依 3.4 節標準程序（等 40 影格、取 24 影格平均）："
        "先量有目標；再於同條件緊接量一次作雜訊參考（兩者差＝該次工作階段的量測雜訊底）；"
        "最後刪除目標量無目標背景。訊噪比依 3.3 節公式由三段波形計算。",
    ),
    (
        "Content",
        "四區結果摘要如下。"
        "A 區（無干擾、水平指向，距離×尺寸）14/15 可偵測；"
        "唯一失格是 0.04 m 目標在 1.2 m，SNR 僅 3.9。"
        "B 區（俯仰 30°／60°）6/10 可偵測：30° 各距離皆可偵測，60° 僅最近 0.15 m 可偵測。"
        "C 區（含桌面，距離×俯仰）9/15 可偵測：水平指向全過（SNR 約 67–148）；"
        "30° 僅到約 0.5 m；60° 僅最近距離。"
        "D 區（桌面加手臂，距離×尺寸×俯仰）7/12 可偵測：水平指向全過（SNR 約 67–293）；"
        "60° 僅 0.3 m、0.20 m 目標可偵測。"
        "合計 36/52 格可偵測："
        "已否定「包絡完全不可行」的最壞情境，"
        "也未觸發預先寫定的中止條件「有手臂干擾就全數失敗」。",
    ),
    (
        "Content",
        "關鍵發現一：主宰因子是「指向」，不是「干擾物」。"
        "水平指向（0°）時，即使桌面與手臂都在，0.10 m 目標在 0.3–0.8 m 仍全可偵測（SNR ≥ 67）。"
        "反之俯仰到 60°，即使無干擾物，可偵測格也幾乎全滅。"
        "較合理的解釋是立方體近似鏡面反射：大俯角時回波偏離接收方向，"
        "而不是被桌、臂「擋住或吸掉」。",
    ),
    (
        "Content",
        "關鍵發現二：感測器後方物體對前向路徑的貢獻，在此確定性模擬下為零。"
        "C 區（僅桌）與 D 區（桌加臂）在相同距離與俯仰下，SNR 數值逐位相同"
        "（例如 148.36、94.58、67.67 兩區一致）。"
        "因此後方手臂對前向聲學路徑沒有可量測干擾。"
        "工程含義：抓取場景若偵測失敗，根因未必是桌臂本身，"
        "更常是擺位與指向不對——此點直接指導第五章場景設計。",
    ),
    (
        "Content",
        "關鍵發現三：目標尺寸有明確下限。"
        "0.04 m 目標超過約 1.2 m 即失效（SNR 掉到個位數）；"
        "0.10 m 在 0.15–1.2 m 皆可偵測；"
        "0.20 m 訊噪比最高（約 216–311）。"
        "後續量化特性與閉環實驗因此採用 0.10 m："
        "兼顧可偵測性與接近實際工件尺度。",
    ),
    ("Header2", "4.2 S2 感測器量化特性表"),
    (
        "Content",
        "S2 在已確認可偵測的包絡內，量化距離編碼精度、近距邊界與可重複性。"
        "實驗代號 S2。量測刻意採用 S1 的 D 區實戰幾何"
        "（桌面加手臂、水平指向、0.10 m 目標），"
        "讓後續控制器用的校正來自任務場景，而不是無干擾理想場景。"
        "內容含：視軸高度與桌面高度的距離掃描、垂直於視軸的側向掃描（13 點），"
        "以及關閉再開模擬器十次的重複性測試。",
    ),
    (
        "Content",
        "距離編碼結果。"
        "視軸高度（z＝0.65 m）三遍獨立掃描："
        "峰值位置與已知距離的皮爾森 r＝0.9994，RMSE 1.21 cm，自校樣本週期 103.09 µs。"
        "目標在桌面高度（z＝0.45 m，垂直差固定 0.20 m）："
        "r＝0.9998，在可用距離約 0.32 m 以上 RMSE 為 5.3 mm，自校週期 100.77 µs。"
        "0.15–0.26 m 近距點因俯視角超過約 50°，被平穩性稽核排除，符合幾何預期。"
        "三遍掃描結果逐位相同（含被排除點）；"
        "十次重啟重複性測試中，峰值位置完全相同，能量特徵變異係數為 0。"
        "這表示在固定種子下管線高度確定——"
        "它證明的是可重算、可稽核，不是「感測器很抗真實雜訊」。",
    ),
    (
        "Content",
        "校正斜率跨批次表面可差約一成，做了兩層歸因。"
        "統計層：小樣本批次的 95% 信賴區間仍涵蓋主校正值，差異落在估計不確定度內。"
        "實驗層：同一場景「移目標、感測器固定」與「手臂載感測器移動、目標固定」"
        "各量 13 點，斜率差小於合併標準誤的兩倍，量測方式本身無系統偏差。"
        "附帶發現：不同外形目標物解得的斜率相差不到百分之一——"
        "斜率主要反映傳播與取樣，外形主要影響截距。",
    ),
    ("Header2", "4.3 負結果與失效機制"),
    (
        "Content",
        "以下誠實記錄三類「做不成或沒有的東西」：單次輸出的左右資訊、頻率參數、腕載陰影；"
        "它們直接決定第五章設計與第六章宣稱邊界。",
    ),
    (
        "Content",
        "側向（左右）四重證偽——四項獨立實驗全數未過預設判準。"
        "（一）能量：兩接收器能量差與目標橫移（±0.15 m）無單調關係"
        "（斯皮爾曼 ρ＝0.357，遠低於預設 0.9）。"
        "（二）時間：兩路互相關時差幾乎恆定（與橫移的 r 僅約 0.002），且左右對稱，"
        "像是管線固定偏移，不是幾何方位。"
        "（三）身分：逐幀路徑標籤恆為 (0,0,0)，資料裡沒有分開的接收器身分。"
        "（四）分組：改成兩獨立接收群組後，第一路與原組態相同，第二路變成亂跳峰值，"
        "引擎不支援分組分離渲染。"
        "結論：缺的是引擎輸出層的側向資訊，不是量測太吵或演算法選錯；"
        "任何靠單次量測「左右差」的方法在此模擬器中都不可行。"
        "第五章因此改走「手臂移動多點測距」的多點定位路線。",
    ),
    (
        "Content",
        "中心頻率參數無效。"
        "同一場景用 20、30、40、60、80、100 kHz 做相同距離掃描："
        "峰值位置序列六組逐位相同；能量僅有約千萬分之一的浮點尾差，"
        "遠小於真實聲學「頻率加倍應出現的吸收／繞射變化」。"
        "故引擎未建模頻率相關物理。"
        "本文仍選 40 kHz，是因為它對應常見商用超音波頻段與實務距離尺度"
        "（波長約 8.6 mm，小於目標；量測距離約 0.4–1.2 m 在該頻段常見射程內），"
        "不是模擬內「掃頻選最優」的結果；真實頻段效應須待硬體驗證。",
    ),
    (
        "Content",
        "腕載聲影。"
        "同一目標在「含手臂與桌面」場景中，配對移除得到的目標貢獻低於約 2.2×10⁻⁵"
        "（相對背景能量約 135，已落入浮點雜訊）；"
        "在「無手臂」簡單場景中，同尺寸目標貢獻約 1.24×10⁵ 量級（訊號對背景約 6000:1）。"
        "落差主因是感測器貼在腕部、被夾爪網格幾何擋住而形成聲學陰影，並非該場景天生不能做聲學。"
        "這與 4.1「指向與擺位優先」互相印證，"
        "並直接決定第五章採用腕部前伸掛載，使感測器離開陰影區。",
    ),
]


# ---------------------------------------------------------------------------
# Chapter 5 (rewritten — new core result)
# ---------------------------------------------------------------------------


# 2026-07-13 CH5 深修：節首「本節回答」、去飛行感測器、5.4 中文標點、表5.2 標明首輪、數字不改。
CH5 = [
    ("Header1", "第五章、聲學閉環接近"),
    ("Header2", "5.1 D1：未掛臂感測器三臂對照（感測＋控制隔離驗證）"),
    (
        "Content",
        "D1 對應第一章 RQ2 的前半：不讀目標真實座標的純聲學控制，在排除手臂運動學後是否仍成立。"
        "若一開始就把感測器掛上 UR10e，估距誤差會與關節限制、IK 等交雜，難以歸因。"
        "因此先做 D1：感測器是可獨立放置、未掛在實體手臂上的物件，"
        "只檢驗「感測＋控制」這一環；通過後再於 5.2 節掛上手臂做 D1.5。",
    ),
    (
        "Content",
        "場景：0.10 m 方塊目標在桌面 x 介於 0.45–1.10 m 間隨機放置"
        "（避開第四章約 0.32 m 以內的桌高失效區，並保留停止距離裕度）。"
        "感測器離地 0.65 m、水平朝前，沿 X 軸逐步接近，步長 0.05 m，預定停止距離 0.35 m。"
        "每回合：抽目標位置並寫入場景；感測器回到起點 0.60 m；"
        "進入 3.5 節「量測—預測—決策—執行—稽核」迴圈——"
        "估距大於停止距離則前進，小於等於則停。"
        "聲學／盲走／開環三臂各 30 回合，同種子、同目標組；"
        "目標真值只進記錄欄，不進控制。"
        "正式實驗前先做 D0 閘門：目標固定，感測器沿軸移到 13 個已知位置量估距並迴歸，"
        "用來確認「移動感測器」本身聲學有效（此前只驗證過移目標）。"
        "D0 得 r＝0.9958，通過預設門檻 r≥0.99。",
    ),
    (
        "Content",
        "D1 結果。"
        "聲學臂：停止位置與目標 r＝0.9970，誤差平均 2.1 cm、RMSE 2.5 cm，"
        "30/30 誤差≤10 cm，且全數聲學觸發停止。"
        "盲走臂：估距被換成無資訊值後完全失能，恆停在走廊護欄 1.15 m，"
        "與目標無關，RMSE 79.3 cm，0/30 落入 10 cm。"
        "開環臂：固定行程停在 0.425 m，RMSE 16.9 cm，14/30 落入 10 cm。"
        "解讀：同一量測與控制管線，只要拿掉聲學資訊就完全失能——"
        "接近成功來自估距，不是走廊幾何碰巧造成。",
    ),
    ("Header2", "5.2 D1.5：手臂載具三臂對照（主結果）"),
    (
        "Content",
        "D1.5 把感測器掛上真實 UR10e，檢驗主結果是否仍成立。"
        "D1 通過後進入 D1.5（本文主結果）。"
        "感測器掛在腕部第三連桿（軟體名稱 wrist_3_link），並前伸 0.25 m，"
        "以離開夾爪機構的聲學陰影（見 4.3 節）。"
        "每步以逆向運動學（IK）把末端目標解成六軸角，"
        "以上一步解為初值（暖啟動）並限制單步關節變化，減少多解之間亂跳。",
    ),
    (
        "Content",
        "因重視非物理姿態（連桿穿桌、穿地），D1.5 每步做兩類稽核："
        "姿態稽核——前臂與腕部高度是否低於安全裕度，違規則整回合無效；"
        "感測器位姿稽核——水平視軸與掛載高度是否仍符合校正前提。"
        "三臂設計同 D1，各 30 回合、同種子同目標。"
        "正式前 D0.5 閘門（手臂沿固定路徑接近已知目標）r＝0.9918，"
        "姿態與位姿違規皆 0，通過預設閘門。",
    ),
    (
        "Content",
        "D1.5 主結果。"
        "聲學臂：r＝0.9856，誤差平均 2.5 cm、RMSE 2.8 cm，30/30 誤差≤10 cm，"
        "平均 5.0 步，全數聲學觸發停止。"
        "盲走臂：恆停護欄 0.95 m，RMSE 18.9 cm，僅 4/30 落入 10 cm。"
        "開環臂：固定行程停 0.80 m，RMSE 7.8 cm，22/30 落入 10 cm——"
        "看起來不差，但 r＝0，並未逐回合追蹤目標。"
        "聲學對盲走的停止誤差以 Welch t 檢定（不假設兩組變異數相等）得 t＝−10.6、p＜0.001"
        "（D1 對應 t＝−25.1、p＜0.001）。"
        "三臂共 90 回合全有效，姿態／位姿違規 0，IK 失敗 0。"
        "相對 D1：r 由 0.997 略降至 0.9856，RMSE 由 2.5 cm 略增至 2.8 cm——"
        "掛上真實手臂運動學後精度幾乎無損。",
    ),
    (
        "TableCaption",
        "表5.1  D1／D1.5 三臂對照（停止位置 vs 目標位置）",
    ),
    (
        "Table",
        ["實驗／臂", "r(停止, 目標)", "停止誤差 mean/RMSE", "P(誤差≤10 cm)", "終止方式"],
        [
            ["D1 聲學臂（未掛臂感測器）", "0.9970", "2.1 / 2.5 cm", "30/30", "聲學觸發"],
            ["D1 盲走臂", "0（恆停 1.15 m）", "77.6 / 79.3 cm", "0/30", "撞護欄"],
            ["D1 開環臂", "0（恆停 0.425 m）", "13.4 / 16.9 cm", "14/30", "固定行程"],
            ["D1.5 聲學臂（手臂載具）", "0.9856", "2.5 / 2.8 cm", "30/30", "聲學觸發"],
            ["D1.5 盲走臂", "0（恆停 0.95 m）", "17.4 / 18.9 cm", "4/30", "撞護欄"],
            ["D1.5 開環臂", "0（恆停 0.80 m）", "6.2 / 7.8 cm", "22/30", "固定行程"],
        ],
    ),
    ("Header2", "5.3 D3：端到端夾取整合"),
    (
        "Content",
        "D3 檢驗接近之後能否靠聲學估距完成對位夾取，以及夾持機制能宣稱到哪裡。"
        "D3 在停止後執行固定夾取序列："
        "（一）用停止當下估距一次推算目標位置（停止位置加估距）——"
        "因校正有效距離下限約 0.32 m，更近處不再聲學閉環；"
        "（二）抬到安全高度（指尖高於目標頂約 3 cm），以 2 cm 步長移到推算點上方；"
        "（三）下降到夾取高度；（四）閉合並偵測接觸；"
        "（五）附著後升舉 0.10 m（5 mm 小步），維持 60 影格後看目標升高是否≥0.05 m。"
        "全程不讀目標真值；真值只供事後評對位。"
        "目標改為截面 0.06 m、高 0.12 m 的直立長方柱。"
        "因換目標，正式前三道閘門皆過："
        "可偵測（三距離 SNR 31.9–82.3，皆＞10）、"
        "測距重校 r＝0.996、"
        "已知位置完整夾取序列十次姿態零違規。",
    ),
    (
        "Content",
        "夾持機制如實說明。"
        "開發中發現：模擬器指墊摩擦不足以在升舉中持續抓持"
        "（二十餘輪排查尺寸、摩擦、質量、閉合速度、高度等皆失敗），"
        "故改為接觸觸發之附著：閉合時若手指因物體擋住而未完全閉死"
        "（純物理接觸訊號，不含目標座標），則在夾爪與物體間建立剛性附著再升舉。"
        "對照區辨力仍在：盲走停錯位時抓空、手指閉死、不觸發附著即失敗。"
        "可宣稱「聲學對位＋接觸觸發附著」；不可宣稱「物理摩擦夾持」。",
    ),
    (
        "Content",
        "首輪三臂結果（各 30 回合，目標在 1.00–1.20 m）："
        "聲學臂夾取中心與目標 r＝0.9885，對位 RMSE 1.9 cm，"
        "對位率（±2 cm 預鎖容差）18/30（60%）；"
        "盲走與開環皆恆停名義點，對位率同為 7/30（23%）——"
        "約等於固定停點與隨機目標偶然落入容差的機率，r＝0。"
        "聲學對盲走對位率以費雪精確檢定（小樣本 2×2 精確機率）單尾 p＝0.004。"
        "對位成功者升舉 15/18（83%）。"
        "對位率與升舉率分開陳報，不相乘成總成功率，以免混淆聲學與力學。",
    ),
    (
        "Content",
        "首輪四項預設判準：三項過，一項不過——"
        "姿態全淨因 90 回合中 3 回合升舉階段 IK 無解。"
        "該 3 回合目標約在走廊最遠端 1.19 m，手臂已近可達極限，無法再上抬 0.10 m；"
        "接近段 130 控制步姿態與位姿稽核仍乾淨。"
        "此屬可達性邊界，與聲學無關。"
        "依預設判準字面如實判不通過，不放寬；"
        "修正方向為縮走廊遠端或降升舉高度。",
    ),
    (
        "Content",
        "複驗：走廊上限由 1.20 m 縮至 1.15 m"
        "（計入夾取進場點為估計值、可能比真目標再遠約 2 cm），"
        "全新目錄重跑三臂各 30 回合，原始數據保留。"
        "四項判準全過：對位 r＝0.978，對位率 80%（24/30）對盲走 33%"
        "（費雪 p＜0.001），升舉|對位 79%，90 回合零違規、無 IK 失敗。"
        "這是完整的失效—歸因—修正—驗證：不改判準，改被證實錯估的物理邊界後重跑。",
    ),
    (
        "Content",
        "對位成功但升舉失敗（首輪 3/18、複驗 5/24）逐回合檢視："
        "IK 皆正常；失敗多落在對位誤差靠近容差邊緣"
        "（複驗失敗平均 |誤差| 1.4 cm，成功 0.9 cm，且多為停短方向）。"
        "機制：統計容差 ±2 cm 略寬於夾爪物理捕捉窗約 ±1.5 cm，"
        "縫隙內「判準算對位成功」但咬合不穩、附著不成。"
        "屬容差與夾爪幾何議題，與聲學對位能力無關；"
        "分開陳報正是為了讓此縫隙可見。",
    ),
    (
        "TableCaption",
        "表5.2  D3 三臂對照（夾取對位與附著升舉；首輪 1.00–1.20 m 走廊）",
    ),
    (
        "Table",
        ["臂", "r(夾取中心, 目標)", "對位率（±2 cm）", "對位誤差 RMSE", "P(升舉|對位)"],
        [
            ["聲學臂", "0.9885", "18/30（60%）", "1.9 cm", "15/18（83%）"],
            ["盲走臂", "0（恆停名義點）", "7/30（23%）", "5.4 cm", "4/7"],
            ["開環臂", "0（恆停名義點）", "7/30（23%）", "5.3 cm", "5/7"],
        ],
    ),
    ("Header2", "5.4 D2：二維多點定位與閉環接近"),
    (
        "Content",
        "D2 接續第四章：單次輸出無左右資訊時，能否用手臂移動把側向找回來並閉環。"
        "感測器單次量測沒有左右，但手臂可以動——"
        "在多個已知位置各量一次距離，用幾何交會解二維座標"
        "（多點定位，見 Kapoor 等（2016）；"
        "移動感測器合成觀測基線的譜系見 Hayes 等（2009）；詳 2.7 節）。"
        "D2 程序：目標在縱向 1.00–1.18 m、側向 ±0.15 m 隨機；"
        "聲學臂先側向掃描五個量測位置（橫向基線 0.30 m），每點量一次距離，"
        "五筆距離以最小平方法做圓交會得 (x̂, ŷ)；"
        "再沿估計方位接近，估距小於預定停止距離即停。"
        "盲走臂掃描與接近流程相同，但估距在進定位前就換成無窮大——定位無解，只能直走護欄；"
        "開環臂不量測，走固定名義點。三臂各 30 回合、同種子配對。",
    ),
    (
        "Content",
        "D2 結果：四項預設判準全過。"
        "聲學臂定位：縱向 r＝0.979（RMSE 1.2 cm），側向 r＝0.950（RMSE 3.3 cm）——"
        "側向確由單軸測距經演算法恢復。"
        "二維停止誤差（平面距離相對預定停止距離的偏差）RMSE 1.9 cm，30/30 聲學觸發；"
        "盲走二維停止 RMSE 15.0 cm（Welch t＝−12.1、p＜0.001），開環 5.5 cm；"
        "90 回合（含側向移動步）姿態與位姿零違規。"
        "因果鏈因此閉合：第四章證明原生輸出無側向（負結果），"
        "本節證明運動合成可恢復並驅動二維閉環（正結果）——"
        "負結果界定問題，正結果回答問題，同一硬體限制下互為因果。",
    ),
    (
        "TableCaption",
        "表5.3  D2 三臂對照（二維定位與停止誤差）",
    ),
    (
        "Table",
        ["臂", "側向定位 r(ŷ, y)", "二維停止誤差 RMSE", "終止方式", "無效回合"],
        [
            ["聲學臂", "0.950（RMSE 3.3 cm）", "1.9 cm", "30/30 聲學觸發", "0"],
            ["盲走臂", "—（定位無解）", "15.0 cm", "30/30 撞護欄", "0"],
            ["開環臂", "—（不量測）", "5.5 cm", "固定停點", "0"],
        ],
    ),
    ("Header2", "5.5 盲走臂與開環臂的角色"),
    (
        "Content",
        "三臂裡的盲走與開環對照都不能省（定義見 3.3 節）。"
        "盲走檢驗的是：拿掉估距後同一管線是否失能（Meyes 等（2019））。"
        "若盲走也成功，聲學臂的成功就可能只是幾何巧合或量測節奏副作用；"
        "各閉環實驗中盲走失能，才把因果鎖在聲學資訊上。",
    ),
    (
        "Content",
        "開環是完全不量測的固定行程基準線。"
        "D1.5 目標帶寬僅約 0.30 m 時，開環 RMSE 7.8 cm、73% 回合誤差≤10 cm，看起來不差；"
        "但它只能停在一個固定點，無法逐回合追蹤隨機目標（r＝0）。"
        "聲學臂的優勢因此不只在誤差數字（2.8 cm 對 7.8 cm），"
        "而在相關性本質差異（0.9856 對 0）與到達一致性（100% 對 73%）："
        "只有聲學臂對每個隨機目標做出對應停止。"
        "此觀察是第六章 6.1 節「為何不能只看到達率」的主要論據。",
    ),
]


# ---------------------------------------------------------------------------
# Chapter 6 (rewritten)
# ---------------------------------------------------------------------------


# 2026-07-13 CH6 深修：節首「本節回答」、Claim boundary 中文化、清單可讀、與前五章口徑對齊。
CH6 = [
    ("Header1", "第六章、討論與限制"),
    ("Header2", "6.1 實驗效度設計之討論"),
    (
        "Content",
        "閉環評估為何必須用三臂與相關性，而不能只報告到達率，可由 D1.5 開環數據直接說明。"
        "這一點不必靠假設，D1.5 的開環臂數據就夠說明。"
        "開環臂完全不量測、只走固定行程；目標帶寬僅約 0.30 m 時，"
        "仍可達到 22/30（73%）「誤差≤10 cm」的到達率。"
        "若研究只以到達率當主指標，會嚴重高估閉環能力——"
        "因為成功大半來自走廊寬度與固定停點的幾何巧合，不是感測回授。"
        "開環臂停止位置與目標的相關 r＝0，正顯示它並未逐回合追蹤目標。",
    ),
    (
        "Content",
        "要判斷接近是否真由聲學驅動，至少同時看："
        "停止位置與目標的相關性（聲學臂 r＝0.9856，開環 r＝0），"
        "以及盲走消融是否失能（4/30 對聲學 30/30；三臂定義見 3.3 節）。"
        "以相關性／誤差而非到達率為主指標（預註冊見 Nosek 等（2018）；"
        "消融邏輯見 Meyes 等（2019）），是閉環宣稱能否成立的方法主幹。"
        "同一框架也可推廣到其他模擬感測回授研究，"
        "用以區分「真被感測驅動」與「幾何寬度撐出表面分數」。",
    ),
    ("Header2", "6.2 可宣稱與不可宣稱之範圍界定"),
    (
        "Content",
        "全文能講什麼、不能講什麼，收成下列兩張清單。"
        "可宣稱者如下。"
        "（一）純聲學閉環（全程不讀目標真實世界座標）可驅動 UR10e 接近隨機桌面目標；"
        "主結果停止 RMSE 2.8 cm、30 回合全數聲學觸發停止，"
        "並經三臂消融與逐步姿態稽核（範圍：模擬、單一隨機種子、確定性引擎）。"
        "（二）聲學閉環可延伸到端到端夾取對位："
        "夾取中心與目標 r＝0.9885、對位 RMSE 1.9 cm、對位率顯著優於盲走；"
        "對位成功後接觸觸發附著之升舉率約 83%（首輪陳報；複驗見 5.3 節）。"
        "（三）WPM 感測包絡可系統量測（52 格配對移除，36/52 可偵測），"
        "並建立距離編碼量化表（r＝0.9994；桌高目標 RMSE 5.3 mm）。"
        "（四）側向資訊可經手臂運動合成之多點定位自單軸測距恢復，"
        "且正式三臂已驗證（側向 r＝0.950、二維停止 RMSE 1.9 cm、盲走失能、90 回合零違規；5.4 節）。"
        "（五）對腕載聲影、側向四重證偽、頻率參數無效等失效機制有實證歸因，"
        "並提出可推廣的三臂效度設計。",
    ),
    (
        "Content",
        "不可宣稱者如下。"
        "（一）物理摩擦夾持：本文夾持為接觸觸發附著，摩擦保真度限制見 5.3 節，實體夾爪未驗證。"
        "（二）單次量測內的左右感知：4.3 節四項實驗已證偽；"
        "5.4 節側向是多次量測的演算法合成，不代表單次輸出本身含側向。"
        "（三）任何部署級或實機效能。"
        "（四）預定停止距離 0.35 m 以內的聲學閉環控制："
        "桌高目標約 0.32 m 以內距離編碼失效，夾取段採停止當下估距的一次推算。"
        "（五）跨隨機種子的統計穩健性：本文為單一種子、確定性引擎，尚未完成跨種子補跑。"
        "（六）二維定位之後再夾取的完整任務鏈：D2 已驗證定位與二維閉環接近，"
        "但側向 RMSE 3.3 cm 仍大於夾取物理容差約 ±1.5 cm（見 5.4、6.3 節）。",
    ),
    ("Header2", "6.3 未來工作"),
    (
        "Content",
        "主線完成後的自然延伸如下；皆應在新預註冊或實機協定下進行，而非事後放寬本文判準。",
    ),
    (
        "Content",
        "夾持與對位容差。"
        "聲學對位已驗證；首輪遠端 IK 可達性問題亦已由縮走廊複驗修正（5.3 節）。"
        "仍待補完者：其一，在更新版模擬器或實體夾爪上重驗摩擦夾持；"
        "其二，對位成功但仍有約二成因誤差落在 ±2 cm 統計容差與約 ±1.5 cm 物理捕捉窗縫隙而未能咬合，"
        "可收緊容差或加寬指墊捕捉窗——須在新預註冊下重跑，而非事後改舊判準。",
    ),
    (
        "Content",
        "二維夾取與動態目標。"
        "D2 已完成二維定位與閉環接近，但側向誤差 3.3 cm 仍大於夾取容差，"
        "故「定位後再夾取」需先壓低側向誤差。"
        "可能方向：增加視點、加大基線、或融合完整波形的學習模型"
        "（任何路線都應重新預註冊閘門，而非放寬舊門檻）。"
        "動態目標可把接近路徑本身當連續視點，以滑動窗口對最近數步距離做圓交會、邊走邊更新；"
        "此為純演算法升級，對照仍建議沿用三臂框架。",
    ),
    (
        "Content",
        "學習型控制與狀態估計。"
        "現行控制器是規則式固定步長前進，優點是透明、可歸因；"
        "不一定是最有效率或最泛化的策略。"
        "Isaac Lab（NVIDIA, 2026d）與大規模並行訓練方法（Rudin 等（2022））"
        "及 PPO（Schulman 等（2017））提供未來換學習控制器的基礎設施；"
        "換上之後仍須用本文三臂與稽核框架檢驗，"
        "避免策略學到場景捷徑而非使用聲學資訊。"
        "此外，各實驗已留存逐步波形與位姿（數百回合），"
        "可直接作監督式狀態估計（如回歸位姿、偵測目標是否存在）的訓練資料，"
        "無需額外模擬成本即可展開。",
    ),
    (
        "Content",
        "實機對應。"
        "後續可規劃對應 TDK CH201 等商用超音波測距等級的任務級驗證協定，"
        "檢驗包絡地圖與三臂對照結論在真實聲學、真實硬體與真實手臂動力學下能否轉移。"
        "重點是任務級指標與效度框架能否成立，而不是先追求波形與模擬逐點相同。",
    ),
]


# Apply numeric-citation replacement to the remaining chapters too (CH2_FIXED
# was already done above, at definition time). No author-year citations are
# currently known outside Chapter 2, but this keeps the mechanism general.
CH1 = _replace_citations_in_items(CH1)
CH3 = _replace_citations_in_items(CH3)
CH4 = _replace_citations_in_items(CH4)
CH5 = _replace_citations_in_items(CH5)
CH6 = _replace_citations_in_items(CH6)

# ---------------------------------------------------------------------------
# Mechanism (copied from rebuild_thesis_six_chapters.py)
# ---------------------------------------------------------------------------


def insert_paragraph_after(paragraph: Paragraph, style: str = "Content") -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.style = style
    return new_para


def remove_paragraph(paragraph: Paragraph) -> None:
    paragraph._element.getparent().remove(paragraph._element)


def set_text(p: Paragraph, text: str, style: str, center: bool = False) -> None:
    p.style = style
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        r._element.getparent().remove(r._element)
    if text:
        p.add_run(text)


def paragraph_after_element(parent, element) -> Paragraph:
    after = OxmlElement("w:p")
    element.addnext(after)
    return Paragraph(after, parent)


def append_items(doc: Document, anchor: Paragraph, items: list[tuple[Any, ...]]) -> Paragraph:
    prev = anchor
    for item in items:
        kind = item[0]
        if kind == "Table":
            _, headers, rows = item
            table = doc.add_table(rows=1 + len(rows), cols=len(headers))
            for j, h in enumerate(headers):
                table.rows[0].cells[j].text = h
            for i, row in enumerate(rows, start=1):
                for j, val in enumerate(row):
                    table.rows[i].cells[j].text = val
            prev._element.addnext(table._element)
            prev = paragraph_after_element(prev._parent, table._element)
            continue
        new_p = insert_paragraph_after(prev)
        if kind == "Header1":
            set_text(new_p, item[1], "Header1")
        elif kind == "Header2":
            set_text(new_p, item[1], "Header2")
        elif kind == "Content":
            set_text(new_p, item[1], "Content")
        elif kind == "Image":
            _, key, caption = item
            path = FIG[key]
            new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            from docx.shared import Inches

            new_p.add_run().add_picture(str(path), width=Inches(5.2))
            cap = insert_paragraph_after(new_p)
            set_text(cap, caption, "Content", center=True)
            prev = cap
            continue
        elif kind == "TableCaption":
            set_text(new_p, item[1], "Content", center=True)
        prev = new_p
    return prev


def paragraph_text(element) -> str:
    texts = []
    for node in element.iter():
        if node.tag.endswith("}t") and node.text:
            texts.append(node.text)
    return "".join(texts).strip()


def remove_body_between(doc: Document, start_marker: str, end_marker: str) -> None:
    body = doc.element.body
    start_el = end_el = None
    for child in list(body):
        if child.tag.endswith("}p"):
            text = paragraph_text(child)
            if text == start_marker and start_el is None:
                start_el = child
            if text == end_marker and start_el is not None:
                end_el = child
                break
    if start_el is None or end_el is None:
        raise RuntimeError("Cannot find body block")

    removing = False
    for child in list(body):
        if child is start_el:
            removing = True
            body.remove(child)
            continue
        if child is end_el:
            break
        if removing:
            body.remove(child)


def replace_body(doc: Document) -> None:
    start = end = None
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t == "第一章、緒論" and start is None:
            start = i
        if t == "參考文獻" and start is not None:
            end = i
            break
    if start is None or end is None:
        raise RuntimeError("Cannot find body block")

    anchor = doc.paragraphs[start - 1] if start > 0 else doc.paragraphs[0]
    remove_body_between(doc, "第一章、緒論", "參考文獻")

    body = CH1 + [("Header1", "第二章、文獻探討")] + CH2_FIXED + CH3 + CH4 + CH5 + CH6
    append_items(doc, anchor, body)

    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "參考文獻":
            nxt = doc.paragraphs[i + 1] if i + 1 < len(doc.paragraphs) else None
            if nxt is not None:
                set_text(nxt, REFERENCES_NUMBERED, "Content")
                # The v1 template stores each reference entry as its own
                # paragraph (one "w:p" per entry), not as one paragraph with
                # embedded line breaks; `nxt` is only the first of these.
                # Rewriting just `nxt` leaves the remaining old, unnumbered
                # entry paragraphs trailing after it (they are the last
                # thing in the document body), so remove them.
                for extra in list(doc.paragraphs[i + 2:]):
                    remove_paragraph(extra)
            break


def _norm(t: str) -> str:
    return t.replace(" ", "").replace("　", "").strip()


def replace_front_matter(doc: Document) -> None:
    """Replace 摘要/Abstract body paragraphs and keyword lines (V2 narrative).

    Locates Header1 "摘  要" / "Abstract" and replaces the immediately
    following Normal paragraph's text; keyword lines are matched by prefix.
    """
    paras = doc.paragraphs
    for i, p in enumerate(paras):
        t = _norm(p.text)
        if p.style.name == "Header1" and t == "摘要":
            if i + 1 < len(paras):
                set_text(paras[i + 1], ABSTRACT_ZH, "Normal")
        elif p.style.name == "Header1" and t == "Abstract":
            if i + 1 < len(paras):
                set_text(paras[i + 1], ABSTRACT_EN, "Normal")

    for p in doc.paragraphs:
        t = p.text.strip()
        if t.startswith("關鍵詞："):
            set_text(p, KEYWORDS_ZH, "Normal")
        elif t.startswith("Keywords:"):
            set_text(p, KEYWORDS_EN, "Normal")


def _toc_entries() -> list[tuple[str, str]]:
    """Derive the static TOC from the V2 body structure (the template's TOC
    paragraphs are plain text, not a Word field, so they must be rebuilt
    whenever the body changes — the v1 template's TOC still listed the OLD
    chapter structure and contained corrupted concatenated lines)."""
    entries: list[tuple[str, str]] = [
        ("toc 1", "誌  謝"),
        ("toc 1", "摘  要"),
        ("toc 1", "Abstract"),
    ]
    body = CH1 + [("Header1", "第二章、文獻探討")] + CH2_FIXED + CH3 + CH4 + CH5 + CH6
    for item in body:
        if item[0] == "Header1":
            entries.append(("toc 1", item[1]))
        elif item[0] == "Header2":
            entries.append(("toc 2", item[1]))
    entries.append(("toc 1", "參考文獻"))
    return entries


def _set_toc_text(p: Paragraph, text: str, style: str) -> None:
    """Hard-replace a TOC paragraph's content at the XML level.

    Template TOC paragraphs contain w:hyperlink elements; python-docx's run
    APIs do not clear runs nested inside hyperlinks, so naive set_text APPENDS
    after the surviving old text (this is also how the v1 template's TOC got
    its concatenated corrupted lines). Remove every child except w:pPr, then
    add a fresh run."""
    pel = p._p
    for child in list(pel):
        if not child.tag.endswith("}pPr"):
            pel.remove(child)
    p.style = style
    p.add_run(text)


def rebuild_toc(doc: Document) -> None:
    """Replace all existing toc-styled paragraphs with the V2 structure
    (page numbers are patched in a second pass from the rendered PDF)."""
    toc_paras = [p for p in doc.paragraphs if p.style.name in ("toc 1", "toc 2")]
    if not toc_paras:
        print("WARN: no toc paragraphs found; skipping TOC rebuild")
        return
    entries = _toc_entries()
    # Reuse existing toc paragraphs in place; add/remove to match count.
    n_reuse = min(len(toc_paras), len(entries))
    for p, (style, title) in zip(toc_paras[:n_reuse], entries[:n_reuse]):
        _set_toc_text(p, title, style)
    for p in toc_paras[n_reuse:]:
        remove_paragraph(p)
    anchor = toc_paras[n_reuse - 1]
    for style, title in entries[n_reuse:]:
        anchor = insert_paragraph_after(anchor, style)
        _set_toc_text(anchor, title, style)
    print(f"TOC rebuilt: {len(entries)} entries")


def _pdf_page_of(headings: list[str], pdf_path: Path) -> dict[str, int]:
    """Map each heading to the first PDF page containing it (monotonic scan).

    Two gotchas handled here (observed 2026-07-09, first pass patched only
    17/36): pdftotext page text contains newlines inside wrapped headings, so
    normalization must strip ALL whitespace; and the TOC page itself contains
    every heading, so pages holding many entry titles at once (>=8) are
    treated as TOC pages and excluded from the search."""
    import re
    import subprocess
    result = subprocess.run(
        ["pdftotext", "-layout", str(pdf_path), "-"],
        capture_output=True, text=True, check=True,
    )
    pages = result.stdout.split("\f")
    norm_pages = [re.sub(r"\s+", "", pg) for pg in pages]
    keys = [_norm(h) for h in headings]
    toc_like = set()
    for idx, pg in enumerate(norm_pages):
        hits = sum(1 for k in keys if k and k in pg)
        if hits >= 8:
            toc_like.add(idx)
    mapping: dict[str, int] = {}
    page_ptr = 0
    for h, key in zip(headings, keys):
        for idx in range(page_ptr, len(norm_pages)):
            if idx in toc_like:
                continue
            if key and key in norm_pages[idx]:
                mapping[h] = idx + 1
                page_ptr = idx
                break
    return mapping


def _convert_pdf() -> Path:
    import subprocess
    subprocess.run(
        ["soffice", "--headless", "--convert-to", "pdf",
         "--outdir", str(DOC_OUT.parent), str(DOC_OUT)],
        capture_output=True, text=True, check=True,
    )
    return DOC_OUT.with_suffix(".pdf")


def patch_toc_pages() -> None:
    """Second pass: render PDF, locate each TOC heading's page, write the
    page numbers back into the TOC lines (tab-separated, matching the
    template's toc style), then re-render the final PDF."""
    pdf_path = _convert_pdf()
    entries = _toc_entries()
    mapping = _pdf_page_of([t for _, t in entries], pdf_path)
    doc = Document(str(DOC_OUT))
    patched = 0
    for p in doc.paragraphs:
        if p.style.name in ("toc 1", "toc 2"):
            title = p.text.split("\t")[0].strip()
            if title in mapping:
                _set_toc_text(p, f"{title}\t{mapping[title]}", p.style.name)
                patched += 1
    doc.save(str(DOC_OUT))
    _convert_pdf()
    print(f"TOC page numbers patched: {patched}/{len(entries)} (final PDF regenerated)")


def main() -> None:
    if not DOC_IN.exists():
        raise FileNotFoundError(f"Input template not found: {DOC_IN}")
    doc = Document(str(DOC_IN))
    replace_front_matter(doc)
    # Body FIRST, then TOC: freshly rebuilt TOC lines carry no page suffix
    # yet, so their text exactly equals the chapter anchors replace_body
    # searches for — running rebuild_toc first makes replace_body anchor onto
    # the TOC block and splice the chapters there (observed 2026-07-09).
    replace_body(doc)
    rebuild_toc(doc)
    doc.save(str(DOC_OUT))
    print(f"Wrote V2 thesis draft: {DOC_OUT}")
    patch_toc_pages()


if __name__ == "__main__":
    main()
