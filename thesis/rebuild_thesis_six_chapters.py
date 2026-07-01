#!/usr/bin/env python3
"""Rebuild thesis docx: six chapters aligned with 2026-07-01 Physical AI summary."""

from __future__ import annotations

import re
import subprocess
import sys
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches
from docx.text.paragraph import Paragraph

THESIS = Path(__file__).resolve().parent
ROOT = THESIS.parent
DOC = THESIS / "THESIS_DRAFT_FCU_v1.docx"

FIG = {
    "3.1": THESIS / "figures/fig3_1_research_architecture.png",
    "4.1": ROOT / "runtime/outputs/phase3_rtx_pra_comparison_fixed_tcp_repeatability_v1/rtx_amplitude_max_vs_distance.png",
    "4.2": THESIS / "figures/fig4_2_rtx_early_energy_vs_distance.png",
    "4.4": THESIS / "figures/fig4_4_material_early_energy_0p5m.png",
}

ABSTRACT_ZH = (
    "本研究從應用電聲角度，探討 RTX Acoustic 超音波特徵能否在 Isaac Sim 可控幾何下支撐趨勢級距離推理，"
    "並進一步驅動機械手臂閉迴路接近與離線 Physical AI 狀態估計。"
    "研究平台為 NVIDIA Isaac Sim 6.0.0-rc.59 host standalone，Phase A 採 UR10 固定 TCP、移動目標設計，"
    "於 0.5–3.0 m 進行 6×5 次正式擷取，全部通過驗證（30/30 PASS）；"
    "primary_sgw_early_energy 優於飽和振幅峰值作為距離代理（Spearman ρ≈−0.66, n=6）。"
    "Phase B/C 以 UR10e + Robotiq 2F-85 與腕部 RTX 感測器建構閉環接近管線，"
    "於隨機化目標位姿之 v9 資料集上，閉環組 approach ≤0.45 m 達 84.0%（25 trials），"
    "顯著優於 open-loop baseline 之 29.2%（24 trials）；"
    "離線 stop_region 分類中，acoustic_only 特徵 F1≈0.598，顯示聲學觀測含可測量狀態信號。"
    "最終夾取成功率約 20%（兩組相近），瓶頸在下游 PhysX 接觸／夾爪整合，非聲學接近本身。"
    "本研究貢獻為可審計模擬可行性管線與明確 claim boundary，非部署級測距或穩定端到端夾取系統。"
)

ABSTRACT_EN = (
    "This thesis examines, from an applied electro-acoustics perspective, whether RTX Acoustic ultrasonic "
    "features support trend-level distance reasoning and closed-loop robotic approach in Isaac Sim. "
    "Phase A uses a fixed-TCP, moving-target protocol on a UR10 end-effector; all 30 formal captures "
    "(6 distances × 5 repeats) passed validation, with primary_sgw_early_energy outperforming saturated "
    "amplitude peaks (Spearman ρ≈−0.66, n=6). Phase B/C builds a UR10e + Robotiq closed-loop pipeline "
    "with wrist-mounted RTX sensing. On the randomized v9 dataset, closed-loop approach within 0.45 m "
    "reached 84.0% (25 trials) versus 29.2% for an open-loop baseline (24 trials). Offline stop-region "
    "classification with acoustic-only features achieved F1≈0.598. Final grasp success remained near 20% "
    "in both modes, indicating contact/gripper integration—not acoustic approach—as the limiting stage. "
    "The contribution is an auditable feasibility pipeline with explicit claim boundaries, not "
    "deployment-grade ranging or stable end-to-end grasping."
)

KEYWORDS_ZH = "關鍵詞：RTX Acoustic；閉迴路接近；Physical AI；early energy；Isaac Sim；可重複性"
KEYWORDS_EN = "Keywords: RTX Acoustic; closed-loop approach; Physical AI; early energy; Isaac Sim; repeatability"

sys.path.insert(0, str(THESIS))
from build_chapter2_docx import CHAPTER2, REFERENCES  # noqa: E402

CH2_FIXED = []
for item in CHAPTER2:
    if item[0] == "Content" and "PyRoom" in item[1] and "趨勢對照" in item[1]:
        continue
    if item[0] == "Content" and "智能製造" in item[1]:
        text = item[1].replace("呼應本學程之智能製造脈絡", "呼應電聲與機器人應用交叉之模擬驗證需求")
        CH2_FIXED.append((item[0], text))
    elif item[0] == "Content" and "Zhou" in item[1] and "智能製造" in item[1]:
        text = item[1].replace("呼應智能製造數位驗證需求", "支撐機器人應用場景之模擬驗證")
        CH2_FIXED.append((item[0], text))
    else:
        CH2_FIXED.append(item)

CH1 = [
    ("Header1", "第一章、緒論"),
    ("Header2", "1.1 研究背景與問題意識"),
    (
        "Content",
        "室內主動聲學系統中，多徑傳播與殘響使接收信號同時承載幾何、材質與距離資訊。"
        "在機器人末端非視覺 last-meter 接近中，如何從可控模擬條件擷取可重現聲學觀測，"
        "並據以驅動閉環運動與離線狀態判斷，是應用電聲與機器人感知之交叉問題。"
        "NVIDIA Isaac Sim 6.0 提供 RTX Acoustic 實驗性模組，可輸出 Generic Model Output（GMO）signal-way 序列，"
        "為系統化檢驗上述問題提供工具。",
    ),
    (
        "Content",
        "與相機 + 視覺語言模型（VLM）端到端操作相比，本研究聚焦非視覺超聲閉環接近，"
        "定位為互補而非取代：VLM 擅長語義與粗定位，超聲擅長走廊內距離趨勢接近。",
    ),
    ("Header2", "1.2 研究目的與問題陳述"),
    ("Content", "RQ1（Phase A）：固定 TCP 下 RTX 特徵是否可重現且具距離趨勢？"),
    ("Content", "RQ2（Phase B）：閉環控制器能否改善目標區到達率（相對 open-loop）？"),
    ("Content", "RQ3（Phase B/C）：離線 Physical AI 是否含可測量聲學狀態信號？"),
    ("Content", "RQ4（Phase C，限制）：Tier B 接觸級夾取是否趨勢級可行？"),
    ("Header2", "1.3 研究範圍與限制"),
    (
        "Content",
        "納入：Isaac Sim 6.0 host standalone、UR10/UR10e 官方資產、六面牆房間、"
        "Phase A 30/30 特徵實驗、Phase B/C 閉環接近與 v9 隨機化 Physical AI 資料集、"
        "Tier B contact-only（--skip-lift）夾取評估。"
        "排除：CH201 實機量測、波形級跨模擬器等價、穩定物理抬升、可部署學習控制器。",
    ),
    (
        "Content",
        "限制：模擬 only；experimental API；控制器不讀目標世界座標作前進決策，"
        "但 supervisor 可用 oracle 距離作安全包絡；最終夾取成功率約 20%，應歸因 PhysX 接觸整合。",
    ),
    ("Header2", "1.4 名詞解釋與研究貢獻"),
    (
        "Content",
        "Passport：可版本化之幾何／材質／任務實驗護照。"
        "trend-level：趨勢級，允許單調相關但不宣稱部署精度。"
        "claim boundary：可宣稱與不可宣稱之對照邊界（詳表6.1）。",
    ),
    ("Content", "貢獻一（主）：Geometry/Material Passport + 30/30 可審計 RTX 特徵管線（Phase A）。"),
    ("Content", "貢獻二（次）：超聲閉環接近 + supervisor v1；v9 上 84% vs 29% 區域到達改善（Phase B）。"),
    ("Content", "貢獻三（次）：隨機化 Sim 下 Physical AI 狀態估計基線（acoustic_only F1≈0.598）。"),
    ("Content", "貢獻四（附）：Tier B contact-only 示範與階段化評估框架；夾取瓶頸列為限制（Phase C）。"),
]

CH3 = [
    ("Header1", "第三章、研究方法"),
    ("Header2", "3.1 研究架構"),
    (
        "Content",
        "本研究採 Phase A→B→C 三階架構（圖3.1）："
        "Phase A 驗證特徵可審計性；Phase B 驗證閉環接近；"
        "Phase C 評估 Tier B 接觸與離線 Physical AI 狀態估計。"
        "Isaac Lab 動態觀測與 SL/RL 示範保留為附錄延伸，非主貢獻。",
    ),
    ("Image", "3.1", "圖3.1  研究架構（Phase A→B→C + Physical AI）"),
    ("Header2", "3.2 實驗平台與可重現執行"),
    (
        "Content",
        "Isaac Sim 6.0 host standalone（6.0.0-rc.59），以 scripts/env_host_isolated.sh 隔離執行。"
        "Phase A 使用官方 UR10；Phase B/C 使用 UR10e + Robotiq 2F-85，RTX 感測器掛載於腕部連桿。",
    ),
    ("Header2", "3.3 Phase A：fixed_tcp_moving_target"),
    (
        "Content",
        "固定 TCP（xy 半徑≈0.816 m、z≈0.65 m），僅移動 8 cm×8 cm×2 cm 目標板；ee_link 運動量 0 m。"
        "距離 waypoints：0.5–3.0 m（6 點）。設計目的為隔離幾何變量，使重複性與趨勢分析可解釋。",
    ),
    ("Header2", "3.4 Passport 體系"),
    (
        "Content",
        "Geometry Passport v1.0：房間、感測器掛載、目標軌跡。"
        "Material Passport v1.0：NonVisualMaterial 條件 A/B/C。"
        "Grasp Passport v1.0：搜尋走廊、standoff、UR10e/Robotiq 任務幾何。",
    ),
    ("Header2", "3.5 RTX GMO 擷取與 signal-way 特徵"),
    (
        "Content",
        "以 Replicator Writer 擷取 GMO；rtx_acoustic_factory.py 解析 signal-way，"
        "萃取 primary_sgw_early_energy、TOF、雙 RX 平衡等特徵。",
    ),
    ("Header2", "3.6 Phase B：UltrasonicClosedLoopController"),
    (
        "Content",
        "純 Python 狀態機（ultrasonic_closed_loop_controller.py），"
        "融合 early_energy、TOF 與對齊分數估計距離趨勢並驅動逐步接近。"
        "控制器不消費目標世界座標作前進決策；oracle 距離僅供評估記錄。",
    ),
    ("Header2", "3.7 Phase C：supervisor、contact-only 與隨機化協定"),
    (
        "Content",
        "approach_supervisor_v1.py 於 fusion 飽和或前進上限時以 oracle 距離作安全仲裁，"
        "建議進入 standoff／夾取階段。"
        "預設 --skip-lift（FixedCuboid）以隔離 PhysX lift 不穩定。"
        "v8/v9 管線隨機化 search start 與 wrench 橫向位置，打破 sensor_x/y 捷徑。",
    ),
    ("Header2", "3.8 評估指標與 claim boundary"),
    (
        "Content",
        "Phase A：30/30 PASS、Spearman ρ、跨 repeat CV。"
        "Phase B/C：approach/near/final 階段成功率、closed-loop vs open-loop 對照、"
        "Physical AI ablation（F1、balanced accuracy）。"
        "全書以表6.1 集中界定可宣稱與不可宣稱項目。",
    ),
]

CH4 = [
    ("Header1", "第四章、特徵驗證結果（Phase A）"),
    ("Header2", "4.1 可重複性與管線可行性"),
    (
        "Content",
        "正式實驗 30/30 PASS；gmo_valid_rate=1.0；max_ee_motion_m=0。"
        "2.0–3.0 m 區間 amplitude_max 與 early_energy 之跨 repeat CV 近乎 0，顯示擷取管線穩定。",
    ),
    ("Header2", "4.2 距離趨勢"),
    (
        "Content",
        "材質 B、5 repeats 平均：amplitude_max 於 1.0 m 後飽和；"
        "primary_sgw_early_energy 與距離 Spearman ρ≈−0.66（n=6），優於 peak 作為距離 proxy。",
    ),
    ("Image", "4.1", "圖4.1  RTX amplitude_max 與目標距離（材質 B）"),
    ("Image", "4.2", "圖4.2  RTX primary_sgw_early_energy 與目標距離（材質 B）"),
    ("Header2", "4.3 材質敏感度（A/B/C）"),
    (
        "Content",
        "primary_sgw_peak 於 0.5 m 與 3.0 m 幾乎無差；"
        "early_energy @0.5 m 可區分 C（≈190.6）與 A/B（≈165.4），顯示材質吸收對早期能量之影響。",
    ),
    ("Image", "4.4", "圖4.3  材質 A/B/C 於 0.5 m 之 early_energy 比較"),
    ("Header2", "4.4 穩健性驗證（P1）"),
    (
        "Content",
        "P1 smoke 驗證 rtx_acoustic_factory 對 GMO 之結構檢查：signal-way 維度、ACOUSTIC 模態、"
        "primary/ref/all peak 一致性與非平坦波形條件。P1 確保特徵萃取前資料有效。",
    ),
]

CH5 = [
    ("Header1", "第五章、閉環接近、Physical AI 與夾取評估（Phase B/C）"),
    ("Header2", "5.1 實驗協定與資料集"),
    (
        "Content",
        "Canonical 資料集：physical_ai_v9_skip_lift_clean（49 trial 目錄、284 step rows；"
        "closed_loop 25 trials、open_loop_baseline 24 trials）。"
        "接觸模式：--skip-lift / FixedCuboid，避免 PhysX lift 污染。"
        "控制器：closed_loop 使用超聲融合距離估計；open_loop_baseline 作對照。",
    ),
    ("Header2", "5.2 閉環 vs open-loop 接近成功率"),
    (
        "Content",
        "階段化審計顯示：閉環組在 approach ≤0.45 m 與 near ≤0.35 m 均達 84.0%，"
        "open-loop 分別僅 29.2% 與 4.2%。"
        "最終 success 兩組皆約 20%，表明聲學閉環主要改善區域到達，而非下游夾取執行。",
    ),
    ("TableCaption", "表5.1  隨機化 v9 階段成功率（contact-only）"),
    (
        "Table",
        ["指標", "Closed-loop", "Open-loop"],
        [
            ["Approach ≤ 0.45 m", "84.0% (21/25)", "29.2% (7/24)"],
            ["Near ≤ 0.35 m", "84.0% (21/25)", "4.2% (1/24)"],
            ["Final success", "20.0% (5/25)", "20.8% (5/24)"],
        ],
    ),
    ("Header2", "5.3 規則監管員 v1"),
    (
        "Content",
        "當融合距離飽和（約 0.73 m）而 tool0 已達前進上限時，supervisor 以 oracle 距離判斷是否足夠接近，"
        "建議進入 standoff／夾取。此設計屬安全包絡，非逐步神經策略；"
        "論文 claim 明確區分「控制器不讀 target pose」與「系統層安全仲裁」。",
    ),
    ("Header2", "5.4 離線 Physical AI 特徵消融"),
    (
        "Content",
        "以 physical_ai_v9_skip_lift_clean 步級資料訓練 logistic regression / decision tree 基線。"
        "stop_region_label 上，all_features 最佳（F1=0.684）；"
        "acoustic_only 達 F1=0.598，優於 pose_only（F1=0.533），顯示聲學特徵含狀態資訊；"
        "pose 仍為混淆因子。結果支持離線狀態估計可行性，非可部署控制器。",
    ),
    ("TableCaption", "表5.2  stop_region_label 特徵消融（v9）"),
    (
        "Table",
        ["Feature set", "F1", "Balanced accuracy"],
        [
            ["All features", "0.684", "0.665"],
            ["Acoustic only", "0.598", "0.590"],
            ["Pose only", "0.533", "0.650"],
        ],
    ),
    ("Header2", "5.5 Tier B 夾取與限制"),
    (
        "Content",
        "contact-only 模式下，PhysX 非有限態多來自 Robotiq 接觸物理或誤入 lift 路徑，"
        "已以 --skip-lift 明確旗標修復 headless 建立 DynamicCuboid 之 bug。"
        "SurfaceGripper 官方整合嘗試未成功（runtime 註冊失敗），列為限制與未來工作。",
    ),
    ("Header2", "5.6 與 VLM 端到端路線之討論"),
    (
        "Content",
        "本研究不與 VLM 全任務成功率硬比較，而強調 last-meter 非視覺閉環之互補角色。"
        "未來可探索 VLM 粗定位 + 超聲精接近之 hybrid 架構。",
    ),
]

CH6 = [
    ("Header1", "第六章、結論與建議"),
    ("Header2", "6.1 研究結論"),
    (
        "Content",
        "（1）RQ1：固定 TCP 下 RTX 聲學管線具高可重複性（30/30），early_energy 為有效距離 proxy。"
        "（2）RQ2：閉環超聲接近顯著改善隨機化場景下目標區到達率（84% vs 29%）。"
        "（3）RQ3：離線 Physical AI 顯示 acoustic_only 含可測量狀態信號（F1≈0.598）。"
        "（4）RQ4：最終夾取約 20%，應列為下游 PhysX／夾爪限制，非否定聲學主線。",
    ),
    ("Header2", "6.2 綜合討論"),
    (
        "Content",
        "電聲觀點：室內多徑使距離資訊體現在早期能量之弱趨勢；飽和 peak 不適合作距離特徵。"
        "控制觀點：閉環融合 + 走廊幾何可將趨勢級特徵轉為可重現接近行為。"
        "Physical AI 觀點：隨機化協定必要，以避免 pose 捷徑；多特徵優於單一閾值。"
        "操作觀點：接觸／抬升整合仍是 Sim2Real 前主要瓶頸。",
    ),
    ("Header2", "6.3 研究限制與 claim boundary"),
    ("TableCaption", "表6.1  Claim boundary 總表"),
    (
        "Table",
        ["範疇", "可宣稱", "不可宣稱"],
        [
            [
                "Phase A",
                "30/30；early_energy 距離趨勢（ρ≈−0.66, n=6）",
                "厘米級測距；波形等價",
            ],
            [
                "Phase B",
                "閉環 approach ≤0.45 m：84% vs open-loop 29%",
                "純超聲端到端夾取；零 oracle 系統",
            ],
            [
                "Physical AI",
                "acoustic_only stop_region F1≈0.598",
                "可部署學習控制器",
            ],
            [
                "Phase C",
                "Tier B contact-only 階段化評估",
                "穩定最終夾取；部署級 grasp rate",
            ],
            [
                "整體",
                "可審計 Sim 可行性管線",
                "CH201 實機已驗證；優於 VLM 全任務",
            ],
        ],
    ),
    ("Header2", "6.4 後續建議"),
    (
        "Content",
        "建議：（1）CH201 實機 task-level 協定；（2）SurfaceGripper isolated smoke 後再整合 UR10；"
        "（3）擴充隨機化資料並分離 approach/contact/lift 指標；"
        "（4）VLM + 超聲 hybrid；（5）以 replication package 支援口試重現。",
    ),
    ("Header2", "附錄說明（Isaac Lab 延伸）"),
    (
        "Content",
        "Isaac Lab 動態 smoke（ρ≈−0.48）、Sim→Lab SL（r≈0.47）、in-sim RL 閉環示範已實作，"
        "保留為附錄或未來工作，不計入本版主貢獻。",
    ),
]

TEMPLATE_PAT = re.compile(r"^\(.*標楷體|^\(.*粗體|^\(22,|請在此寫入|目錄頁內文")


def insert_paragraph_after(paragraph: Paragraph, style: str = "Content") -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.style = style
    return new_para


def remove_paragraph(paragraph: Paragraph) -> None:
    paragraph._element.getparent().remove(paragraph._element)


def set_text(p: Paragraph, text: str, style: str, center: bool = False) -> None:
    p.style = style
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        r._element.getparent().remove(r._element)
    if text:
        p.add_run(text)


def paragraph_after_element(parent, element) -> Paragraph:
    after = OxmlElement("w:p")
    element.addnext(after)
    return Paragraph(after, parent)


def append_items(doc: Document, anchor: Paragraph, items: list[tuple[Any, ...]]) -> Paragraph:
    prev = anchor
    for item in items:
        kind = item[0]
        if kind == "Table":
            _, headers, rows = item
            table = doc.add_table(rows=1 + len(rows), cols=len(headers))
            for j, h in enumerate(headers):
                table.rows[0].cells[j].text = h
            for i, row in enumerate(rows, start=1):
                for j, val in enumerate(row):
                    table.rows[i].cells[j].text = val
            prev._element.addnext(table._element)
            prev = paragraph_after_element(prev._parent, table._element)
            continue
        new_p = insert_paragraph_after(prev)
        if kind == "Header1":
            set_text(new_p, item[1], "Header1")
        elif kind == "Header2":
            set_text(new_p, item[1], "Header2")
        elif kind == "Content":
            set_text(new_p, item[1], "Content")
        elif kind == "Image":
            _, key, caption = item
            path = FIG[key]
            new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            new_p.add_run().add_picture(str(path), width=Inches(5.2))
            cap = insert_paragraph_after(new_p)
            set_text(cap, caption, "Content", center=True)
            prev = cap
            continue
        elif kind == "TableCaption":
            set_text(new_p, item[1], "Content", center=True)
        prev = new_p
    return prev


def cleanup_front_matter(doc: Document) -> None:
    to_remove = []
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if TEMPLATE_PAT.search(t) or t in ("(20, 粗體, 置中, 字與字之間空一字元)",):
            to_remove.append(i)
    for i in reversed(to_remove):
        remove_paragraph(doc.paragraphs[i])

    for p in doc.paragraphs:
        t = p.text.strip()
        if t.startswith("本研究評估") or t.startswith("本研究從應用電聲"):
            set_text(p, ABSTRACT_ZH, "Normal")
        elif t.startswith("This study evaluates") or t.startswith("This thesis examines"):
            set_text(p, ABSTRACT_EN, "Normal")
        elif t.startswith("關鍵詞："):
            set_text(p, KEYWORDS_ZH, "Normal")
        elif t.startswith("Keywords:"):
            set_text(p, KEYWORDS_EN, "Normal")

    has_student = any("研究生" in p.text for p in doc.paragraphs)
    if not has_student:
        for p in doc.paragraphs:
            if p.text.strip() == "指導教授：蔡鈺鼎 教授":
                sp = insert_paragraph_after(p, "Normal")
                set_text(sp, "研究生：（請填入姓名）", "Normal")
                break


def paragraph_text(element) -> str:
    texts = []
    for node in element.iter():
        if node.tag.endswith("}t") and node.text:
            texts.append(node.text)
    return "".join(texts).strip()


def remove_body_between(doc: Document, start_marker: str, end_marker: str) -> None:
    body = doc.element.body
    start_el = end_el = None
    for child in list(body):
        if child.tag.endswith("}p"):
            text = paragraph_text(child)
            if text == start_marker and start_el is None:
                start_el = child
            if text == end_marker and start_el is not None:
                end_el = child
                break
    if start_el is None or end_el is None:
        raise RuntimeError("Cannot find body block")

    removing = False
    for child in list(body):
        if child is start_el:
            removing = True
            body.remove(child)
            continue
        if child is end_el:
            break
        if removing:
            body.remove(child)


def replace_body(doc: Document) -> None:
    start = end = None
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t == "第一章、緒論" and start is None:
            start = i
        if t == "參考文獻" and start is not None:
            end = i
            break
    if start is None or end is None:
        raise RuntimeError("Cannot find body block")

    anchor = doc.paragraphs[start - 1] if start > 0 else doc.paragraphs[0]
    remove_body_between(doc, "第一章、緒論", "參考文獻")

    body = CH1 + [("Header1", "第二章、文獻探討")] + CH2_FIXED + CH3 + CH4 + CH5 + CH6
    append_items(doc, anchor, body)

    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "參考文獻":
            nxt = doc.paragraphs[i + 1] if i + 1 < len(doc.paragraphs) else None
            if nxt is not None:
                set_text(nxt, REFERENCES, "Content")
            break


def ensure_figures() -> None:
    subprocess.run([sys.executable, str(THESIS / "generate_thesis_figures.py")], check=True)
    subprocess.run([sys.executable, str(THESIS / "generate_fig31.py")], check=True)
    for k, p in FIG.items():
        if not p.exists():
            raise FileNotFoundError(f"Missing figure {k}: {p}")


def main() -> None:
    ensure_figures()
    doc = Document(str(DOC))
    cleanup_front_matter(doc)
    replace_body(doc)
    doc.save(str(DOC))
    print(f"Rebuilt six-chapter thesis: {DOC}")


if __name__ == "__main__":
    main()