#!/usr/bin/env python3
"""Streamline thesis Chapter 2 §2.2–§2.4: Gao (2026) platform focus, trim acoustics jargon."""

from __future__ import annotations

import re
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph

THESIS = Path(__file__).resolve().parent
SRC = Path("/home/lab109/下載/2THESIS_DRAFT_FCU_v1_no_pyroom.docx")
OUT_PATHS = [
    THESIS / "THESIS_DRAFT_FCU_v1.docx",
    THESIS / "2THESIS_DRAFT_FCU_v1.docx",
    SRC,
]

SECTION_22 = [
    (
        "Gao 等（2026）為第一篇專題綜述 NVIDIA Isaac Sim，指出模擬已成為機器人研究的核心基礎設施。"
        "平台以 Omniverse 為底，整合 GPU 加速 PhysX 物理、RTX 光追渲染與 USD 場景表示，"
        "並內建 RGB／深度、LiDAR、IMU 等感測模擬與合成資料管線。"
        "該綜述亦強調 Isaac Lab 可支援大規模並行強化學習；Mittal 等（2023）提出之 Orbit 即為此系譜前身。"
        "本研究以 Isaac Sim 6.0 為主實驗平台，並保留與 Isaac Lab 後續延伸的相容性。"
    ),
    (
        "近年 Isaac Sim 持續用於動態場景建置（GRADE, 2025）、Sim-to-Real 策略遷移（Salimpour 等, 2025）"
        "與工業操作驗證（Zhou 等, 2024）。然而 Gao 等（2026）亦未專論 RTX Acoustic；"
        "模擬器輸出仍非物理真值。Höfer 等（2021）強調應以任務級指標驗證遷移，"
        "不宜假設模擬信號與實機波形等價。本研究遵循此邊界："
        "RTX Acoustic 特徵僅作趨勢級距離推理之可行性證據，而非 CH201 實機波形對照標準。"
    ),
]

SECTION_23_HEADER = "2.3 室內環境與距離特徵（簡述）"
SECTION_23 = [
    (
        "封閉工業環境中，超音波回波會受牆面與工件反射影響，距離資訊往往不是單一路徑量測。"
        "Liu 等（2020）指出，室內環境下早期反射能量可作距離之弱趨勢指標，但仍受房間幾何與材質制約。"
        "本研究因此不宣稱厘米級測距，而以早期能量等摘要特徵檢驗「距離趨勢是否可用於感測回授」；"
        "特徵定義與擷取流程見第三章。"
    ),
]

SECTION_24 = [
    (
        "Isaac Sim 6.0 提供 RTX Acoustic 超音波感測模組（NVIDIA, 2026a），屬實驗性功能："
        "在模擬場景中以 GPU 產生回波資料（signal-way 格式），不輸出點雲。"
        "本研究把它視為「可控幾何下的合成超音波觀測」，用來檢驗感測回授接近與離線狀態判斷是否可行。"
    ),
    (
        "官方文件說明，輸出為發射端、接收端、通道與振幅取樣值的組合（NVIDIA, 2026a）；"
        "本研究據此整理早期能量、峰值與雙接收端平衡等特徵，程式細節列於附錄。"
        "Gao 等（2026）平台綜述未涵蓋此模組；學術上相近案例為 Song 等（2025）OceanSim，"
        "顯示 Isaac Sim 可延伸為專用 ray-traced 感測管線。"
    ),
    (
        "RTX Acoustic 不保證與實機 CH201 波形一致。本文所有結論均以任務級、趨勢級指標表述，"
        "不宣稱部署級測距精度。"
    ),
]

SECTION_27_TAIL = (
    "綜合前述各面向，現有文獻多覆蓋子集而非完整交集："
    "Gao 等（2026）雖系統整理 Isaac Sim 平台，但未涵蓋 RTX Acoustic 與工業手臂感測回授接近；"
    "聲學機器人研究多未提供完整的幾何、材質與任務設定紀錄；"
    "VLM 操作研究則少見以 RTX 超音波作 last-meter 非視覺感測回授。"
    "在 UR10e 工業手臂與腕部超音波感測場景下，同時整合上述要素之公開工作仍有限。"
)


def remove_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)


def insert_paragraph_after(paragraph: Paragraph, text: str, style: str) -> Paragraph:
    new_p = deepcopy(paragraph._element)
    paragraph._element.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.style = style
    if new_para.runs:
        for run in new_para.runs:
            run._element.getparent().remove(run._element)
    new_para.add_run(text)
    return new_para


def find_header_index(doc: Document, prefix: str) -> int:
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t.startswith(prefix) and p.style.name == "Header2":
            return i
    raise RuntimeError(f"Header not found: {prefix}")


def replace_section(
    doc: Document,
    header_prefix: str,
    next_header_prefix: str,
    new_header: str | None,
    bodies: list[str],
) -> None:
    start = find_header_index(doc, header_prefix)
    end = find_header_index(doc, next_header_prefix)
    header_para = doc.paragraphs[start]
    if new_header:
        header_para.text = new_header
    for idx in range(end - 1, start, -1):
        remove_paragraph(doc.paragraphs[idx])
    prev = header_para
    for text in bodies:
        prev = insert_paragraph_after(prev, text, "Content")


def fix_toc_duplicates(doc: Document) -> None:
    for p in doc.paragraphs:
        if p.style.name.startswith("toc") and "2.5" in p.text and p.text.count("2.5") > 1:
            p.text = re.sub(
                r"2\.5.*",
                "2.5 感測回授式接近與視覺語義操作對照11",
                p.text,
            )
        if p.style.name.startswith("toc") and p.text.startswith("2.3 聲學模擬"):
            p.text = p.text.replace("2.3 聲學模擬與多徑效應", "2.3 室內環境與距離特徵（簡述）")


def patch_section_27(doc: Document) -> None:
    for p in doc.paragraphs:
        if p.style.name == "Content" and p.text.strip().startswith("綜合前述各面向，現有文獻多覆蓋子集"):
            p.text = SECTION_27_TAIL
            return


def update_doc(path: Path) -> None:
    doc = Document(str(path))
    replace_section(doc, "2.2", "2.3", None, SECTION_22)
    replace_section(doc, "2.3", "2.4", SECTION_23_HEADER, SECTION_23)
    replace_section(doc, "2.4", "2.5", None, SECTION_24)
    patch_section_27(doc)
    fix_toc_duplicates(doc)
    doc.save(str(path))
    print(f"Updated {path}")


def main() -> None:
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    for out in OUT_PATHS:
        if out == SRC:
            update_doc(out)
        else:
            import shutil

            shutil.copy2(SRC, out)
            update_doc(out)


if __name__ == "__main__":
    main()