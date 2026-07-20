#!/usr/bin/env python3
"""Insert missing method tables + schematic figures into THESIS_DRAFT_FCU_v2.docx.

Idempotent: skips if caption already present.
Backup: .bak_before_tables_schematics
"""
from __future__ import annotations

import copy
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
DOCX = ROOT / "THESIS_DRAFT_FCU_v2.docx"
SCH = ROOT / "figures" / "schematic"
BACKUP = ROOT / "THESIS_DRAFT_FCU_v2.docx.bak_before_tables_schematics"


def find_para(doc: Document, prefix: str, max_len: int = 120):
    import re

    candidates = []
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if not t.startswith(prefix):
            continue
        if len(t) > max_len:
            continue
        # TOC lines often glue page number at end: "...結構29"
        if re.search(r"\d{1,3}$", t) and not t.endswith("cm") and "r＝" not in t:
            continue
        candidates.append((i, p, t))
    if not candidates:
        # fallback: first startswith
        for i, p in enumerate(doc.paragraphs):
            if p.text.strip().startswith(prefix):
                return i, p
        raise KeyError(prefix)
    # prefer shorter heading-like (section titles)
    candidates.sort(key=lambda x: (len(x[2]), x[0]))
    return candidates[0][0], candidates[0][1]


def insert_paragraph_after(paragraph, text: str = "", *, bold: bool = False, center: bool = False):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    # wrap as paragraph
    from docx.text.paragraph import Paragraph

    para = Paragraph(new_p, paragraph._parent)
    if text:
        run = para.add_run(text)
        run.font.size = Pt(11)
        run.font.name = "Times New Roman"
        r = run._element
        rPr = r.get_or_add_rPr()
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:eastAsia"), "標楷體")
        rPr.insert(0, rFonts)
        if bold:
            run.bold = True
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # copy style from content if possible
    try:
        para.style = paragraph.style
    except Exception:
        pass
    return para


def insert_picture_after(paragraph, image_path: Path, width_cm: float = 14.5):
    cap_holder = insert_paragraph_after(paragraph, "", center=True)
    run = cap_holder.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    cap_holder.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return cap_holder


def set_cell_shading(cell, fill_hex: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def add_table_after(paragraph, headers: list[str], rows: list[list[str]], col_widths_cm: list[float]):
    # Walk up to Document
    body = paragraph._parent
    while not hasattr(body, "add_table"):
        body = body._parent
    total_w = Cm(sum(col_widths_cm))
    table = body.add_table(rows=1 + len(rows), cols=len(headers), width=total_w)
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
        set_cell_shading(cell, "D9E2F3")
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = val
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    # widths
    for row in table.rows:
        for j, w in enumerate(col_widths_cm):
            row.cells[j].width = Cm(w)
    # move table XML after paragraph
    tbl = table._tbl
    paragraph._p.addnext(tbl)
    return table


def has_caption(doc: Document, caption: str) -> bool:
    return any(caption in p.text for p in doc.paragraphs)


def insert_block_after_anchor(
    doc: Document,
    anchor_prefix: str,
    after_contains: str | None,
    *,
    intro: str | None,
    image: Path | None,
    caption: str | None,
    table: tuple | None = None,
):
    """Find section; optionally find paragraph containing after_contains within next N paras."""
    if caption and has_caption(doc, caption):
        print(f"SKIP exists: {caption}")
        return
    idx, p = find_para(doc, anchor_prefix)
    # walk forward to after_contains
    target = p
    if after_contains:
        for j in range(idx, min(idx + 25, len(doc.paragraphs))):
            if after_contains in doc.paragraphs[j].text:
                target = doc.paragraphs[j]
                break
    # insert from bottom to top so order preserves: intro, image, caption, table caption, table
    # Actually insert after target sequentially using returned para as next anchor
    cur = target
    if intro:
        cur = insert_paragraph_after(cur, intro)
    if image is not None and image.is_file():
        cur = insert_picture_after(cur, image)
    if caption:
        cur = insert_paragraph_after(cur, caption, bold=True, center=True)
    if table is not None:
        t_cap, headers, rows, widths = table
        if not has_caption(doc, t_cap):
            cur = insert_paragraph_after(cur, t_cap, bold=True, center=True)
            add_table_after(cur, headers, rows, widths)
            print(f"  + table {t_cap}")
    if caption:
        print(f"  + {caption}")


def main():
    if not DOCX.is_file():
        raise SystemExit(f"missing {DOCX}")
    shutil.copy2(DOCX, BACKUP)
    doc = Document(str(DOCX))

    # ── 3.1 GMO ─────────────────────────────────────────────
    insert_block_after_anchor(
        doc,
        "3.1 WPM",
        "timeOffsetNs",
        intro=(
            "表 3.1 整理 GMO 在超音波語意下的欄位含義（與 LiDAR 點雲語意不同）；"
            "圖 3.1 示意 signal way 緩衝區排列與正確重建方式。距離換算採當輪自校，"
            "不以 timeOffsetNs 為準（Isaac Sim 6.0 實測恆為 0）。"
        ),
        image=SCH / "fig_gmo_structure.png",
        caption="圖 3.1  GMO 欄位語意與 signal way 排列示意",
        table=(
            "表 3.1  GMO 主要欄位在超音波下之含義（Isaac Sim 6.0／RTX Acoustic）",
            ["欄位", "LiDAR 常見語意", "本研究超音波語意", "備註"],
            [
                ["x[i]", "空間 X", "TX 掛載 ID", "非公尺座標"],
                ["y[i]", "空間 Y", "RX 掛載 ID", "非公尺座標"],
                ["z[i]", "空間 Z", "通道 ID", "不可當時間軸"],
                ["scalar[i]", "強度等", "振幅樣本", "組成波形"],
                ["timeOffsetNs[i]", "到達時間", "實測恆 0", "不可算 ToF"],
                ["numSamplesPerSgw", "—", "每條 way 樣本數", "重建 stride"],
                ["numElements", "點數", "總樣本數", "檢查有效輸出"],
            ],
            [2.2, 3.2, 4.0, 3.0],
        ),
    )

    # ── 3.3 pillars + three arm + paired + lit table ───────
    insert_block_after_anchor(
        doc,
        "3.3 實驗方法學",
        "支柱四",
        intro=(
            "圖 3.2 總覽四支柱；圖 3.3 對應支柱二之三臂消融；"
            "圖 3.4 為配對移除（支柱一）之操作流程。表 3.2 對照方法族與文獻譜系"
            "（組合落地為本研究，分項非憑空設計）。"
        ),
        image=SCH / "fig_four_pillars.png",
        caption="圖 3.2  實驗方法學四支柱示意",
        table=(
            "表 3.2  方法／名詞與文獻譜系對照（摘要）",
            ["方法／名詞", "本研究角色", "文獻譜系（文中標號族）"],
            [
                ["包絡優先", "先測可偵測幾何再放任務", "Liu 等；動機 Valin／Tsuchiya"],
                ["配對移除", "有／無目標分離回波", "Valin 等；Tsuchiya 等；Liu 等"],
                ["三臂消融", "聲學／盲走／開環歸因", "Meyes 等（消融）"],
                ["預註冊判準", "執行前鎖定通過標準", "Nosek 等"],
                ["peak／ToF 估距", "peak 索引→距離自校", "Zhmud 等；He 等"],
                ["多點定位", "五視點交會恢復側向", "Kapoor 等；Hayes 等"],
                ["GMO／WPM", "引擎聲學輸出容器", "NVIDIA RTX Acoustic／GMO 文件"],
                ["weld-on-stall 升舉", "接觸觸發附著（非摩擦宣稱）", "模擬器限制下工程邊界（6.2）"],
            ],
            [3.0, 4.5, 7.0],
        ),
    )
    # more figures after 3.3 table block — anchor to 圖 3.2 caption if exists
    insert_block_after_anchor(
        doc,
        "圖 3.2",
        None,
        intro=None,
        image=SCH / "fig_three_arm_ablation.png",
        caption="圖 3.3  三臂資訊消融對照（closed／blind／open）",
        table=None,
    )
    insert_block_after_anchor(
        doc,
        "圖 3.3",
        None,
        intro=None,
        image=SCH / "fig_paired_removal.png",
        caption="圖 3.4  配對移除量測流程示意",
        table=None,
    )

    # ── 3.5 / 3.6 control pipeline + experiment chain ──────
    insert_block_after_anchor(
        doc,
        "3.5 聲學特徵",
        "可歸因",
        intro="圖 3.5 將 GMO 波形至控制決策的管線示意如下。",
        image=SCH / "fig_acoustic_range_pipeline.png",
        caption="圖 3.5  聲學距離估測與控制管線示意",
        table=None,
    )
    insert_block_after_anchor(
        doc,
        "3.6 章節間",
        "資料鏈",
        intro="圖 3.6 對應全文章節間實驗鏈（S1→…→D4）。",
        image=SCH / "fig_pipeline_s1_to_d4.png",
        caption="圖 3.6  實驗鏈：感測包絡→閉環接近→夾取整合→策略串聯",
        table=None,
    )

    # ── 5.3 D3 sequence + D4 ───────────────────────────────
    insert_block_after_anchor(
        doc,
        "5.3 D3",
        "夾持機制如實說明",
        intro="圖 5.4 將端到端夾取序列以示意圖整理（非 Isaac 視窗截圖）。",
        image=SCH / "fig_d3_grasp_sequence.png",
        caption="圖 5.4  D3 端到端夾取序列示意（聲學停靠→一次推算→附著升舉）",
        table=None,
    )
    # D4 paragraph (long body sentence — match substring)
    if not has_caption(doc, "圖 5.5"):
        target = None
        for p in doc.paragraphs:
            if "代號 D4" in p.text and "不是要改寫" in p.text:
                target = p
                break
        if target is None:
            for p in doc.paragraphs:
                if "代號 D4" in p.text:
                    target = p
                    break
        if target is None:
            raise KeyError("D4 body paragraph not found")
        cur = insert_paragraph_after(
            target,
            "圖 5.5、圖 5.6 分別為 D4 雙軌架構與同場景串聯結果示意。",
        )
        cur = insert_picture_after(cur, SCH / "fig_d4_dual_track.png")
        cur = insert_paragraph_after(
            cur, "圖 5.5  D4 雙軌：規則狀態機（A）與策略接近／合爪（B）", bold=True, center=True
        )
        cur = insert_picture_after(cur, SCH / "fig_same_scene_policy_n90.png")
        insert_paragraph_after(
            cur, "圖 5.6  同場景策略串聯結果示意（n=90）", bold=True, center=True
        )
        print("  + 圖 5.5 / 5.6 D4")
    else:
        print("SKIP 圖 5.5")

    # ── 5.4 multilateration ────────────────────────────────
    insert_block_after_anchor(
        doc,
        "5.4 D2",
        "多點定位",
        intro="圖 5.7 示意五視點測距交會以恢復側向（文獻譜系：Kapoor 等；Hayes 等）。",
        image=SCH / "fig_multilateration.png",
        caption="圖 5.7  二維多點定位（運動合成基線）示意",
        table=None,
    )

    # ── 6.2 claim boundary ─────────────────────────────────
    insert_block_after_anchor(
        doc,
        "6.2 推論範圍",
        "推論不及於",
        intro="圖 6.1 將可支持與不支持之結論並列，便於口試快速對照。",
        image=SCH / "fig_claim_boundary.png",
        caption="圖 6.1  推論範圍：可支持與不支持之結論示意",
        table=None,
    )

    doc.save(str(DOCX))
    print(f"saved {DOCX}")
    print(f"backup {BACKUP}")


if __name__ == "__main__":
    main()
