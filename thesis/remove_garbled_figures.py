#!/usr/bin/env python3
"""Remove Chinese-labeled matplotlib figures that render as tofu boxes in Word."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph

PATHS = [
    Path(__file__).resolve().parent / "THESIS_DRAFT_FCU_v1.docx",
    Path(__file__).resolve().parent / "2THESIS_DRAFT_FCU_v1.docx",
    Path("/home/lab109/下載/2THESIS_DRAFT_FCU_v1_no_pyroom.docx"),
]

REMOVE_CAPTIONS = {
    "圖1  本研究歷程與系統演進",
    "圖2  模擬場景示意（手臂、超音波感測與目標工件）",
    "圖5  感測回授與對照組之階段成功率比較",
    "圖6  停止前進區域判斷：不同特徵組合之比較",
}


def has_image(paragraph: Paragraph) -> bool:
    return any(run._element.xpath(".//a:blip") for run in paragraph.runs)


def remove_paragraph(paragraph: Paragraph) -> None:
    paragraph._element.getparent().remove(paragraph._element)


def strip_figures(path: Path) -> None:
    doc = Document(str(path))
    to_remove: list[Paragraph] = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t in REMOVE_CAPTIONS:
            prev = para._element.getprevious()
            if prev is not None:
                prev_para = Paragraph(prev, para._parent)
                if has_image(prev_para):
                    to_remove.append(prev_para)
            to_remove.append(para)

    for para in to_remove:
        remove_paragraph(para)

    doc.save(str(path))
    imgs = sum(1 for p in doc.paragraphs for r in p.runs if r._element.xpath(".//a:blip"))
    print(f"{path.name}: removed {len(to_remove)} blocks, images left={imgs}")


def main() -> None:
    for path in PATHS:
        if path.exists():
            strip_figures(path)


if __name__ == "__main__":
    main()